#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试翻译遗漏问题的脚本
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import logging

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def analyze_document_structure(docx_path: str):
    """分析文档结构，找出可能遗漏翻译的内容"""
    try:
        doc = Document(docx_path)
        
        print("=" * 80)
        print(f"文档结构分析: {docx_path}")
        print("=" * 80)
        
        # 分析段落
        print(f"\n📄 段落分析 (共 {len(doc.paragraphs)} 个段落):")
        print("-" * 60)
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                print(f"段落 {i+1:3d}: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        # 分析表格
        print(f"\n📊 表格分析 (共 {len(doc.tables)} 个表格):")
        print("-" * 60)
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                print(f"  行 {row_idx + 1}:")
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        print(f"    单元格 {cell_idx + 1}: {text[:50]}{'...' if len(text) > 50 else ''}")
        
        # 检查文档末尾的内容
        print(f"\n🔍 文档末尾内容检查:")
        print("-" * 60)
        
        # 获取最后10个段落
        last_paragraphs = doc.paragraphs[-10:]
        print(f"最后10个段落:")
        for i, para in enumerate(last_paragraphs):
            text = para.text.strip()
            if text:
                para_num = len(doc.paragraphs) - 10 + i + 1
                print(f"  段落 {para_num}: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        # 检查最后一个表格的内容
        if doc.tables:
            last_table = doc.tables[-1]
            print(f"\n最后一个表格的详细内容:")
            for row_idx, row in enumerate(last_table.rows):
                print(f"  行 {row_idx + 1}:")
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    print(f"    单元格 {cell_idx + 1}: '{text}'")
        
        # 分析可能的翻译对
        print(f"\n🔄 可能的翻译对分析:")
        print("-" * 60)
        
        # 检查段落中的翻译对
        for i in range(len(doc.paragraphs) - 1):
            para1 = doc.paragraphs[i]
            para2 = doc.paragraphs[i + 1]
            text1 = para1.text.strip()
            text2 = para2.text.strip()
            
            if text1 and text2:
                # 简单的翻译对检测
                has_chinese1 = any('\u4e00' <= char <= '\u9fff' for char in text1)
                has_english1 = any('a' <= char.lower() <= 'z' for char in text1)
                has_chinese2 = any('\u4e00' <= char <= '\u9fff' for char in text2)
                has_english2 = any('a' <= char.lower() <= 'z' for char in text2)
                
                if (has_chinese1 and not has_chinese2 and has_english2) or (has_english1 and not has_english2 and has_chinese2):
                    print(f"  可能的翻译对 (段落 {i+1} & {i+2}):")
                    print(f"    第一行: {text1[:80]}...")
                    print(f"    第二行: {text2[:80]}...")
                    print()
        
        # 检查表格中的翻译对
        for table_idx, table in enumerate(doc.tables):
            print(f"表格 {table_idx + 1} 中的可能翻译对:")
            for row_idx, row in enumerate(table.rows):
                cells = row.cells
                for cell_idx in range(len(cells) - 1):
                    cell1 = cells[cell_idx]
                    cell2 = cells[cell_idx + 1]
                    text1 = cell1.text.strip()
                    text2 = cell2.text.strip()
                    
                    if text1 and text2:
                        # 简单的翻译对检测
                        has_chinese1 = any('\u4e00' <= char <= '\u9fff' for char in text1)
                        has_english1 = any('a' <= char.lower() <= 'z' for char in text1)
                        has_chinese2 = any('\u4e00' <= char <= '\u9fff' for char in text2)
                        has_english2 = any('a' <= char.lower() <= 'z' for char in text2)
                        
                        if (has_chinese1 and not has_chinese2 and has_english2) or (has_english1 and not has_english2 and has_chinese2):
                            print(f"    行 {row_idx + 1}, 单元格 {cell_idx + 1} & {cell_idx + 2}:")
                            print(f"      第一个: {text1[:50]}...")
                            print(f"      第二个: {text2[:50]}...")
        
        print("\n" + "=" * 80)
        print("分析完成")
        print("=" * 80)
        
    except Exception as e:
        logger.error(f"分析文档时出错: {str(e)}")
        raise

def find_untranslated_content(docx_path: str):
    """查找可能未翻译的内容"""
    try:
        doc = Document(docx_path)
        
        print("\n🔍 查找可能未翻译的内容:")
        print("-" * 60)
        
        untranslated_paragraphs = []
        untranslated_cells = []
        
        # 检查段落
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # 检查是否只包含中文（可能需要翻译）
                has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
                has_english = any('a' <= char.lower() <= 'z' for char in text)
                
                if has_chinese and not has_english:
                    untranslated_paragraphs.append((i + 1, text))
        
        # 检查表格单元格
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        # 检查是否只包含中文（可能需要翻译）
                        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in text)
                        has_english = any('a' <= char.lower() <= 'z' for char in text)
                        
                        if has_chinese and not has_english:
                            untranslated_cells.append((table_idx + 1, row_idx + 1, cell_idx + 1, text))
        
        print(f"可能未翻译的段落 ({len(untranslated_paragraphs)} 个):")
        for para_num, text in untranslated_paragraphs:
            print(f"  段落 {para_num}: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        print(f"\n可能未翻译的表格单元格 ({len(untranslated_cells)} 个):")
        for table_num, row_num, cell_num, text in untranslated_cells:
            print(f"  表格 {table_num}, 行 {row_num}, 单元格 {cell_num}: {text[:50]}{'...' if len(text) > 50 else ''}")
        
    except Exception as e:
        logger.error(f"查找未翻译内容时出错: {str(e)}")
        raise

if __name__ == "__main__":
    # 分析最近翻译的文档
    import glob
    
    # 查找最新的翻译文档
    output_dir = "输出"
    if os.path.exists(output_dir):
        docx_files = glob.glob(os.path.join(output_dir, "*_带翻译_*.docx"))
        if docx_files:
            # 按修改时间排序，获取最新的文件
            latest_file = max(docx_files, key=os.path.getmtime)
            print(f"分析最新的翻译文档: {latest_file}")
            
            analyze_document_structure(latest_file)
            find_untranslated_content(latest_file)
        else:
            print("未找到翻译文档")
    else:
        print("输出目录不存在")
