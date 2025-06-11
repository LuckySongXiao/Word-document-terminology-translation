#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复表格翻译遗漏问题
"""

import os
import sys
from docx import Document
import logging
from services.translation_detector import TranslationDetector
from services.zhipuai_translator import ZhipuAITranslator
from services.siliconflow_translator import SiliconFlowTranslator
from services.ollama_translator import OllamaTranslator
import json

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def load_api_config():
    """加载API配置"""
    try:
        with open('API_config/zhipu_api.json', 'r', encoding='utf-8') as f:
            zhipu_config = json.load(f)
        return zhipu_config
    except Exception as e:
        logger.error(f"加载API配置失败: {e}")
        return None

def get_translator():
    """获取翻译器"""
    # 尝试使用智谱AI
    config = load_api_config()
    if config and config.get('api_key'):
        try:
            translator = ZhipuAITranslator(config['api_key'])
            # 测试连接
            test_result = translator.translate("测试", {}, "zh", "en")
            if test_result and test_result != "测试":
                logger.info("使用智谱AI翻译器")
                return translator
        except Exception as e:
            logger.warning(f"智谱AI翻译器初始化失败: {e}")

    # 回退到本地Ollama
    try:
        # 从配置文件加载Ollama配置
        with open('config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)

        ollama_config = config.get("fallback_translator", {})
        model = ollama_config.get("model", "qwen2.5:7b")
        api_url = ollama_config.get("api_url", "http://localhost:11434")

        translator = OllamaTranslator(
            model=model,
            api_url=api_url,
            model_list_timeout=10,
            translate_timeout=60
        )
        logger.info("使用本地Ollama翻译器")
        return translator
    except Exception as e:
        logger.error(f"Ollama翻译器初始化失败: {e}")
        return None

def fix_document_translation():
    """修复文档翻译"""
    
    # 文档路径
    output_dir = "输出"
    doc_name = "单晶电阻率管控技术标准_带翻译_20250610201907.docx"
    doc_path = os.path.join(output_dir, doc_name)
    
    if not os.path.exists(doc_path):
        logger.error(f"文档不存在: {doc_path}")
        return False
    
    # 获取翻译器
    translator = get_translator()
    if not translator:
        logger.error("无法初始化翻译器")
        return False
    
    try:
        # 打开文档
        doc = Document(doc_path)
        logger.info(f"打开文档: {doc_path}")
        
        # 找到表格4（最后一个表格）
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        table = doc.tables[3]  # 第4个表格（索引3）
        logger.info(f"处理表格4，行数: {len(table.rows)}")
        
        # 找到最后一行
        last_row = table.rows[-1]
        logger.info(f"最后一行单元格数: {len(last_row.cells)}")
        
        # 检查并修复最后两个单元格
        cells_to_fix = []
        if len(last_row.cells) >= 2:
            # 倒数第二个单元格
            second_last_cell = last_row.cells[-2]
            last_cell = last_row.cells[-1]
            
            cells_to_fix.append(("倒数第二个", second_last_cell))
            cells_to_fix.append(("最后一个", last_cell))
        
        # 修复每个单元格
        for cell_name, cell in cells_to_fix:
            cell_text = cell.text.strip()
            if not cell_text:
                logger.info(f"{cell_name}单元格为空，跳过")
                continue
            
            logger.info(f"检查{cell_name}单元格: {cell_text[:50]}...")
            
            # 检查是否需要翻译（包含中文但没有完整英文翻译）
            import re
            has_chinese = bool(re.search(r'[\u4e00-\u9fff]', cell_text))
            
            # 更严格的英文检测：至少要有完整的英文单词
            english_words = re.findall(r'\b[A-Za-z]{2,}\b', cell_text)
            has_meaningful_english = len(english_words) >= 3  # 至少3个英文单词才算有意义的英文内容
            
            logger.info(f"{cell_name}单元格 - 包含中文: {has_chinese}, 有意义的英文单词: {len(english_words)}")
            
            if has_chinese and not has_meaningful_english:
                logger.info(f"{cell_name}单元格需要翻译")
                
                # 翻译内容
                try:
                    logger.info(f"正在翻译{cell_name}单元格...")
                    translated_text = translator.translate(cell_text, {}, "zh", "en")

                    if translated_text and translated_text != cell_text:
                        logger.info(f"翻译成功: {translated_text[:100]}...")
                        
                        # 添加翻译到单元格（双语对照格式）
                        # 清空现有内容
                        for para in cell.paragraphs:
                            para.clear()
                        
                        # 添加原文
                        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        para.add_run(cell_text)
                        
                        # 添加换行和翻译
                        para.add_run("\n")
                        para.add_run(translated_text)
                        
                        logger.info(f"{cell_name}单元格翻译完成")
                    else:
                        logger.warning(f"{cell_name}单元格翻译失败或无变化")
                        
                except Exception as e:
                    logger.error(f"翻译{cell_name}单元格时出错: {e}")
            else:
                logger.info(f"{cell_name}单元格不需要翻译")
        
        # 保存修复后的文档
        fixed_doc_name = doc_name.replace(".docx", "_修复翻译.docx")
        fixed_doc_path = os.path.join(output_dir, fixed_doc_name)
        doc.save(fixed_doc_path)
        logger.info(f"修复后的文档已保存: {fixed_doc_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"修复文档时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    logger.info("开始修复表格翻译遗漏问题...")
    
    success = fix_document_translation()
    
    if success:
        logger.info("✅ 修复完成！")
        logger.info("请检查输出文件夹中的修复后文档")
    else:
        logger.error("❌ 修复失败！")

if __name__ == "__main__":
    main()
