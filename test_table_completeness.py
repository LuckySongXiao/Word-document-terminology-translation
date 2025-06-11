#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试表格处理完整性的脚本
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import logging

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def analyze_table_structure(docx_path: str):
    """分析Word文档中表格的详细结构"""
    try:
        doc = Document(docx_path)
        
        print("=" * 80)
        print(f"表格结构详细分析: {docx_path}")
        print("=" * 80)
        
        if not doc.tables:
            print("❌ 文档中没有表格")
            return
        
        print(f"📊 文档包含 {len(doc.tables)} 个表格")
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n{'='*60}")
            print(f"表格 {table_idx + 1}")
            print(f"{'='*60}")
            
            # 统计表格信息
            total_rows = len(table.rows)
            total_cells = 0
            non_empty_cells = 0
            chinese_cells = 0
            english_cells = 0
            mixed_cells = 0
            number_cells = 0
            
            print(f"总行数: {total_rows}")
            
            # 逐行逐列分析
            for row_idx, row in enumerate(table.rows):
                print(f"\n第 {row_idx + 1} 行 (共 {len(row.cells)} 个单元格):")
                
                for cell_idx, cell in enumerate(row.cells):
                    total_cells += 1
                    cell_text = cell.text.strip()
                    
                    if cell_text:
                        non_empty_cells += 1
                        
                        # 分析单元格内容类型
                        has_chinese = any('\u4e00' <= char <= '\u9fff' for char in cell_text)
                        has_english = any('a' <= char.lower() <= 'z' for char in cell_text)
                        has_numbers = any(char.isdigit() for char in cell_text)
                        
                        if has_chinese and has_english:
                            content_type = "混合"
                            mixed_cells += 1
                        elif has_chinese:
                            content_type = "中文"
                            chinese_cells += 1
                        elif has_english:
                            content_type = "英文"
                            english_cells += 1
                        elif has_numbers and not has_chinese and not has_english:
                            content_type = "数字"
                            number_cells += 1
                        else:
                            content_type = "其他"
                        
                        print(f"  单元格 [{row_idx + 1}, {cell_idx + 1}] ({content_type}): {cell_text[:50]}{'...' if len(cell_text) > 50 else ''}")
                    else:
                        print(f"  单元格 [{row_idx + 1}, {cell_idx + 1}] (空): <空单元格>")
            
            # 表格统计信息
            print(f"\n表格 {table_idx + 1} 统计:")
            print(f"  总单元格数: {total_cells}")
            print(f"  非空单元格数: {non_empty_cells}")
            print(f"  中文单元格数: {chinese_cells}")
            print(f"  英文单元格数: {english_cells}")
            print(f"  混合单元格数: {mixed_cells}")
            print(f"  数字单元格数: {number_cells}")
            print(f"  空单元格数: {total_cells - non_empty_cells}")
        
        print(f"\n{'='*80}")
        print("分析完成")
        print("=" * 80)
        
    except Exception as e:
        logger.error(f"分析表格结构时出错: {str(e)}")
        raise

def simulate_table_processing(docx_path: str):
    """模拟表格处理过程，检查是否有遗漏"""
    try:
        doc = Document(docx_path)
        
        print("\n" + "=" * 80)
        print("模拟表格处理过程")
        print("=" * 80)
        
        if not doc.tables:
            print("❌ 文档中没有表格")
            return
        
        for table_idx, table in enumerate(doc.tables):
            print(f"\n🔄 模拟处理表格 {table_idx + 1}")
            print("-" * 60)
            
            # 模拟双行检测处理
            processed_cells = set()
            translation_pairs = []
            skip_cells = []
            translate_cells = []
            
            for row_idx, row in enumerate(table.rows):
                cells = row.cells
                cell_idx = 0
                
                while cell_idx < len(cells):
                    current_cell = cells[cell_idx]
                    current_text = current_cell.text.strip()
                    cell_position = (table_idx, row_idx, cell_idx)
                    
                    if not current_text:
                        cell_idx += 1
                        continue
                    
                    if cell_position in processed_cells:
                        cell_idx += 1
                        continue
                    
                    # 检查与右侧单元格的翻译对
                    translation_pair_found = False
                    if cell_idx + 1 < len(cells):
                        next_cell = cells[cell_idx + 1]
                        next_text = next_cell.text.strip()
                        next_position = (table_idx, row_idx, cell_idx + 1)
                        
                        if next_text and next_position not in processed_cells:
                            # 简单的翻译对检测
                            has_chinese1 = any('\u4e00' <= char <= '\u9fff' for char in current_text)
                            has_english1 = any('a' <= char.lower() <= 'z' for char in current_text)
                            has_chinese2 = any('\u4e00' <= char <= '\u9fff' for char in next_text)
                            has_english2 = any('a' <= char.lower() <= 'z' for char in next_text)
                            
                            if (has_chinese1 and not has_chinese2 and has_english2) or (has_english1 and not has_english2 and has_chinese2):
                                translation_pairs.append((
                                    f"[{row_idx + 1}, {cell_idx + 1}]",
                                    f"[{row_idx + 1}, {cell_idx + 2}]",
                                    current_text[:30],
                                    next_text[:30]
                                ))
                                processed_cells.add(cell_position)
                                processed_cells.add(next_position)
                                translation_pair_found = True
                                cell_idx += 2
                                continue
                    
                    # 检查是否应该跳过
                    if not translation_pair_found:
                        # 简单的跳过检测（纯数字）
                        if current_text.replace(' ', '').replace('.', '').replace('-', '').replace('>', '').replace('<', '').replace('=', '').replace('μ', '').replace('s', '').replace('Ω', '').replace('cm', '').replace('ohm', '').isdigit():
                            skip_cells.append((f"[{row_idx + 1}, {cell_idx + 1}]", current_text[:30], "纯数字/代码"))
                        else:
                            # 检查是否需要翻译
                            has_chinese = any('\u4e00' <= char <= '\u9fff' for char in current_text)
                            if has_chinese:
                                translate_cells.append((f"[{row_idx + 1}, {cell_idx + 1}]", current_text[:30]))
                            else:
                                skip_cells.append((f"[{row_idx + 1}, {cell_idx + 1}]", current_text[:30], "非中文内容"))
                        
                        processed_cells.add(cell_position)
                        cell_idx += 1
            
            # 输出处理结果
            print(f"✅ 检测到的翻译对 ({len(translation_pairs)} 对):")
            for pos1, pos2, text1, text2 in translation_pairs:
                print(f"  {pos1} & {pos2}: '{text1}' ↔ '{text2}'")
            
            print(f"\n⏭️ 跳过的单元格 ({len(skip_cells)} 个):")
            for pos, text, reason in skip_cells:
                print(f"  {pos}: '{text}' (原因: {reason})")
            
            print(f"\n🔄 需要翻译的单元格 ({len(translate_cells)} 个):")
            for pos, text in translate_cells:
                print(f"  {pos}: '{text}'")
            
            # 验证完整性
            total_non_empty = sum(1 for row in table.rows for cell in row.cells if cell.text.strip())
            total_processed = len(translation_pairs) * 2 + len(skip_cells) + len(translate_cells)
            
            print(f"\n📊 处理统计:")
            print(f"  总非空单元格: {total_non_empty}")
            print(f"  已处理单元格: {total_processed}")
            print(f"  处理完整性: {'✅ 完整' if total_processed == total_non_empty else '❌ 有遗漏'}")
            
            if total_processed != total_non_empty:
                print(f"  ⚠️ 遗漏了 {total_non_empty - total_processed} 个单元格")
        
    except Exception as e:
        logger.error(f"模拟表格处理时出错: {str(e)}")
        raise

if __name__ == "__main__":
    # 分析测试文档
    test_file = "测试表格翻译完整性.docx"

    if os.path.exists(test_file):
        print(f"分析测试文档: {test_file}")
        analyze_table_structure(test_file)
        simulate_table_processing(test_file)
    else:
        print(f"测试文档不存在: {test_file}")

        # 查找最近翻译的文档作为备选
        import glob
        output_dir = "输出"
        if os.path.exists(output_dir):
            docx_files = glob.glob(os.path.join(output_dir, "*_带翻译_*.docx"))
            if docx_files:
                # 按修改时间排序，获取最新的文件
                latest_file = max(docx_files, key=os.path.getmtime)
                print(f"分析最新的翻译文档: {latest_file}")

                analyze_table_structure(latest_file)
                simulate_table_processing(latest_file)
            else:
                print("未找到翻译文档")
        else:
            print("输出目录不存在")
