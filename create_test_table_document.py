#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建包含复杂表格的测试Word文档
"""

from docx import Document
from docx.shared import Inches
import os

def create_test_document():
    """创建包含复杂表格的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('表格翻译完整性测试文档', 0)
    
    # 添加说明段落
    doc.add_paragraph('这是一个用于测试表格翻译完整性的文档，包含多种类型的表格内容。')
    
    # 创建第一个表格：简单的中英文对照表格
    doc.add_heading('表格1：中英文对照表格', level=1)
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    
    # 表头
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '中文'
    hdr_cells[1].text = 'English'
    
    # 添加数据行
    data1 = [
        ('产品名称', 'Product Name'),
        ('技术规格', 'Technical Specification'),
        ('质量标准', 'Quality Standard'),
        ('测试方法', 'Test Method'),
        ('验收标准', 'Acceptance Criteria')
    ]
    
    for chinese, english in data1:
        row_cells = table1.add_row().cells
        row_cells[0].text = chinese
        row_cells[1].text = english
    
    # 创建第二个表格：混合内容表格
    doc.add_heading('表格2：混合内容表格', level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    
    # 表头
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = 'Item'
    hdr_cells[2].text = '数值'
    hdr_cells[3].text = '单位'
    
    # 添加数据行
    data2 = [
        ('电阻', 'Resistance', '< 0.2Ω.cm', 'Ω.cm'),
        ('正常循环（A级）', 'Normal Cycle (Grade A)', '≥50μs', 'μs'),
        ('单晶管控使用（B级）', 'Single Crystal Control Usage (B Level)', '20 < x > 50μs', 'μs'),
        ('单晶管控试投（C级）', 'Single Crystal Control Pilot (C-Level)', '5 < x > 20μs', 'μs'),
        ('报废（D级）', 'Scrap (D-grade)', '1、整根少子寿命带：≤50μs部分报废；', ''),
        ('', '', '2、尾部少子寿命 < 5μs报废。', ''),
    ]
    
    for item_cn, item_en, value, unit in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = item_cn
        row_cells[1].text = item_en
        row_cells[2].text = value
        row_cells[3].text = unit
    
    # 创建第三个表格：需要翻译的纯中文表格
    doc.add_heading('表格3：需要翻译的纯中文表格', level=1)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = 'Table Grid'
    
    # 表头
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '检测项目'
    hdr_cells[2].text = '技术要求'
    
    # 添加数据行
    data3 = [
        ('1', '外观检查', '表面无明显缺陷'),
        ('2', '尺寸测量', '符合图纸要求'),
        ('3', '电气性能', '满足技术规范'),
        ('4', '机械强度', '通过振动测试'),
        ('5', '环境适应性', '耐高温低温循环')
    ]
    
    for seq, item, requirement in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = seq
        row_cells[1].text = item
        row_cells[2].text = requirement
    
    # 创建第四个表格：复杂的混合表格
    doc.add_heading('表格4：复杂混合表格', level=1)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = 'Table Grid'
    
    # 表头
    hdr_cells = table4.rows[0].cells
    hdr_cells[0].text = '参数类型'
    hdr_cells[1].text = '规格范围'
    hdr_cells[2].text = '备注说明'
    
    # 添加数据行
    data4 = [
        ('温度范围', '-40°C ~ +85°C', '工作环境温度'),
        ('湿度要求', '5% ~ 95% RH', '相对湿度，无凝露'),
        ('电源电压', 'DC 12V ± 10%', '直流电源供电'),
        ('功耗限制', '≤ 5W', '最大功耗不超过5瓦'),
        ('通信接口', 'RS485 / Ethernet', '支持多种通信方式'),
        ('防护等级', 'IP65', '防尘防水等级'),
        ('123', '456', '789'),  # 纯数字行
        ('ABC', 'DEF', 'GHI'),  # 纯字母行
    ]
    
    for param, spec, note in data4:
        row_cells = table4.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = spec
        row_cells[2].text = note
    
    # 添加结尾段落
    doc.add_paragraph('')
    doc.add_paragraph('备注：1、尾部接触面少子进行分类，圆棒按照A面少子进行分类；')
    doc.add_paragraph('2、晶裂部分少子统一按照5 < x > 20μs进行分类；')
    doc.add_paragraph('')
    doc.add_paragraph('这是文档末尾的一段需要翻译的中文内容。')
    
    # 保存文档
    output_path = '测试表格翻译完整性.docx'
    doc.save(output_path)
    print(f"测试文档已创建: {output_path}")
    return output_path

if __name__ == "__main__":
    create_test_document()
