#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试修复后的文档翻译
"""

import os
import sys
import json
import logging
from services.document_factory import DocumentProcessorFactory
from services.zhipuai_translator import ZhipuAITranslator
from services.ollama_translator import OllamaTranslator
from utils.terminology import load_terminology

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def load_api_config():
    """加载API配置"""
    try:
        with open('API_config/zhipu_api.json', 'r', encoding='utf-8') as f:
            zhipu_config = json.load(f)
        return zhipu_config
    except Exception as e:
        logger.error(f"加载API配置失败: {e}")
        return None

def get_translator():
    """获取翻译器"""
    # 尝试使用智谱AI
    config = load_api_config()
    if config and config.get('api_key'):
        try:
            translator = ZhipuAITranslator(config['api_key'])
            # 测试连接
            test_result = translator.translate("测试", {}, "zh", "en")
            if test_result and test_result != "测试":
                logger.info("使用智谱AI翻译器")
                return translator
        except Exception as e:
            logger.warning(f"智谱AI翻译器初始化失败: {e}")
    
    # 回退到本地Ollama
    try:
        # 从配置文件加载Ollama配置
        with open('config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        ollama_config = config.get("fallback_translator", {})
        model = ollama_config.get("model", "qwen2.5:7b")
        api_url = ollama_config.get("api_url", "http://localhost:11434")
        
        translator = OllamaTranslator(
            model=model,
            api_url=api_url,
            model_list_timeout=10,
            translate_timeout=60
        )
        logger.info("使用本地Ollama翻译器")
        return translator
    except Exception as e:
        logger.error(f"Ollama翻译器初始化失败: {e}")
        return None

def test_document_translation():
    """测试文档翻译"""
    
    # 输入文档路径
    input_file = "uploads/单晶电阻率管控技术标准.docx"
    
    if not os.path.exists(input_file):
        logger.error(f"输入文档不存在: {input_file}")
        return False
    
    # 获取翻译器
    translator = get_translator()
    if not translator:
        logger.error("无法初始化翻译器")
        return False
    
    # 加载术语库
    terminology = load_terminology()
    english_terms = terminology.get("英语", {})
    logger.info(f"加载术语库，包含 {len(english_terms)} 个英语术语")
    
    try:
        # 创建翻译服务
        from services.translator import TranslationService
        translation_service = TranslationService()

        # 设置翻译器类型
        if isinstance(translator, ZhipuAITranslator):
            translation_service.set_translator_type("zhipuai")
        elif isinstance(translator, OllamaTranslator):
            translation_service.set_translator_type("ollama")

        # 创建文档处理器
        processor = DocumentProcessorFactory.create_processor(
            file_path=input_file,
            translator=translation_service
        )

        if not processor:
            logger.error("无法创建文档处理器")
            return False

        logger.info(f"开始翻译文档: {input_file}")

        # 执行翻译
        output_file = processor.process_document(
            file_path=input_file,
            target_language="英语",  # 目标语言名称
            terminology=terminology,  # 术语库格式
            source_lang="zh",
            target_lang="en"
        )

        logger.info(f"翻译完成！输出文件: {output_file}")

        # 检查输出文件中的表格
        if output_file and os.path.exists(output_file):
            logger.info("正在检查翻译结果...")
            check_translation_result(output_file)

        return True
            
    except Exception as e:
        logger.error(f"翻译过程中出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_translation_result(output_file):
    """检查翻译结果"""
    try:
        from docx import Document
        
        doc = Document(output_file)
        logger.info(f"检查输出文档: {output_file}")
        
        if len(doc.tables) >= 4:
            table = doc.tables[3]  # 第4个表格
            logger.info(f"检查表格4，行数: {len(table.rows)}")
            
            if table.rows:
                last_row = table.rows[-1]
                logger.info(f"最后一行单元格数: {len(last_row.cells)}")
                
                if len(last_row.cells) >= 2:
                    # 检查最后两个单元格
                    second_last_cell = last_row.cells[-2]
                    last_cell = last_row.cells[-1]
                    
                    second_last_text = second_last_cell.text.strip()
                    last_text = last_cell.text.strip()
                    
                    logger.info(f"倒数第二个单元格内容: {second_last_text[:100]}...")
                    logger.info(f"最后一个单元格内容: {last_text[:100]}...")
                    
                    # 检查是否包含英文翻译
                    import re
                    
                    def has_meaningful_english(text):
                        english_words = re.findall(r'\b[A-Za-z]{2,}\b', text)
                        return len(english_words) >= 3
                    
                    second_last_has_english = has_meaningful_english(second_last_text)
                    last_has_english = has_meaningful_english(last_text)
                    
                    logger.info(f"倒数第二个单元格包含英文翻译: {second_last_has_english}")
                    logger.info(f"最后一个单元格包含英文翻译: {last_has_english}")
                    
                    if second_last_has_english and last_has_english:
                        logger.info("✅ 修复成功！表格最后两个单元格都包含英文翻译。")
                    else:
                        logger.warning("⚠️ 表格最后两个单元格可能仍然缺少英文翻译。")
                else:
                    logger.warning("最后一行单元格数不足")
            else:
                logger.warning("表格4没有行")
        else:
            logger.warning("文档中表格数量不足4个")
            
    except Exception as e:
        logger.error(f"检查翻译结果时出错: {e}")

def main():
    logger.info("开始测试修复后的文档翻译...")
    
    success = test_document_translation()
    
    if success:
        logger.info("✅ 文档翻译测试完成！")
    else:
        logger.error("❌ 文档翻译测试失败！")

if __name__ == "__main__":
    main()
