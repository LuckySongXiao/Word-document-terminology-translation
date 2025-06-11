#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析Word表格中的合并单元格问题
"""

import os
import sys
from docx import Document
import logging

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def analyze_table_structure(doc_path):
    """分析表格结构，特别是合并单元格"""
    
    try:
        doc = Document(doc_path)
        logger.info(f"分析文档: {doc_path}")
        
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        # 重点分析表格4
        table = doc.tables[3]
        logger.info(f"分析表格4，行数: {len(table.rows)}")
        
        # 分析每一行的单元格结构
        for row_idx, row in enumerate(table.rows):
            logger.info(f"\n=== 第{row_idx+1}行分析 ===")
            logger.info(f"单元格数量: {len(row.cells)}")
            
            # 检查每个单元格的内容和属性
            cell_contents = []
            cell_ids = []
            
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_id = id(cell)
                cell_contents.append(cell_text[:50] if cell_text else "[空]")
                cell_ids.append(cell_id)
                
                logger.info(f"  单元格 {cell_idx+1}: ID={cell_id}, 内容='{cell_text[:50]}{'...' if len(cell_text) > 50 else ''}'")
            
            # 检查是否有重复的内容
            content_counts = {}
            for content in cell_contents:
                if content != "[空]":
                    content_counts[content] = content_counts.get(content, 0) + 1
            
            duplicates = {content: count for content, count in content_counts.items() if count > 1}
            if duplicates:
                logger.warning(f"  发现重复内容: {duplicates}")
                
                # 检查重复内容的单元格是否是同一个对象
                for content, count in duplicates.items():
                    matching_cells = []
                    for cell_idx, cell_content in enumerate(cell_contents):
                        if cell_content == content:
                            matching_cells.append((cell_idx, cell_ids[cell_idx]))
                    
                    logger.info(f"    重复内容 '{content}' 出现在单元格: {matching_cells}")
                    
                    # 检查是否是同一个对象引用
                    unique_ids = set(cell_id for _, cell_id in matching_cells)
                    if len(unique_ids) == 1:
                        logger.warning(f"    ⚠️ 所有重复单元格都是同一个对象！这表明存在合并单元格")
                    else:
                        logger.info(f"    不同的对象，可能是翻译过程中的重复")
        
        return True
        
    except Exception as e:
        logger.error(f"分析表格结构时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_merged_cells_xml(doc_path):
    """通过XML结构检查合并单元格"""
    
    try:
        doc = Document(doc_path)
        table = doc.tables[3]  # 表格4
        
        logger.info("\n=== XML结构分析 ===")
        
        for row_idx, row in enumerate(table.rows):
            logger.info(f"\n第{row_idx+1}行 XML 分析:")
            
            for cell_idx, cell in enumerate(row.cells):
                # 获取单元格的XML元素
                tc_element = cell._tc
                
                # 检查合并属性
                grid_span = tc_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                v_merge = tc_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                h_merge = tc_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hMerge')
                
                merge_info = []
                if grid_span:
                    merge_info.append(f"gridSpan={grid_span}")
                if v_merge:
                    merge_info.append(f"vMerge={v_merge}")
                if h_merge:
                    merge_info.append(f"hMerge={h_merge}")
                
                if merge_info:
                    logger.info(f"  单元格 {cell_idx+1}: {', '.join(merge_info)}")
                    logger.info(f"    内容: '{cell.text.strip()[:50]}'")
                
        return True
        
    except Exception as e:
        logger.error(f"XML结构分析时出错: {e}")
        return False

def detect_duplicate_translation_pattern(doc_path):
    """检测重复翻译模式"""
    
    try:
        doc = Document(doc_path)
        table = doc.tables[3]  # 表格4
        
        logger.info("\n=== 重复翻译模式检测 ===")
        
        duplicate_rows = []
        
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) > 1:
                # 检查这一行是否所有单元格都包含相同内容
                first_cell_text = row.cells[0].text.strip()
                
                if first_cell_text and len(first_cell_text) > 20:  # 只检查有实质内容的单元格
                    all_same = True
                    for cell in row.cells[1:]:
                        if cell.text.strip() != first_cell_text:
                            all_same = False
                            break
                    
                    if all_same:
                        duplicate_rows.append(row_idx)
                        logger.warning(f"第{row_idx+1}行: 所有单元格内容相同")
                        logger.info(f"  内容: '{first_cell_text[:100]}{'...' if len(first_cell_text) > 100 else ''}'")
        
        if duplicate_rows:
            logger.error(f"发现 {len(duplicate_rows)} 行存在重复翻译问题: {[r+1 for r in duplicate_rows]}")
            return duplicate_rows
        else:
            logger.info("未发现重复翻译问题")
            return []
        
    except Exception as e:
        logger.error(f"重复翻译模式检测时出错: {e}")
        return []

def main():
    logger.info("开始分析Word表格中的合并单元格问题...")
    
    # 分析指定的文档
    output_dir = "输出"
    doc_name = "单晶电阻率管控技术标准_带翻译_20250610230213.docx"
    doc_path = os.path.join(output_dir, doc_name)
    
    if not os.path.exists(doc_path):
        logger.error(f"文档不存在: {doc_path}")
        return False
    
    # 1. 分析表格结构
    logger.info("1. 分析表格结构...")
    structure_ok = analyze_table_structure(doc_path)
    
    # 2. 检查XML结构中的合并单元格
    logger.info("\n2. 检查XML结构...")
    xml_ok = check_merged_cells_xml(doc_path)
    
    # 3. 检测重复翻译模式
    logger.info("\n3. 检测重复翻译模式...")
    duplicate_rows = detect_duplicate_translation_pattern(doc_path)
    
    # 总结分析结果
    logger.info("\n=== 分析结果总结 ===")
    if duplicate_rows:
        logger.error(f"发现重复翻译问题，涉及 {len(duplicate_rows)} 行")
        logger.info("建议解决方案:")
        logger.info("1. 在翻译前检测合并单元格")
        logger.info("2. 对合并单元格只翻译一次")
        logger.info("3. 避免对已翻译的合并单元格重复处理")
    else:
        logger.info("未发现明显的重复翻译问题")
    
    return structure_ok and xml_ok

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
