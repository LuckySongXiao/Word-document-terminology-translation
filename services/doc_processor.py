import os
import logging
import time
import json
import re
from typing import Dict, List, Any, Tuple
from .translator import TranslationService
import pandas as pd
from datetime import datetime
from utils.term_extractor import TermExtractor
from .translation_detector import TranslationDetector
from docx import Document

logger = logging.getLogger(__name__)

class DOCProcessor:
    """DOC文档处理器（旧版Word文档）"""
    
    def __init__(self, translator: TranslationService):
        self.translator = translator
        self.use_terminology = True  # 默认使用术语库
        self.preprocess_terms = True  # 是否预处理术语（默认启用）
        self.term_extractor = TermExtractor()  # 术语提取器
        self.translation_detector = TranslationDetector()  # 翻译检测器
        self.skip_translated_content = True  # 是否跳过已翻译内容（默认启用）
        self.export_pdf = False  # 是否导出PDF
        self.source_lang = "zh"  # 默认源语言为中文
        self.target_lang = "en"  # 默认目标语言为英文
        self.is_cn_to_foreign = True  # 默认翻译方向为中文→外语
        self.progress_callback = None  # 进度回调函数
        self.retry_count = 3  # 翻译失败重试次数
        self.retry_delay = 1  # 重试延迟（秒）
        self.output_format = "bilingual"  # 输出格式：bilingual（双语对照）或translation_only（仅翻译）

        # 配置日志记录器
        self.logger = logging.getLogger(__name__)
        self.web_logger = logging.getLogger('web_logger')

    def set_progress_callback(self, callback):
        """设置进度回调函数"""
        self.progress_callback = callback

    def _update_progress(self, progress: float, message: str):
        """更新进度"""
        if self.progress_callback:
            self.progress_callback(progress, message)
        
        # 记录进度到日志
        if progress >= 0:
            logger.info(f"进度: {progress*100:.1f}% - {message}")
        else:
            logger.error(f"进度: 错误 - {message}")

    def process_document(self, file_path: str, target_language: str, terminology: Dict, 
                        source_lang: str = "zh", target_lang: str = "en") -> str:
        """
        处理DOC文档翻译

        Args:
            file_path: 文件路径
            target_language: 目标语言名称（用于术语表查找）
            terminology: 术语词典
            source_lang: 源语言代码，默认为中文(zh)
            target_lang: 目标语言代码，默认为英文(en)

        Returns:
            str: 输出文件路径
        """
        # 更新进度：开始处理
        self._update_progress(0.01, "开始处理DOC文档...")

        # 设置翻译方向
        self.source_lang = source_lang
        self.target_lang = target_lang

        # 根据源语言和目标语言确定翻译方向
        self.is_cn_to_foreign = source_lang == "zh" or (source_lang == "auto" and target_lang != "zh")
        logger.info(f"翻译方向: {'中文→外语' if self.is_cn_to_foreign else '外语→中文'}")

        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise Exception("文件不存在")

        # 创建输出目录
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 生成输出文件名（转换为DOCX格式）
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        time_stamp = datetime.now().strftime("%Y%m%d%H%M%S")
        
        if self.output_format == "translation_only":
            output_path = os.path.join(output_dir, f"{file_name}_翻译_{time_stamp}.docx")
        else:
            output_path = os.path.join(output_dir, f"{file_name}_双语对照_{time_stamp}.docx")

        # 更新进度：转换DOC文件
        self._update_progress(0.1, "转换DOC文件为DOCX格式...")

        try:
            # 将DOC文件转换为文本，然后创建新的DOCX文档
            content = self._convert_doc_to_text(file_path)
            
            # 更新进度：准备翻译
            self._update_progress(0.2, "准备翻译...")

            # 获取目标语言的术语库
            target_terminology = terminology.get(target_language, {}) if isinstance(terminology, dict) else {}
            
            # 翻译结果存储
            translation_results = []
            used_terminology = {}

            # 按段落分割文本
            paragraphs = self._split_into_paragraphs(content)
            total_paragraphs = len(paragraphs)
            
            logger.info(f"共找到 {total_paragraphs} 个段落需要翻译")

            # 创建新的Word文档
            doc = Document()

            # 更新进度：开始翻译
            self._update_progress(0.3, f"开始翻译 {total_paragraphs} 个段落...")

            # 翻译每个段落并添加到文档
            for i, paragraph_text in enumerate(paragraphs):
                if paragraph_text.strip():
                    # 更新进度
                    progress = 0.3 + (i / total_paragraphs) * 0.5
                    self._update_progress(progress, f"翻译第 {i+1}/{total_paragraphs} 个段落...")

                    # 翻译段落
                    translated_text = self._translate_paragraph(
                        paragraph_text, target_terminology, translation_results, used_terminology
                    )
                    
                    # 根据输出格式添加到文档
                    if self.output_format == "translation_only":
                        # 仅翻译结果
                        doc.add_paragraph(translated_text)
                    else:
                        # 双语对照格式
                        # 添加原文段落
                        original_para = doc.add_paragraph()
                        original_run = original_para.add_run("【原文】")
                        original_run.bold = True
                        original_para.add_run(paragraph_text)
                        
                        # 添加译文段落
                        translated_para = doc.add_paragraph()
                        translated_run = translated_para.add_run("【译文】")
                        translated_run.bold = True
                        translated_para.add_run(translated_text)
                        
                        # 添加空行分隔
                        doc.add_paragraph()
                else:
                    # 保留空段落
                    doc.add_paragraph()

            # 更新进度：保存文件
            self._update_progress(0.8, "保存翻译后的文件...")

            # 保存翻译后的文件
            doc.save(output_path)
            logger.info(f"文件已保存到: {output_path}")

            # 更新进度：导出结果
            self._update_progress(0.9, "导出翻译结果...")

            # 导出翻译结果到Excel
            if translation_results:
                self.export_to_excel(translation_results, file_path, target_language)

            # 如果使用了术语预处理，导出使用的术语
            if used_terminology:
                self.export_used_terminology(used_terminology, file_path)

            # 更新进度：完成
            self._update_progress(1.0, "DOC翻译完成！")

            return output_path

        except Exception as e:
            logger.error(f"处理DOC文件时出错: {str(e)}")
            raise Exception(f"处理DOC文件失败: {str(e)}")

    def _convert_doc_to_text(self, file_path: str) -> str:
        """将DOC文件转换为文本"""
        try:
            logger.info(f"开始转换DOC文件: {file_path}")

            # 方法1: 尝试使用win32com转换DOC为DOCX，然后读取
            if os.name == 'nt':  # Windows系统
                try:
                    import win32com.client
                    logger.info("尝试使用win32com转换DOC文件...")

                    # 创建临时DOCX文件路径
                    temp_docx_path = file_path.replace('.doc', '_temp.docx')

                    # 创建Word应用程序实例
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = False  # 禁用警告对话框

                    # 打开DOC文档
                    doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=True)

                    # 另存为DOCX格式
                    doc.SaveAs2(os.path.abspath(temp_docx_path), FileFormat=16)  # 16 = docx格式

                    # 关闭文档和应用程序
                    doc.Close(SaveChanges=False)
                    word.Quit()

                    # 使用python-docx读取转换后的DOCX文件
                    try:
                        from docx import Document
                        docx_doc = Document(temp_docx_path)
                        paragraphs = []
                        for para in docx_doc.paragraphs:
                            if para.text.strip():
                                paragraphs.append(para.text)
                        content = '\n\n'.join(paragraphs)

                        # 删除临时文件
                        if os.path.exists(temp_docx_path):
                            os.remove(temp_docx_path)

                        if content and content.strip():
                            logger.info("使用win32com转换后成功提取DOC文件内容")
                            return content
                    except Exception as docx_error:
                        logger.warning(f"读取转换后的DOCX文件失败: {str(docx_error)}")
                        # 删除临时文件
                        if os.path.exists(temp_docx_path):
                            os.remove(temp_docx_path)

                except ImportError:
                    logger.warning("win32com库未安装，无法使用Word COM接口")
                except Exception as e:
                    logger.warning(f"使用win32com转换失败: {str(e)}")
                    # 确保Word应用程序被关闭
                    try:
                        import win32com.client
                        word = win32com.client.Dispatch("Word.Application")
                        word.Quit()
                    except:
                        pass

            # 方法2: 尝试使用textract库（如果可用）
            try:
                import textract
                logger.info("尝试使用textract提取DOC文件内容...")
                content = textract.process(file_path).decode('utf-8')
                if content and content.strip():
                    logger.info("使用textract成功提取DOC文件内容")
                    return content
            except ImportError:
                logger.warning("textract库未安装，跳过此方法")
            except Exception as e:
                logger.warning(f"使用textract提取内容失败: {str(e)}")

            # 方法3: 尝试使用python-docx2txt库（对某些DOC文件可能有效）
            try:
                import docx2txt
                logger.info("尝试使用docx2txt提取DOC文件内容...")
                content = docx2txt.process(file_path)
                if content and content.strip():
                    logger.info("使用docx2txt成功提取DOC文件内容")
                    return content
            except ImportError:
                logger.warning("docx2txt库未安装，跳过此方法")
            except Exception as e:
                logger.warning(f"使用docx2txt提取内容失败: {str(e)}")

            # 方法4: 尝试使用python-docx库（通常不支持.doc，但值得一试）
            try:
                from docx import Document
                logger.info("尝试使用python-docx提取DOC文件内容...")
                doc = Document(file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                content = '\n\n'.join(paragraphs)
                if content and content.strip():
                    logger.info("使用python-docx成功提取DOC文件内容")
                    return content
            except Exception as e:
                logger.warning(f"使用python-docx提取内容失败: {str(e)}")

            # 如果所有方法都失败，提供详细的解决方案
            error_msg = self._generate_doc_error_message()
            logger.error(f"DOC文件处理失败: {error_msg}")

            # 抛出一个更友好的异常，包含具体的解决建议
            friendly_error = (
                "无法处理DOC文件 '" + os.path.basename(file_path) + "'。\n\n" +
                "DOC格式是较旧的二进制格式，处理难度较大。建议：\n" +
                "1. 【推荐】将文件另存为DOCX格式后重新上传\n" +
                "2. 使用Microsoft Word打开文件，选择'另存为'，格式选择'Word文档(.docx)'\n" +
                "3. 或使用在线转换工具将DOC转换为DOCX格式\n\n" +
                "DOCX格式具有更好的兼容性和处理效果。"
            )
            raise Exception(friendly_error)

        except Exception as e:
            logger.error(f"转换DOC文件失败: {str(e)}")
            raise

    def _generate_doc_error_message(self) -> str:
        """生成DOC文件处理失败的详细错误消息"""
        error_msg = "无法读取DOC文件。请尝试以下解决方案：\n\n"

        if os.name == 'nt':  # Windows系统
            error_msg += "【推荐方案】\n"
            error_msg += "1. 确保系统已安装Microsoft Word\n"
            error_msg += "2. 将DOC文件另存为DOCX格式后重新上传\n\n"

            error_msg += "【技术方案】\n"
            error_msg += "3. 安装textract库：pip install textract\n"
            error_msg += "4. 重新安装pywin32：pip install --upgrade pywin32\n"
            error_msg += "5. 运行：python -m win32com.client.makepy Microsoft.Word\n\n"
        else:
            error_msg += "【推荐方案】\n"
            error_msg += "1. 将DOC文件另存为DOCX格式后重新上传\n"
            error_msg += "2. 使用在线转换工具将DOC转换为DOCX\n\n"

            error_msg += "【技术方案】\n"
            error_msg += "3. 安装textract库：pip install textract\n"
            error_msg += "4. 安装antiword：sudo apt-get install antiword (Linux)\n\n"

        error_msg += "【通用方案】\n"
        error_msg += "• 使用Microsoft Word打开文件，另存为DOCX格式\n"
        error_msg += "• 使用LibreOffice Writer打开文件，导出为DOCX格式\n"
        error_msg += "• 使用在线文档转换服务\n\n"

        error_msg += "注意：旧版DOC格式（.doc）是二进制格式，处理难度较大。\n"
        error_msg += "建议优先使用现代DOCX格式（.docx）以获得最佳兼容性。"

        return error_msg

    def _split_into_paragraphs(self, content: str) -> List[str]:
        """将文本分割为段落"""
        # 按双换行符分割段落
        paragraphs = re.split(r'\n\s*\n', content)
        
        # 如果没有双换行符，按单换行符分割
        if len(paragraphs) == 1:
            paragraphs = content.split('\n')
        
        # 过滤空段落
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        return paragraphs

    def _translate_paragraph(self, text: str, terminology: Dict,
                           translation_results: List, used_terminology: Dict) -> str:
        """翻译单个段落"""
        try:
            original_text = text

            # 检查是否应该跳过翻译（已翻译内容检测）
            if self.skip_translated_content:
                should_skip, reason = self.translation_detector.should_skip_translation(
                    text, self.source_lang, self.target_lang
                )
                if should_skip:
                    logger.info(f"跳过DOC段落翻译: {reason} - {text[:50]}...")
                    return text  # 返回原文

            # 使用新的直接术语替换方法（如果启用）
            if self.preprocess_terms and terminology:
                # 直接替换术语为目标语言，避免占位符机制
                text = self._preprocess_terminology_direct(text, terminology)
                logger.info(f"直接替换术语后的文本: {text[:100]}...")

            # 翻译文本（不再需要术语库，因为已经直接替换了）
            translated_text = self.translator.translate(
                text, {}, self.source_lang, self.target_lang
            )

            # 记录翻译结果
            translation_results.append({
                "原文": original_text,
                "译文": translated_text,
                "源语言": self.source_lang,
                "目标语言": self.target_lang
            })

            return translated_text

        except Exception as e:
            logger.error(f"翻译段落失败: {str(e)}")
            # 翻译失败时返回原文
            return original_text

    def _preprocess_terminology_direct(self, text: str, terminology: Dict) -> str:
        """直接替换术语为目标语言，避免占位符机制"""
        if not terminology:
            return text

        processed_text = text

        # 按术语长度排序，优先处理长术语
        sorted_terms = sorted(terminology.items(), key=lambda x: len(x[0]), reverse=True)

        for term, translation in sorted_terms:
            if term in processed_text and translation:
                processed_text = processed_text.replace(term, translation)

        return processed_text

    def _preprocess_terminology(self, text: str, terminology: Dict) -> Tuple[str, Dict]:
        """预处理术语，使用占位符替换（保留原有方法以兼容性）"""
        used_terms = {}
        processed_text = text

        # 按术语长度排序，优先处理长术语
        sorted_terms = sorted(terminology.items(), key=lambda x: len(x[0]), reverse=True)

        for term, translation in sorted_terms:
            if term in processed_text:
                placeholder = f"__TERM_{len(used_terms)}__"
                processed_text = processed_text.replace(term, placeholder)
                used_terms[placeholder] = {"term": term, "translation": translation}

        return processed_text, used_terms

    def export_to_excel(self, translation_results: List[Dict], file_path: str, target_language: str = None):
        """将翻译结果导出到Excel文件"""
        try:
            # 创建输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 获取文件名
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            time_stamp = datetime.now().strftime("%Y%m%d%H%M%S")

            # 根据目标语言生成文件名
            lang_suffix = f"_{target_language}" if target_language else ""
            excel_path = os.path.join(output_dir, f"{file_name}_翻译结果{lang_suffix}_{time_stamp}.xlsx")

            # 转换翻译结果为DataFrame
            df = pd.DataFrame(translation_results)

            # 保存到Excel
            df.to_excel(excel_path, index=False, engine='openpyxl')
            logger.info(f"翻译结果已导出到: {excel_path}")

        except Exception as e:
            logger.error(f"导出翻译结果到Excel失败: {str(e)}")

    def export_used_terminology(self, used_terminology: Dict, file_path: str):
        """导出使用的术语到Excel文件"""
        try:
            if not used_terminology:
                return

            # 创建输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 获取文件名
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            time_stamp = datetime.now().strftime("%Y%m%d%H%M%S")
            excel_path = os.path.join(output_dir, f"{file_name}_使用术语_{time_stamp}.xlsx")

            # 准备术语数据
            terminology_data = []
            for placeholder, term_info in used_terminology.items():
                terminology_data.append({
                    "原术语": term_info["term"],
                    "翻译": term_info["translation"],
                    "占位符": placeholder
                })

            # 转换为DataFrame并保存
            df = pd.DataFrame(terminology_data)
            df.to_excel(excel_path, index=False, engine='openpyxl')
            logger.info(f"使用的术语已导出到: {excel_path}")

        except Exception as e:
            logger.error(f"导出使用术语失败: {str(e)}")
