import os
import logging
import time
import json
import re
import uuid
import pdfplumber
from typing import Dict, List, Tuple
from .translator import TranslationService
import pandas as pd
from datetime import datetime
from utils.term_extractor import TermExtractor
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WD_ALIGN_PARAGRAPH
# 导入PDF转换模块
try:
    from docx2pdf import convert as docx2pdf_convert
except ImportError:
    docx2pdf_convert = None
    logging.warning("docx2pdf模块未安装，PDF导出功能将不可用")

logger = logging.getLogger(__name__)

class PDFProcessor:
    """PDF文件处理器，用于从PDF文件中提取文本并进行翻译"""

    def __init__(self, translator: TranslationService):
        self.translator = translator
        self.use_terminology = True  # 默认使用术语库
        self.preprocess_terms = True  # 是否预处理术语（默认启用）
        self.term_extractor = TermExtractor()  # 术语提取器
        self.export_pdf = False  # 是否导出PDF
        self.output_format = "bilingual"  # 默认双语对照
        self.source_lang = "en"  # 默认源语言为英文
        self.target_lang = "zh"  # 默认目标语言为中文
        self.is_cn_to_foreign = False  # 默认翻译方向为外语→中文
        self.progress_callback = None  # 进度回调函数
        self.retry_count = 3  # 翻译失败重试次数
        self.retry_delay = 1  # 重试延迟（秒）

        # 配置日志记录器 - 添加Web日志记录器以确保日志同步
        self.web_logger = logging.getLogger('web_logger')

        # 初始化反向术语库缓存（用于外语→中文翻译优化）
        self.reversed_terminology = None

        # 初始化日志配置
        logger.info("PDFProcessor initialized")
        self.web_logger.info("PDF translation service ready")
        # 文本处理参数
        self.min_paragraph_length = 10  # 最小段落长度（字符数）
        self.optimal_paragraph_length = 200  # 理想段落长度
        self.max_paragraph_length = 1000  # 最大段落长度
        self.batch_size = 5  # 批量处理段落数量
        self.preserve_structure = True  # 保留原始文档结构
        self.respect_newlines = True  # 尊重原始换行符
        self.preserve_images = True  # 保留图片信息
        # 数学公式正则表达式模式 - 扩展以支持更多LaTeX公式格式
        self.latex_patterns = [
            r'\$\$(.*?)\$\$',  # 行间公式 $$...$$
            r'\$(.*?)\$',      # 行内公式 $...$
            r'\\begin\{equation\*?\}(.*?)\\end\{equation\*?\}',  # equation环境（包括equation*）
            r'\\begin\{align\*?\}(.*?)\\end\{align\*?\}',        # align环境（包括align*）
            r'\\begin\{eqnarray\*?\}(.*?)\\end\{eqnarray\*?\}',  # eqnarray环境（包括eqnarray*）
            r'\\begin\{gather\*?\}(.*?)\\end\{gather\*?\}',      # gather环境（包括gather*）
            r'\\begin\{multline\*?\}(.*?)\\end\{multline\*?\}',  # multline环境（包括multline*）
            r'\\begin\{math\}(.*?)\\end\{math\}',                # math环境
            r'\\begin\{displaymath\}(.*?)\\end\{displaymath\}',  # displaymath环境
            r'\\[(.*?)\\]',                                      # \[...\] 行间公式
            r'\\((.*?)\\)',                                      # \(...\) 行内公式
            r'\\begin\{cases\}(.*?)\\end\{cases\}'               # cases环境
        ]
        # 添加LaTeX命令模式，用于识别可能的公式片段
        self.latex_commands = [
            r'\\frac\{.*?\}\{.*?\}',  # 分数
            r'\\sum_\{.*?\}\^.*?',    # 求和
            r'\\int_\{.*?\}\^.*?',    # 积分
            r'\\prod_\{.*?\}\^.*?',   # 连乘
            r'\\lim_\{.*?\}',         # 极限
            r'\\sqrt\{.*?\}',         # 平方根
            r'\\vec\{.*?\}',          # 向量
            r'\\overrightarrow\{.*?\}', # 箭头
            r'\\mathbf\{.*?\}'        # 粗体数学符号
        ]

    def _create_reversed_terminology(self, terminology: Dict[str, str]) -> Dict[str, str]:
        """
        创建反向术语库缓存，用于外语→中文翻译优化

        Args:
            terminology: 原始术语库 {中文术语: 外语术语}

        Returns:
            Dict[str, str]: 反向术语库 {外语术语: 中文术语}
        """
        if not terminology:
            logger.warning("术语库为空，无法创建反向术语库")
            self.web_logger.warning("Terminology is empty, cannot create reversed terminology")
            return {}

        reversed_dict = {}
        valid_count = 0

        for cn_term, foreign_term in terminology.items():
            if cn_term and cn_term.strip() and foreign_term and foreign_term.strip():
                # 确保外语术语不重复
                if foreign_term not in reversed_dict:
                    reversed_dict[foreign_term] = cn_term
                    valid_count += 1
                else:
                    # 如果外语术语重复，记录警告但保留第一个
                    logger.warning(f"外语术语重复: {foreign_term}, 保留第一个映射: {reversed_dict[foreign_term]}")

        logger.info(f"创建反向术语库完成，有效术语对: {valid_count}/{len(terminology)}")
        self.web_logger.info(f"Created reversed terminology cache with {valid_count} valid term pairs")

        return reversed_dict

    def set_progress_callback(self, callback):
        """设置进度回调函数"""
        self.progress_callback = callback

    def _update_progress(self, progress: float, message: str = ""):
        """更新进度"""
        if self.progress_callback:
            import asyncio
            try:
                # 尝试异步调用回调函数
                if asyncio.iscoroutinefunction(self.progress_callback):
                    # 检查是否已有事件循环在运行
                    try:
                        loop = asyncio.get_event_loop()
                        if loop.is_running():
                            # 如果当前事件循环正在运行，创建一个任务而不是新的事件循环
                            asyncio.create_task(self.progress_callback(progress, message))
                        else:
                            # 如果当前事件循环不在运行，使用它来运行协程
                            loop.run_until_complete(self.progress_callback(progress, message))
                    except RuntimeError:
                        # 如果没有事件循环，创建一个新的
                        loop = asyncio.new_event_loop()
                        asyncio.set_event_loop(loop)
                        try:
                            loop.run_until_complete(self.progress_callback(progress, message))
                        finally:
                            loop.close()
            except Exception as e:
                logger.error(f"更新进度失败: {str(e)}")

        # 记录进度到日志
        if progress >= 0:
            logger.info(f"进度: {progress*100:.1f}% - {message}")
        else:
            logger.error(f"进度: 错误 - {message}")

    def process_document(self, file_path: str, target_language: str, terminology: Dict, source_lang: str = "en", target_lang: str = "zh") -> str:
        """
        处理PDF文档翻译

        Args:
            file_path: 文件路径
            target_language: 目标语言名称（用于术语表查找）
            terminology: 术语词典
            source_lang: 源语言代码，默认为英文(en)
            target_lang: 目标语言代码，默认为中文(zh)

        Returns:
            str: 输出文件路径
        """
        # 更新进度：开始处理
        self._update_progress(0.01, "开始处理PDF文档...")

        # 设置翻译方向
        self.source_lang = source_lang
        self.target_lang = target_lang

        # 根据源语言和目标语言确定翻译方向
        self.is_cn_to_foreign = source_lang == "zh" or (source_lang == "auto" and target_lang != "zh")
        logger.info(f"翻译方向: {'中文→外语' if self.is_cn_to_foreign else '外语→中文'}")

        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise Exception("文件不存在")

        # 更新进度：准备输出目录
        self._update_progress(0.05, "准备输出目录...")

        # 在程序根目录下创建输出目录
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 获取输出路径，添加时间戳避免文件名冲突
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        time_stamp = time.strftime("%Y%m%d%H%M%S")
        output_file_name = f"{file_name}_带翻译_{time_stamp}.docx"
        output_path = os.path.join(output_dir, output_file_name)

        # 更新进度：创建文档
        self._update_progress(0.1, "创建文档...")

        # 创建一个新的Word文档来保存翻译结果
        doc = Document()

        # 更新进度：加载术语库
        self._update_progress(0.15, "加载术语库...")

        # 获取选定语言的术语表
        target_terminology = terminology.get(target_language, {})
        logger.info(f"使用 {target_language} 术语表，包含 {len(target_terminology)} 个术语")

        # 创建一个列表来收集翻译结果
        translation_results = []

        try:
            # 更新进度：打开PDF文件
            self._update_progress(0.2, "打开PDF文件并提取文本...")

            # 打开PDF文件并提取文本
            with pdfplumber.open(file_path) as pdf:
                # 如果启用了术语预处理，先收集所有使用的术语
                used_terminology = {}
                if self.preprocess_terms:
                    # 更新进度：提取术语
                    self._update_progress(0.25, "从PDF中提取术语...")

                    # 添加详细的术语预处理日志
                    logger.info("=== PDF术语预处理已启用，开始收集术语 ===")
                    self.web_logger.info("=== PDF术语预处理已启用，开始收集术语 ===")
                    logger.info(f"翻译方向: {self.source_lang} -> {self.target_lang}")
                    self.web_logger.info(f"翻译方向: {self.source_lang} -> {self.target_lang}")
                    logger.info(f"术语库大小: {len(target_terminology)} 个术语")
                    self.web_logger.info(f"术语库大小: {len(target_terminology)} 个术语")

                    # 显示术语库样本
                    if target_terminology:
                        sample_terms = list(target_terminology.items())[:5]
                        logger.info(f"术语库样本（前5个）: {sample_terms}")
                        self.web_logger.info(f"术语库样本（前5个）: {sample_terms}")

                    # 如果是外语→中文翻译，创建反向术语库缓存以提高效率
                    if self.source_lang != "zh":
                        logger.info("创建反向术语库缓存以优化外语→中文翻译...")
                        self.web_logger.info("Creating reversed terminology cache for foreign→Chinese translation optimization...")
                        self.reversed_terminology = self._create_reversed_terminology(target_terminology)

                        if self.reversed_terminology:
                            logger.info(f"反向术语库缓存创建成功，包含 {len(self.reversed_terminology)} 个术语对")
                            self.web_logger.info(f"Reversed terminology cache created successfully with {len(self.reversed_terminology)} term pairs")
                            # 显示反向术语库样本
                            sample_reversed = list(self.reversed_terminology.items())[:5]
                            logger.info(f"反向术语库样本（前5个）: {sample_reversed}")
                            self.web_logger.info(f"Reversed terminology sample: {sample_reversed}")
                        else:
                            logger.warning("反向术语库缓存创建失败或为空")
                            self.web_logger.warning("Reversed terminology cache creation failed or empty")

                    # 遍历所有页面，收集使用的术语
                    for page_idx, page in enumerate(pdf.pages, 1):
                        self.web_logger.info(f"Processing PDF page {page_idx}/{len(pdf.pages)} for terminology")
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            # 根据翻译方向选择不同的术语提取方法
                            if self.source_lang == "zh":
                                # 中文 → 外语
                                logger.info(f"从中文PDF页面提取术语: 第{page_idx}页...")
                                self.web_logger.info(f"从中文PDF页面提取术语: 第{page_idx}页...")
                                page_terms = self.term_extractor.extract_terms(page_text, target_terminology)
                                logger.info(f"从中文PDF页面提取术语: {len(page_terms)} 个")
                                self.web_logger.info(f"从中文PDF页面提取术语: {len(page_terms)} 个")
                            else:
                                # 外语 → 中文，使用优化的术语提取方法
                                logger.info(f"从外语PDF页面提取术语: 第{page_idx}页...")
                                self.web_logger.info(f"从外语PDF页面提取术语: 第{page_idx}页...")

                                # 检查是否有缓存的反向术语库
                                if self.reversed_terminology:
                                    # 使用缓存的反向术语库进行高效匹配
                                    logger.info("使用缓存的反向术语库进行术语提取")
                                    self.web_logger.info("Using cached reversed terminology for efficient term extraction")
                                    page_terms = self.term_extractor.extract_foreign_terms_from_reversed_dict(page_text, self.reversed_terminology)
                                    logger.info(f"从外语PDF页面提取术语（使用缓存）: {len(page_terms)} 个")
                                    self.web_logger.info(f"从外语PDF页面提取术语（使用缓存）: {len(page_terms)} 个")
                                else:
                                    # 回退到原始方法
                                    logger.info("缓存不可用，使用原始术语提取方法")
                                    self.web_logger.info("Cache unavailable, using original term extraction method")
                                    page_terms = self.term_extractor.extract_foreign_terms_by_chinese_values(page_text, target_terminology)
                                    logger.info(f"从外语PDF页面提取术语（原始方法）: {len(page_terms)} 个")
                                    self.web_logger.info(f"从外语PDF页面提取术语（原始方法）: {len(page_terms)} 个")

                            # 显示提取到的术语
                            if page_terms:
                                logger.info(f"第{page_idx}页提取到的术语: {list(page_terms.items())[:3]}")
                                self.web_logger.info(f"第{page_idx}页提取到的术语: {list(page_terms.items())[:3]}")

                            # 更新使用的术语词典
                            used_terminology.update(page_terms)

                    logger.info(f"从PDF中提取了 {len(used_terminology)} 个术语")
                    self.web_logger.info(f"从PDF中提取了 {len(used_terminology)} 个术语")

                    # 如果有使用的术语，导出到Excel文件
                    if used_terminology:
                        self._export_used_terminology(used_terminology)

                # 更新进度：开始处理页面
                self._update_progress(0.3, "开始处理PDF页面...")

                # 处理每一页
                for i, page in enumerate(pdf.pages):
                    # 更新页面进度
                    page_progress = 0.3 + (0.5 * (i / len(pdf.pages)))
                    self._update_progress(page_progress, f"正在处理第 {i+1} 页，共 {len(pdf.pages)} 页")
                    logger.info(f"正在处理第 {i+1} 页，共 {len(pdf.pages)} 页")

                    # 提取页面文本并优化
                    page_text = self._extract_and_optimize_text(page, i+1)

                    # 提取页面图片信息
                    images = []
                    if self.preserve_images:
                        images = self._extract_image_info(page)
                        if images:
                            logger.info(f"第 {i+1} 页提取到 {len(images)} 张图片")

                    if not page_text.strip() and not images:
                        logger.warning(f"第 {i+1} 页没有可提取的文本和图片")
                        continue

                    # 添加页码标题
                    page_heading = doc.add_heading(f"第 {i+1} 页", level=1)
                    page_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 准备图片处理
                    if images:
                        from docx.shared import Inches

                        # 创建临时图片缓存文件夹
                        pictures_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "pictures")
                        if not os.path.exists(pictures_dir):
                            os.makedirs(pictures_dir)
                            logger.info(f"创建临时图片缓存文件夹: {pictures_dir}")

                        # 清理旧的临时图片文件
                        try:
                            for old_file in os.listdir(pictures_dir):
                                if old_file.startswith("temp_img_"):
                                    os.remove(os.path.join(pictures_dir, old_file))
                            logger.info("清理旧的临时图片文件")
                        except Exception as clean_err:
                            logger.warning(f"清理旧的临时图片文件时出错: {str(clean_err)}")

                        # 生成唯一的会话ID，用于标识本次翻译的图片
                        session_id = str(uuid.uuid4())[:8]

                        for idx, img_info in enumerate(images):
                            # 根据图片类型处理
                            if img_info.get('type') == 'image':
                                # 获取图片数据流
                                image_data = img_info.get('image_data')

                                if image_data:
                                    try:
                                        # 保存图片到临时文件，使用猜测的扩展名
                                        img_ext = img_info.get('image_ext', '.png')  # 获取猜测的扩展名，默认为.png
                                        img_filename = f"temp_img_{session_id}_{idx+1}{img_ext}"
                                        img_path = os.path.join(pictures_dir, img_filename)

                                        # 简化图片数据获取 - 参照Word文档处理器的简洁方式
                                        try:
                                            # 获取图片数据 - 使用最直接的方法
                                            image_bytes = None

                                            # 优先尝试最常用的方法
                                            if hasattr(image_data, 'get_data'):
                                                try:
                                                    image_bytes = image_data.get_data()
                                                    logger.info("成功获取图片数据")
                                                except Exception as e:
                                                    logger.debug(f"get_data()方法失败: {str(e)}")

                                            # 如果第一种方法失败，尝试备用方法
                                            if image_bytes is None and hasattr(image_data, 'get_rawdata'):
                                                try:
                                                    image_bytes = image_data.get_rawdata()
                                                    logger.info("使用备用方法获取图片数据")
                                                except Exception as e:
                                                    logger.debug(f"get_rawdata()方法失败: {str(e)}")

                                            # 检查是否成功获取图片数据
                                            if image_bytes is None or len(image_bytes) == 0:
                                                logger.info(f"无法获取图片 {idx+1} 的数据，使用简洁占位符")
                                                # 创建简洁的图片占位符 - 参照Word文档处理器的风格
                                                img_para = doc.add_paragraph()
                                                img_run = img_para.add_run()
                                                img_run.add_text(f"[图片 {idx+1}]")
                                                img_run.bold = True
                                                img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                continue

                                            # 简化图片处理 - 参照Word文档处理器的简洁方式
                                            logger.info(f"获取到图片数据，大小: {len(image_bytes)} 字节")

                                            # 使用简化的图片保存方法
                                            try:
                                                # 直接保存图片数据到文件
                                                with open(img_path, 'wb') as f:
                                                    f.write(image_bytes)
                                                logger.info(f"图片 {idx+1} 已保存到: {img_path}")

                                                # 验证文件是否成功保存
                                                if not os.path.exists(img_path) or os.path.getsize(img_path) == 0:
                                                    raise Exception("图片文件保存失败或为空")

                                            except Exception as save_err:
                                                logger.warning(f"保存图片失败: {str(save_err)}，使用占位符")
                                                # 如果保存失败，创建简洁的占位符
                                                img_para = doc.add_paragraph()
                                                img_run = img_para.add_run()
                                                img_run.add_text(f"[图片 {idx+1} - 保存失败]")
                                                img_run.bold = True
                                                img_run.font.color.rgb = RGBColor(255, 165, 0)  # 橙色
                                                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                continue
                                        except Exception as e:
                                            logger.error(f"处理图片数据时出错: {str(e)}")
                                            # 创建一个空白图片作为替代
                                            try:
                                                from PIL import Image
                                                img = Image.new('RGB', (300, 200), color=(255, 255, 255))
                                                img.save(img_path, format='JPEG')
                                                logger.info(f"创建了空白替代图片: {img_path}")
                                            except Exception as blank_err:
                                                logger.error(f"创建空白图片失败: {str(blank_err)}")
                                                # 不抛出异常，继续处理其他内容
                                                logger.warning("无法创建图片，将使用占位符替代")

                                        logger.info(f"图片 {idx+1} 已保存到临时文件: {img_path}")

                                        # 简化图片插入 - 参照Word文档处理器的简洁方式
                                        img_para = doc.add_paragraph()
                                        img_run = img_para.add_run()
                                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 统一居中对齐

                                        # 简化图片尺寸计算
                                        max_width = Inches(4)  # 固定最大宽度为4英寸，简化计算

                                        # 尝试插入图片
                                        try:
                                            if os.path.exists(img_path):
                                                img_run.add_picture(img_path, width=max_width)
                                                logger.info(f"成功插入图片 {idx+1}")
                                            else:
                                                # 简洁的占位符
                                                img_run.add_text(f"[图片 {idx+1}]")
                                                img_run.bold = True
                                                img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                                logger.info(f"图片文件不存在，使用占位符")
                                        except Exception as pic_err:
                                            logger.warning(f"插入图片失败: {str(pic_err)}，使用占位符")
                                            # 清空run内容，添加占位符
                                            img_run.clear()
                                            img_run.add_text(f"[图片 {idx+1}]")
                                            img_run.bold = True
                                            img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

                                        # 简化图片标题处理 - 参照Word文档处理器的简洁方式
                                        if img_info.get('caption'):
                                            caption_para = doc.add_paragraph()
                                            caption_text = img_info.get('caption')
                                            caption_run = caption_para.add_run(caption_text)
                                            caption_run.italic = True
                                            caption_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
                                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                    except Exception as img_err:
                                        logger.warning(f"处理图片时出错: {str(img_err)}，使用占位符")
                                        # 简洁的错误处理 - 直接使用占位符
                                        img_para = doc.add_paragraph()
                                        img_run = img_para.add_run(f"[图片 {idx+1}]")
                                        img_run.bold = True
                                        img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    # 如果没有图片数据，使用简洁的占位符
                                    img_para = doc.add_paragraph()
                                    img_run = img_para.add_run(f"[图片 {idx+1}]")
                                    img_run.bold = True
                                    img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    logger.info(f"未找到图片数据，使用占位符 {idx+1}")

                            elif img_info.get('type') == 'figure_caption':
                                # 独立的图片标题文本（没有关联的图片）
                                caption_para = doc.add_paragraph()
                                caption_text = img_info.get('text', f"[图片标题 {idx+1}]")
                                caption_run = caption_para.add_run(caption_text)
                                caption_run.italic = True
                                caption_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

                                # 如果是双语对照模式，添加翻译
                                if self.output_format == "bilingual":
                                    try:
                                        # 翻译图片标题
                                        caption_translation = self.translator.translate_text(
                                            caption_text,
                                            target_terminology,
                                            self.source_lang,
                                            self.target_lang
                                        )

                                        # 添加翻译
                                        trans_para = doc.add_paragraph()
                                        trans_run = trans_para.add_run(caption_translation)
                                        trans_run.italic = True
                                        trans_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色

                                        # 收集翻译结果
                                        translation_results.append({
                                            'original': caption_text,
                                            'translated': caption_translation
                                        })
                                    except Exception as e:
                                        logger.error(f"翻译图片标题时出错: {str(e)}")

                            # 添加空行分隔
                            doc.add_paragraph()

                    # 改进的段落处理逻辑，更好地保留文档结构
                    # 检查文本是否已经包含明确的段落分隔
                    if self.preserve_structure and '\n\n' in page_text:
                        # 文本已经有明确的段落分隔，直接使用
                        logger.info("检测到明确的段落分隔，保留原始文档结构")
                        raw_paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]

                        # 过滤掉过短的段落，但不合并段落（保留原始结构）
                        valid_paragraphs = [p for p in raw_paragraphs if len(p) >= self.min_paragraph_length]

                        logger.info(f"第 {i+1} 页原始段落数: {len(raw_paragraphs)}, 有效段落数: {len(valid_paragraphs)}")
                    else:
                        # 没有明确的段落分隔，使用智能段落处理
                        # 首先按换行符分割文本
                        raw_paragraphs = page_text.split('\n')
                        logger.info(f"第 {i+1} 页原始行数: {len(raw_paragraphs)}")

                        # 智能段落识别
                        paragraphs = self._smart_paragraph_detection(raw_paragraphs)

                        # 过滤掉过短的段落
                        valid_paragraphs = [p for p in paragraphs if p.strip() and len(p.strip()) >= self.min_paragraph_length]

                        logger.info(f"第 {i+1} 页智能段落识别结果: {len(paragraphs)} 段，有效段落: {len(valid_paragraphs)} 段")

                    # 创建一个包含文本段落和图片的有序列表
                    ordered_content = []

                    # 添加文本段落，记录它们的原始位置
                    for idx, para_text in enumerate(valid_paragraphs):
                        ordered_content.append({
                            'type': 'text',
                            'content': para_text,
                            'position': idx
                        })

                    # 添加图片，根据它们在页面中的位置
                    for img_info in images:
                        # 只处理实际的图片，跳过图片标题等
                        if img_info.get('type') == 'image':
                            # 计算图片在文本段落中的位置
                            # 使用y坐标来确定图片应该插入在哪个段落之后
                            # 获取图片的y坐标，用于确定位置
                            img_y_pos = img_info.get('y0', 0)

                            # 找到最接近的段落位置
                            # 这里简化处理，按顺序插入
                            # 实际上可以根据段落在PDF中的位置更精确地确定
                            insert_pos = len(ordered_content) // 2  # 默认放在中间位置

                            ordered_content.append({
                                'type': 'image',
                                'content': img_info,
                                'position': insert_pos
                            })

                    # 按位置排序
                    ordered_content.sort(key=lambda x: x['position'])

                    logger.info(f"创建了有序内容列表，包含 {len(ordered_content)} 个元素")

                    # 处理有序内容列表
                    text_count = sum(1 for item in ordered_content if item['type'] == 'text')
                    logger.info(f"开始处理有序内容，包含 {text_count} 个文本段落和 {len(ordered_content) - text_count} 个图片")

                    # 收集文本段落用于批量处理
                    text_paragraphs = [item['content'] for item in ordered_content if item['type'] == 'text']
                    batch_size = self.batch_size

                    # 预先翻译所有文本段落
                    translated_paragraphs = {}
                    logger.info(f"开始翻译PDF文本段落，共 {len(text_paragraphs)} 个段落")
                    self.web_logger.info(f"开始翻译PDF文本段落，共 {len(text_paragraphs)} 个段落")

                    for i in range(0, len(text_paragraphs), batch_size):
                        # 获取当前批次的段落
                        batch_paragraphs = text_paragraphs[i:i+batch_size]
                        logger.info(f"翻译第 {i//batch_size + 1} 批文本，共 {len(batch_paragraphs)} 个段落")
                        self.web_logger.info(f"Processing batch {i//batch_size + 1}/{(len(text_paragraphs) + batch_size - 1) // batch_size}")

                        # 翻译每个段落并存储结果
                        for para_idx, para_text in enumerate(batch_paragraphs):
                            global_idx = i + para_idx
                            self.web_logger.info(f"Translating paragraph {global_idx + 1}/{len(text_paragraphs)}")

                            # 检查段落是否是图片信息行（通常包含"图"、"Figure"等关键词）
                            is_image_line = False
                            image_keywords = ["图", "Figure", "图片", "Image", "图表", "Chart", "插图", "Illustration"]
                            for keyword in image_keywords:
                                if keyword in para_text and len(para_text) < 200:  # 图片说明通常较短
                                    is_image_line = True
                                    logger.info(f"检测到图片信息行: {para_text[:50]}")
                                    break

                            # 初始化text变量，确保在所有代码路径中都有定义
                            text = para_text

                            # 翻译段落内容
                            try:
                                # 检查段落中是否包含数学公式
                                has_formula = any(pattern in para_text for pattern in ['$', '\\begin', '\\end', '\\[', '\\]', '\\(', '\\)'])

                                # 如果包含数学公式，提取公式并特殊处理
                                if has_formula:
                                    logger.info(f"段落中包含数学公式，进行特殊处理")
                                    # 提取LaTeX公式并用占位符替换
                                    processed_text, formulas = self._extract_latex_formulas(para_text)

                                    # 如果段落主要是公式，直接保持原样不翻译
                                    if len(para_text.strip()) < 100 and formulas:  # 短段落可能主要是公式
                                        logger.info("段落主要是公式，保持原样")
                                        translation = para_text
                                    else:
                                        # 如果段落包含公式但不全是公式，翻译非公式部分
                                        text = processed_text
                                        # 翻译处理后的文本
                                        if self.preprocess_terms:
                                            # 使用术语预处理方式翻译
                                            if self.source_lang == "zh":
                                                # 中文 → 外语
                                                # 检查是否有可用的术语
                                                if not used_terminology:
                                                    # 如果没有术语，使用常规翻译
                                                    logger.info("未找到匹配术语，使用常规翻译（带术语库）")
                                                    translation = self.translator.translate_text(text, target_terminology, self.source_lang, self.target_lang)
                                                else:
                                                    # 直接使用翻译器的内置术语处理功能
                                                    logger.info(f"使用翻译器内置术语处理功能，找到 {len(used_terminology)} 个匹配术语")
                                                    translation = self.translator.translate_text(text, used_terminology, self.source_lang, self.target_lang)
                                            else:
                                                # 外语 → 中文
                                                # 检查是否有可用的术语
                                                if not used_terminology:
                                                    # 如果没有术语，使用常规翻译
                                                    logger.info("未找到匹配术语，使用常规翻译")
                                                    self.web_logger.info(f"No terminology found for paragraph {global_idx + 1}, using regular translation")
                                                    translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)
                                                else:
                                                    # 使用term_extractor的占位符系统进行术语预处理
                                                    # 参照Word文档处理器的流程
                                                    # used_terminology已经是 {外语术语: 中文术语} 格式，直接使用
                                                    reverse_terminology = used_terminology

                                                    logger.info(f"使用术语预处理方式翻译，找到 {len(reverse_terminology)} 个匹配术语")
                                                    self.web_logger.info(f"Using terminology preprocessing for paragraph {global_idx + 1}, found {len(reverse_terminology)} terms")

                                                    # 显示术语预处理的详细信息
                                                    if reverse_terminology:
                                                        sample_terms = list(reverse_terminology.items())[:3]
                                                        logger.info(f"术语样本: {sample_terms}")
                                                        self.web_logger.info(f"Terminology sample: {sample_terms}")

                                                    processed_text = self.term_extractor.replace_foreign_terms_with_placeholders(text, reverse_terminology)
                                                    self.web_logger.info(f"Text after placeholder replacement: {processed_text[:100]}...")

                                                    translated_with_placeholders = self.translator.translate_text(processed_text, None, self.source_lang, self.target_lang)
                                                    self.web_logger.info(f"Translation with placeholders: {translated_with_placeholders[:100]}...")

                                                    translation = self.term_extractor.restore_placeholders_with_chinese_terms(translated_with_placeholders)
                                                    self.web_logger.info(f"Final translation: {translation[:100]}...")
                                        else:
                                            # 使用常规方式翻译
                                            translation = self.translator.translate_text(text, target_terminology, self.source_lang, self.target_lang)

                                        # 恢复公式
                                        if formulas:
                                            translation = self._restore_latex_formulas_with_formatting(translation, formulas)
                                else:
                                    # 普通文本，直接翻译
                                    if self.preprocess_terms:
                                        # 使用术语预处理方式翻译
                                        if self.source_lang == "zh":
                                            # 中文 → 外语
                                            # 检查是否有可用的术语
                                            if not used_terminology:
                                                # 如果没有术语，使用常规翻译
                                                logger.info("未找到匹配术语，使用常规翻译（带术语库）")
                                                translation = self.translator.translate_text(text, target_terminology, self.source_lang, self.target_lang)
                                            else:
                                                # 直接使用翻译器的内置术语处理功能
                                                logger.info(f"使用翻译器内置术语处理功能，找到 {len(used_terminology)} 个匹配术语")
                                                translation = self.translator.translate_text(text, used_terminology, self.source_lang, self.target_lang)
                                        else:
                                            # 外语 → 中文
                                            # 检查是否有可用的术语
                                            if not used_terminology:
                                                # 如果没有术语，使用常规翻译
                                                logger.info("未找到匹配术语，使用常规翻译")
                                                self.web_logger.info(f"No terminology found for regular text paragraph {global_idx + 1}, using regular translation")
                                                translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)
                                            else:
                                                # 使用term_extractor的占位符系统进行术语预处理
                                                # 参照Word文档处理器的流程
                                                # used_terminology已经是 {外语术语: 中文术语} 格式，直接使用
                                                reverse_terminology = used_terminology

                                                logger.info(f"使用术语预处理方式翻译，找到 {len(reverse_terminology)} 个匹配术语")
                                                self.web_logger.info(f"Using terminology preprocessing for regular text paragraph {global_idx + 1}, found {len(reverse_terminology)} terms")

                                                # 显示术语预处理的详细信息
                                                if reverse_terminology:
                                                    sample_terms = list(reverse_terminology.items())[:3]
                                                    logger.info(f"术语样本: {sample_terms}")
                                                    self.web_logger.info(f"Terminology sample: {sample_terms}")

                                                processed_text = self.term_extractor.replace_foreign_terms_with_placeholders(text, reverse_terminology)
                                                self.web_logger.info(f"Text after placeholder replacement: {processed_text[:100]}...")

                                                translated_with_placeholders = self.translator.translate_text(processed_text, None, self.source_lang, self.target_lang)
                                                self.web_logger.info(f"Translation with placeholders: {translated_with_placeholders[:100]}...")

                                                translation = self.term_extractor.restore_placeholders_with_chinese_terms(translated_with_placeholders)
                                                self.web_logger.info(f"Final translation: {translation[:100]}...")
                                    else:
                                        # 使用常规方式翻译
                                        translation = self.translator.translate_text(text, target_terminology, self.source_lang, self.target_lang)
                            except Exception as e:
                                logger.error(f"翻译段落时出错: {str(e)}")
                                self.web_logger.error(f"Translation failed for paragraph {global_idx + 1}: {str(e)}")
                                # 如果翻译失败，使用原文
                                translation = para_text
                                # 添加错误提示
                                translation += f"\n（翻译失败: {str(e)}）"

                            # 存储翻译结果
                            translated_paragraphs[global_idx] = {
                                'original': para_text,
                                'translated': translation,
                                'is_image_line': is_image_line
                            }

                            # 更新翻译进度
                            progress = 0.3 + (global_idx + 1) / len(text_paragraphs) * 0.4  # 30%-70%的进度用于翻译
                            self._update_progress(progress, f"已翻译 {global_idx + 1}/{len(text_paragraphs)} 个段落")

                        # 每批次处理完后添加延迟，避免API请求过快
                        time.sleep(0.5)

                    # 现在按照有序内容列表的顺序处理内容
                    logger.info("按照原始顺序处理内容...")

                    for item in ordered_content:
                        if item['type'] == 'text':
                            # 处理文本段落
                            para_idx = item['position']
                            if para_idx in translated_paragraphs:
                                para_data = translated_paragraphs[para_idx]
                                para_text = para_data['original']
                                translation = para_data['translated']
                                is_image_line = para_data['is_image_line']

                                if is_image_line:
                                    # 图片信息行，直接保留原格式
                                    img_para = doc.add_paragraph()
                                    img_run = img_para.add_run(para_text)
                                    img_run.italic = True
                                    img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

                                    logger.info(f"保留原格式的图片信息行: {para_text[:50]}")

                                    # 如果是双语对照模式，添加翻译
                                    if self.output_format == "bilingual":
                                        # 添加翻译
                                        trans_para = doc.add_paragraph()
                                        trans_run = trans_para.add_run(translation)
                                        trans_run.italic = True
                                        trans_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色

                                        # 收集翻译结果
                                        translation_results.append({
                                            'original': para_text,
                                            'translated': translation
                                        })

                                    # 添加空行
                                    doc.add_paragraph()
                                else:
                                    # 普通文本段落
                                    # 根据输出格式决定是否添加原文段落
                                    if self.output_format == "bilingual":
                                        # 双语对照模式：添加原文段落
                                        doc.add_paragraph(para_text)

                                    # 添加翻译段落
                                    trans_para = doc.add_paragraph()

                                    # 检查翻译结果中是否包含公式标记
                                    if "【数学公式】" in translation and "【公式结束】" in translation:
                                        # 包含公式，需要特殊处理
                                        parts = re.split(r'(【数学公式】.*?【公式结束】)', translation)

                                        for part in parts:
                                            if part.startswith("【数学公式】") and part.endswith("【公式结束】"):
                                                # 这是公式部分
                                                # 提取公式内容
                                                formula = part[len("【数学公式】"):-len("【公式结束】")]

                                                # 创建公式段落，使用更好的格式
                                                self._add_formatted_formula_to_paragraph(trans_para, formula)

                                                # 添加结束标记
                                                end_mark = trans_para.add_run("【公式结束】")
                                                end_mark.bold = True
                                                end_mark.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色
                                            else:
                                                # 这是普通文本部分
                                                text_run = trans_para.add_run(part)
                                                # 在双语对照模式下使用斜体区分译文
                                                if self.output_format == "bilingual":
                                                    text_run.italic = True
                                                text_run.font.size = Pt(11)  # 设置字体大小
                                    else:
                                        # 普通文本，直接添加
                                        trans_run = trans_para.add_run(translation)
                                        # 在双语对照模式下使用斜体区分译文
                                        if self.output_format == "bilingual":
                                            trans_run.italic = True
                                        trans_run.font.size = Pt(11)  # 设置字体大小

                                    # 收集翻译结果
                                    translation_results.append({
                                        'original': para_text,
                                        'translated': translation
                                    })

                                    # 添加空行
                                    doc.add_paragraph()
                        elif item['type'] == 'image':
                            # 处理图片
                            img_info = item['content']

                            # 创建段落用于插入图片
                            img_para = doc.add_paragraph()
                            img_run = img_para.add_run()

                            # 根据图片位置信息设置段落对齐方式
                            position = img_info.get('position', 'unknown')
                            if position == 'left_of_text':
                                img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif position == 'right_of_text':
                                img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # 计算合适的图片尺寸
                            original_width = img_info.get('width', 300)

                            # 限制最大宽度为5英寸
                            max_width = Inches(5)
                            width = min(Inches(original_width / 96), max_width)  # 假设96 DPI

                            # 获取图片路径
                            img_path = img_info.get('img_path', '')

                            # 从文件插入图片
                            try:
                                if img_path and os.path.exists(img_path):
                                    img_run.add_picture(img_path, width=width)
                                    logger.info(f"成功插入图片")
                                else:
                                    # 如果没有图片路径或文件不存在，使用占位符
                                    img_run.add_text(f"[图片]")
                                    img_run.bold = True
                                    img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                    logger.info(f"使用占位符替代图片")
                            except Exception as pic_err:
                                logger.error(f"插入图片时出错: {str(pic_err)}")
                                # 如果插入失败，使用占位符替代
                                img_run.add_text(f"[图片]")
                                img_run.bold = True
                                img_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

                            # 如果图片有标题，添加标题
                            if img_info.get('caption'):
                                caption_para = doc.add_paragraph()
                                caption_text = img_info.get('caption')
                                caption_run = caption_para.add_run(caption_text)
                                caption_run.italic = True
                                caption_run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                # 如果是双语对照模式，添加翻译
                                if self.output_format == "bilingual":
                                    try:
                                        # 翻译图片标题
                                        caption_translation = self.translator.translate_text(
                                            caption_text,
                                            target_terminology,
                                            self.source_lang,
                                            self.target_lang
                                        )

                                        # 添加翻译
                                        trans_para = doc.add_paragraph()
                                        trans_run = trans_para.add_run(caption_translation)
                                        trans_run.italic = True
                                        trans_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                                        trans_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                        # 收集翻译结果
                                        translation_results.append({
                                            'original': caption_text,
                                            'translated': caption_translation
                                        })
                                    except Exception as e:
                                        logger.error(f"翻译图片标题时出错: {str(e)}")

                            # 添加空行
                            doc.add_paragraph()

                        # 每批次处理完后添加延迟，避免API请求过快
                        time.sleep(0.5)

            # 更新进度：保存文档
            self._update_progress(0.8, "保存文档...")

            # 保存文档
            try:
                doc.save(output_path)
                logger.info(f"文件已保存到: {output_path}")
                self.web_logger.info(f"PDF translation completed, file saved to: {output_path}")
            except PermissionError:
                new_output_path = os.path.join(output_dir, f"{file_name}_带翻译_retry_{time_stamp}.docx")
                logger.warning(f"保存文件失败，尝试使用新文件名: {new_output_path}")
                self.web_logger.warning(f"File save failed, trying new filename: {new_output_path}")
                doc.save(new_output_path)
                logger.info(f"文件已保存到: {new_output_path}")
                self.web_logger.info(f"PDF translation completed, file saved to: {new_output_path}")
                output_path = new_output_path

            # 保存JSON结果
            json_output = os.path.join(output_dir, f"{file_name}_翻译结果_{time_stamp}.json")
            with open(json_output, 'w', encoding='utf-8') as f:
                json.dump({"status": "success", "file": output_path}, f, ensure_ascii=False, indent=2)

            # 更新进度：导出Excel
            self._update_progress(0.85, "导出翻译对照表...")

            # 翻译完成后，导出Excel文件
            if translation_results:
                self.export_to_excel(translation_results, file_path, target_language)

            # 如果需要导出PDF，则将Word文档转换为PDF
            if self.export_pdf:
                if docx2pdf_convert is not None:
                    # 更新进度：导出PDF
                    self._update_progress(0.9, "导出PDF文件...")

                    pdf_path = self.export_to_pdf(output_path)
                    logger.info(f"PDF文件已保存到: {pdf_path}")
                    # 更新JSON结果，添加PDF文件路径
                    with open(json_output, 'r', encoding='utf-8') as f:
                        json_data = json.load(f)
                    json_data["pdf_file"] = pdf_path
                    with open(json_output, 'w', encoding='utf-8') as f:
                        json.dump(json_data, f, ensure_ascii=False, indent=2)
                else:
                    logger.warning("PDF导出功能不可用，请安装docx2pdf模块")

            # 更新进度：完成
            self._update_progress(1.0, "翻译完成！")

            # 记录翻译完成的详细信息
            logger.info(f"PDF翻译任务完成！共处理 {len(translation_results)} 个翻译段落")
            self.web_logger.info(f"PDF translation task completed! Processed {len(translation_results)} translation segments")
            self.web_logger.info(f"Output file: {output_path}")

            return output_path

        except Exception as e:
            logger.error(f"PDF处理过程出错: {str(e)}")
            self.web_logger.error(f"PDF processing failed: {str(e)}")
            # 更新进度：出错
            self._update_progress(-1, f"翻译出错: {str(e)}")
            raise

    def _extract_latex_formulas(self, text: str) -> Tuple[str, List[Tuple[str, str]]]:
        """
        提取文本中的LaTeX公式，并用占位符替换

        Args:
            text: 原始文本

        Returns:
            Tuple[str, List[Tuple[str, str]]]: 替换后的文本和公式列表(占位符, 公式)
        """
        formulas = []
        processed_text = text

        # 首先检查文本中是否包含可能的LaTeX命令
        has_latex_commands = False
        for cmd_pattern in self.latex_commands:
            if re.search(cmd_pattern, processed_text):
                has_latex_commands = True
                logger.info(f"检测到可能的LaTeX命令: {cmd_pattern}")
                break

        # 使用更强大的公式识别逻辑
        # 1. 首先识别完整的环境公式（如\begin{equation}...\end{equation}）
        env_patterns = [
            (r'\\begin\{equation\*?\}(.*?)\\end\{equation\*?\}', '\\begin{equation}%s\\end{equation}'),
            (r'\\begin\{align\*?\}(.*?)\\end\{align\*?\}', '\\begin{align}%s\\end{align}'),
            (r'\\begin\{eqnarray\*?\}(.*?)\\end\{eqnarray\*?\}', '\\begin{eqnarray}%s\\end{eqnarray}'),
            (r'\\begin\{gather\*?\}(.*?)\\end\{gather\*?\}', '\\begin{gather}%s\\end{gather}'),
            (r'\\begin\{multline\*?\}(.*?)\\end\{multline\*?\}', '\\begin{multline}%s\\end{multline}'),
            (r'\\begin\{math\}(.*?)\\end\{math\}', '\\begin{math}%s\\end{math}'),
            (r'\\begin\{displaymath\}(.*?)\\end\{displaymath\}', '\\begin{displaymath}%s\\end{displaymath}'),
            (r'\\begin\{cases\}(.*?)\\end\{cases\}', '\\begin{cases}%s\\end{cases}')
        ]

        # 处理环境公式
        for i, (pattern, _) in enumerate(env_patterns):
            try:
                # 使用非贪婪匹配查找所有公式
                for match in re.finditer(pattern, processed_text, re.DOTALL):
                    full_match = match.group(0)
                    inner_content = match.group(1) if match.groups() else ""

                    # 创建唯一的占位符
                    placeholder = f"[LATEX_ENV_{i}_{len(formulas)}]"

                    # 记录找到的公式
                    logger.info(f"找到LaTeX环境公式: {full_match[:50]}{'...' if len(full_match) > 50 else ''}")

                    # 保存公式和占位符
                    formulas.append((placeholder, full_match))

                    # 替换文本中的公式为占位符
                    processed_text = processed_text.replace(full_match, placeholder, 1)
            except Exception as e:
                logger.error(f"处理LaTeX环境公式模式 {pattern} 时出错: {str(e)}")
                continue

        # 2. 识别行间公式（$$...$$和\[...\]）
        display_patterns = [
            (r'\$\$(.*?)\$\$', '$$%s$$'),
            (r'\\[(.*?)\\]', '\\[%s\\]')
        ]

        # 处理行间公式
        for i, (pattern, _) in enumerate(display_patterns):
            try:
                for match in re.finditer(pattern, processed_text, re.DOTALL):
                    full_match = match.group(0)
                    inner_content = match.group(1) if match.groups() else ""

                    # 创建唯一的占位符
                    placeholder = f"[LATEX_DISPLAY_{i}_{len(formulas)}]"

                    # 记录找到的公式
                    logger.info(f"找到LaTeX行间公式: {full_match[:50]}{'...' if len(full_match) > 50 else ''}")

                    # 保存公式和占位符
                    formulas.append((placeholder, full_match))

                    # 替换文本中的公式为占位符
                    processed_text = processed_text.replace(full_match, placeholder, 1)
            except Exception as e:
                logger.error(f"处理LaTeX行间公式模式 {pattern} 时出错: {str(e)}")
                continue

        # 3. 识别行内公式（$...$和\(...\)）
        inline_patterns = [
            (r'(?<!\$)\$([^\$]+?)\$(?!\$)', '$%s$'),  # 避免匹配到$$..$$
            (r'\\[(](.*?)\\[)]', '\\(%s\\)')
        ]

        # 处理行内公式
        for i, (pattern, _) in enumerate(inline_patterns):
            try:
                for match in re.finditer(pattern, processed_text, re.DOTALL):
                    full_match = match.group(0)
                    inner_content = match.group(1) if match.groups() else ""

                    # 检查是否是真正的公式（包含数学符号或LaTeX命令）
                    if not self._is_valid_formula(inner_content):
                        continue

                    # 创建唯一的占位符
                    placeholder = f"[LATEX_INLINE_{i}_{len(formulas)}]"

                    # 记录找到的公式
                    logger.info(f"找到LaTeX行内公式: {full_match[:50]}{'...' if len(full_match) > 50 else ''}")

                    # 保存公式和占位符
                    formulas.append((placeholder, full_match))

                    # 替换文本中的公式为占位符
                    processed_text = processed_text.replace(full_match, placeholder, 1)
            except Exception as e:
                logger.error(f"处理LaTeX行内公式模式 {pattern} 时出错: {str(e)}")
                continue

        # 4. 如果检测到LaTeX命令但没有找到完整公式，尝试识别独立的LaTeX命令
        if has_latex_commands and len(formulas) < 3:  # 如果找到的公式很少
            logger.info("尝试识别独立的LaTeX命令...")

            # 常见的数学命令模式
            command_patterns = [
                r'\\frac\{[^{}]*\}\{[^{}]*\}',  # \frac{...}{...}
                r'\\sum_\{[^{}]*\}(\^[^{}]*)?',  # \sum_{...}^{...}
                r'\\int_\{[^{}]*\}(\^[^{}]*)?',  # \int_{...}^{...}
                r'\\prod_\{[^{}]*\}(\^[^{}]*)?',  # \prod_{...}^{...}
                r'\\lim_\{[^{}]*\}',  # \lim_{...}
                r'\\sqrt\{[^{}]*\}',  # \sqrt{...}
                r'\\overrightarrow\{[^{}]*\}',  # \overrightarrow{...}
                r'\\mathbf\{[^{}]*\}',  # \mathbf{...}
                r'\\mathrm\{[^{}]*\}',  # \mathrm{...}
                r'\\mathcal\{[^{}]*\}',  # \mathcal{...}
                r'\\hat\{[^{}]*\}',  # \hat{...}
                r'\\vec\{[^{}]*\}'   # \vec{...}
            ]

            for i, pattern in enumerate(command_patterns):
                try:
                    for match in re.finditer(pattern, processed_text):
                        command = match.group(0)

                        # 检查这个命令是否已经在之前识别的公式中
                        is_in_formula = False
                        for _, formula in formulas:
                            if command in formula:
                                is_in_formula = True
                                break

                        if is_in_formula:
                            continue

                        # 创建唯一的占位符
                        placeholder = f"[LATEX_CMD_{i}_{len(formulas)}]"

                        # 记录找到的命令
                        logger.info(f"找到独立LaTeX命令: {command}")

                        # 将命令包装为行内公式
                        formula = f"${command}$"

                        # 保存公式和占位符
                        formulas.append((placeholder, formula))

                        # 替换文本中的命令为占位符
                        processed_text = processed_text.replace(command, placeholder, 1)
                except Exception as e:
                    logger.error(f"处理独立LaTeX命令 {pattern} 时出错: {str(e)}")
                    continue

        # 如果找到了公式，记录日志
        if formulas:
            logger.info(f"共找到 {len(formulas)} 个LaTeX公式")

        return processed_text, formulas

    def _is_valid_formula(self, content: str) -> bool:
        """
        检查内容是否是有效的LaTeX公式

        Args:
            content: 公式内容

        Returns:
            bool: 是否是有效的公式
        """
        # 如果内容为空，不是有效公式
        if not content or content.isspace():
            return False

        # 检查是否包含数学符号
        math_symbols = ['+', '-', '=', '>', '<', '\\times', '\\div', '\\cdot', '\\alpha', '\\beta', '\\gamma', '\\delta', '\\sum', '\\int', '\\prod', '\\frac', '\\sqrt']
        for symbol in math_symbols:
            if symbol in content:
                return True

        # 检查是否包含LaTeX命令
        if '\\' in content:
            return True

        # 检查是否包含上标或下标
        if '^' in content or '_' in content:
            return True

        # 检查是否只包含数字和简单符号（可能不是公式）
        if re.match(r'^[\d\s\.,]+$', content):
            return False

        # 默认情况下，如果内容看起来复杂，认为是公式
        return len(content) > 3

    def _restore_latex_formulas(self, text: str, formulas: List[Tuple[str, str]]) -> str:
        """
        将占位符替换回LaTeX公式，确保公式正确显示

        Args:
            text: 包含占位符的文本
            formulas: 公式列表(占位符, 公式)

        Returns:
            str: 恢复公式后的文本
        """
        result = text

        if not formulas:
            return result

        logger.info(f"正在恢复 {len(formulas)} 个LaTeX公式...")

        # 按照占位符长度降序排序，避免部分替换问题
        # 例如，如果有 [FORMULA_1] 和 [FORMULA_10]，应该先替换 [FORMULA_10]
        sorted_formulas = sorted(formulas, key=lambda x: len(x[0]), reverse=True)

        # 替换所有占位符为原始公式
        for placeholder, formula in sorted_formulas:
            # 确保公式是有效的LaTeX格式
            normalized_formula = self._normalize_latex_formula(formula)

            # 在Word文档中，公式需要特殊处理以确保正确显示
            if placeholder in result:
                # 检查占位符前后是否已有空格
                pattern = r'(\S)' + re.escape(placeholder) + r'(\S)'
                if re.search(pattern, result):
                    # 如果占位符前后有非空白字符，添加空格
                    result = re.sub(pattern, r'\1 ' + normalized_formula + r' \2', result)
                else:
                    # 否则直接替换
                    result = result.replace(placeholder, normalized_formula)

                logger.info(f"已恢复公式: {normalized_formula[:50]}{'...' if len(normalized_formula) > 50 else ''}")
            else:
                logger.warning(f"未找到占位符: {placeholder}")

        # 检查是否所有公式都已恢复
        missing_placeholders = []
        for placeholder, _ in formulas:
            if placeholder in result:
                missing_placeholders.append(placeholder)

        if missing_placeholders:
            logger.warning(f"以下占位符未被替换: {', '.join(missing_placeholders)}")

        # 最后检查并修复可能的格式问题
        result = self._fix_latex_formatting(result)

        return result

    def _restore_latex_formulas_with_formatting(self, text: str, formulas: List[Tuple[str, str]]) -> str:
        """
        将占位符替换回LaTeX公式，并添加特殊格式标记

        Args:
            text: 包含占位符的文本
            formulas: 公式列表(占位符, 公式)

        Returns:
            str: 恢复公式后的文本，带有特殊格式标记
        """
        result = text

        if not formulas:
            return result

        logger.info(f"正在恢复 {len(formulas)} 个LaTeX公式并添加特殊格式...")

        # 按照占位符长度降序排序，避免部分替换问题
        sorted_formulas = sorted(formulas, key=lambda x: len(x[0]), reverse=True)

        # 替换所有占位符为带格式的公式
        for placeholder, formula in sorted_formulas:
            # 确保公式是有效的LaTeX格式
            normalized_formula = self._normalize_latex_formula(formula)

            # 添加特殊格式标记
            formatted_formula = f"【数学公式】{normalized_formula}【公式结束】"

            # 替换占位符
            if placeholder in result:
                # 检查占位符前后是否已有空格
                pattern = r'(\S)' + re.escape(placeholder) + r'(\S)'
                if re.search(pattern, result):
                    # 如果占位符前后有非空白字符，添加空格
                    result = re.sub(pattern, r'\1 ' + formatted_formula + r' \2', result)
                else:
                    # 否则直接替换
                    result = result.replace(placeholder, formatted_formula)

                logger.info(f"已恢复公式并添加格式: {normalized_formula[:30]}...")
            else:
                logger.warning(f"未找到占位符: {placeholder}")

        return result

    def _normalize_latex_formula(self, formula: str) -> str:
        """
        规范化LaTeX公式格式

        Args:
            formula: 原始公式

        Returns:
            str: 规范化后的公式
        """
        # 去除首尾空白
        formula = formula.strip()

        # 检查公式类型
        is_display_math = False
        is_inline_math = False
        is_environment = False

        # 检查是否是行间公式
        if formula.startswith('$$') and formula.endswith('$$'):
            is_display_math = True
        elif formula.startswith('\\[') and formula.endswith('\\]'):
            is_display_math = True
        # 检查是否是行内公式
        elif formula.startswith('$') and formula.endswith('$') and not (formula.startswith('$$') or formula.endswith('$$')):
            is_inline_math = True
        elif formula.startswith('\\(') and formula.endswith('\\)'):
            is_inline_math = True
        # 检查是否是环境公式
        elif '\\begin{' in formula and '\\end{' in formula:
            is_environment = True

        # 如果没有明确的分隔符，添加行内公式分隔符
        if not (is_display_math or is_inline_math or is_environment):
            # 检查是否包含LaTeX命令
            if '\\' in formula:
                formula = f"${formula}$"
                is_inline_math = True

        # 确保公式周围有适当的空白
        if is_display_math:
            # 行间公式前后添加换行
            formula = f"\n{formula}\n"
        elif is_inline_math:
            # 行内公式前后添加空格
            formula = f" {formula} "
        elif is_environment:
            # 环境公式前后添加换行
            formula = f"\n{formula}\n"

        return formula

    def _fix_latex_formatting(self, text: str) -> str:
        """
        修复LaTeX格式问题

        Args:
            text: 包含LaTeX公式的文本

        Returns:
            str: 修复格式后的文本
        """
        # 修复连续的空格，但保留换行
        text = re.sub(r'[ \t]+', ' ', text)

        # 修复行间公式周围的空格
        text = re.sub(r'\s*\$\$\s*', '$$', text)
        text = re.sub(r'\$\$([^$]+?)\$\$', r'\n$$\1$$\n', text)

        # 修复环境公式周围的空格
        text = re.sub(r'\\begin\{([^}]+)\}', r'\n\\begin{\1}', text)
        text = re.sub(r'\\end\{([^}]+)\}', r'\\end{\1}\n', text)

        # 修复多余的换行
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 修复可能的LaTeX命令错误
        text = self._fix_latex_commands(text)

        return text

    def _fix_latex_commands(self, text: str) -> str:
        """
        修复LaTeX命令错误

        Args:
            text: 包含LaTeX命令的文本

        Returns:
            str: 修复后的文本
        """
        # 修复常见的LaTeX命令错误

        # 1. 修复花括号不匹配的问题
        # 计算左右花括号数量
        left_braces = text.count('{')
        right_braces = text.count('}')

        # 如果不匹配，尝试修复简单情况
        if left_braces > right_braces:
            # 缺少右花括号，添加到末尾
            text += '}' * (left_braces - right_braces)
        elif right_braces > left_braces:
            # 缺少左花括号，这种情况比较复杂，可能需要更复杂的逻辑
            pass

        # 2. 修复常见的命令拼写错误
        common_typos = {
            r'\\fration': r'\\fraction',
            r'\\alph': r'\\alpha',
            r'\\bata': r'\\beta',
            r'\\lamda': r'\\lambda',
            r'\\simga': r'\\sigma',
            r'\\thetha': r'\\theta',
            r'\\omaga': r'\\omega'
        }

        for typo, correction in common_typos.items():
            text = text.replace(typo, correction)

        # 3. 修复缺少空格的数学运算符
        # 例如，将"a+b"改为"a + b"以提高可读性
        operators = ['+', '-', '=', '<', '>', '\\times', '\\div', '\\cdot']
        for op in operators:
            if op in ['+', '-', '=', '<', '>']:
                # 简单运算符，添加空格
                text = re.sub(f'([^\\s{op}])\\{op}([^\\s{op}])', f'\\1 {op} \\2', text)
            else:
                # LaTeX命令运算符，添加空格
                text = re.sub(f'([^\\s]){op}([^\\s])', f'\\1 {op} \\2', text)

        # 4. 修复嵌套公式问题
        # 检测并修复嵌套的$符号
        text = self._fix_nested_delimiters(text, '$')

        # 5. 修复可能的环境嵌套问题
        text = self._fix_nested_environments(text)

        return text

    def _fix_nested_delimiters(self, text: str, delimiter: str) -> str:
        """
        修复嵌套的分隔符问题

        Args:
            text: 文本
            delimiter: 分隔符

        Returns:
            str: 修复后的文本
        """
        # 特殊处理$符号，因为它既用于行内公式又用于行间公式
        if delimiter == '$':
            # 首先处理行间公式$$...$$
            # 找到所有的$$...$$对
            display_matches = list(re.finditer(r'\$\$(.*?)\$\$', text, re.DOTALL))

            # 如果内部有单个$，替换为\$
            for match in reversed(display_matches):  # 从后向前处理，避免位置变化
                inner_content = match.group(1)
                if '$' in inner_content and not '$$' in inner_content:
                    fixed_content = inner_content.replace('$', '\\$')
                    text = text[:match.start()] + '$$' + fixed_content + '$$' + text[match.end():]

            # 然后处理行内公式$...$
            # 找到所有的$...$对
            inline_matches = []
            in_formula = False
            start_pos = -1

            for i, char in enumerate(text):
                if char == '$' and (i == 0 or text[i-1] != '\\'):  # 非转义的$
                    if not in_formula:
                        # 开始一个新公式
                        in_formula = True
                        start_pos = i
                    else:
                        # 结束当前公式
                        in_formula = False
                        inline_matches.append((start_pos, i))

            # 检查是否有未闭合的公式
            if in_formula:
                # 添加一个结束的$
                text += '$'
                inline_matches.append((start_pos, len(text) - 1))

            # 处理可能的嵌套问题
            for start, end in reversed(inline_matches):
                inner_content = text[start+1:end]
                if '$' in inner_content:
                    # 有嵌套的$，替换为\$
                    fixed_content = inner_content.replace('$', '\\$')
                    text = text[:start+1] + fixed_content + text[end:]

        return text

    def _fix_nested_environments(self, text: str) -> str:
        """
        修复嵌套的LaTeX环境问题

        Args:
            text: 文本

        Returns:
            str: 修复后的文本
        """
        # 查找所有的\begin{...}...\end{...}对
        begin_pattern = r'\\begin\{([^}]+)\}'
        end_pattern = r'\\end\{([^}]+)\}'

        begin_matches = list(re.finditer(begin_pattern, text))
        end_matches = list(re.finditer(end_pattern, text))

        # 如果begin和end数量不匹配，尝试修复
        if len(begin_matches) > len(end_matches):
            # 缺少\end，添加缺失的\end
            for i in range(len(end_matches), len(begin_matches)):
                env_name = begin_matches[i].group(1)
                text += f'\\end{{{env_name}}}'
        elif len(end_matches) > len(begin_matches):
            # 缺少\begin，这种情况比较复杂，可能需要更复杂的逻辑
            pass

        # 检查环境名称是否匹配
        stack = []
        for i, match in enumerate(begin_matches):
            env_name = match.group(1)
            stack.append((env_name, match.start()))

            # 如果有对应的\end
            if i < len(end_matches):
                end_env_name = end_matches[i].group(1)
                if stack and stack[-1][0] != end_env_name:
                    # 环境名称不匹配，修复\end的环境名称
                    correct_name = stack[-1][0]
                    end_pos = end_matches[i].start()
                    text = text[:end_pos] + f'\\end{{{correct_name}}}' + text[end_matches[i].end():]

                # 弹出栈顶元素
                if stack:
                    stack.pop()

        return text

    def _add_formatted_formula_to_paragraph(self, paragraph, formula: str):
        """
        向段落添加格式化的数学公式

        Args:
            paragraph: Word段落对象
            formula: 公式内容
        """
        try:
            # 添加公式开始标记
            start_mark = paragraph.add_run("【数学公式】")
            start_mark.bold = True
            start_mark.font.color.rgb = RGBColor(0, 100, 200)  # 蓝色
            start_mark.font.size = Pt(9)

            # 添加换行，让公式独占一行
            paragraph.add_run("\n")

            # 添加公式内容
            formula_run = paragraph.add_run(formula)

            # 设置公式的格式
            formula_run.font.name = "Cambria Math"  # 数学字体
            formula_run.font.size = Pt(11)
            formula_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

            # 检查公式类型并应用相应格式
            if formula.startswith('$$') and formula.endswith('$$'):
                # 行间公式，居中显示
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                formula_run.font.size = Pt(12)
            elif formula.startswith('$') and formula.endswith('$'):
                # 行内公式，正常对齐
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                formula_run.font.size = Pt(10)
            elif '\\begin{' in formula and '\\end{' in formula:
                # 环境公式，居中显示，稍大字体
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                formula_run.font.size = Pt(12)
                formula_run.bold = True
            else:
                # 其他公式，默认格式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                formula_run.font.size = Pt(10)

            # 添加换行
            paragraph.add_run("\n")

            # 添加公式结束标记
            end_mark = paragraph.add_run("【公式结束】")
            end_mark.bold = True
            end_mark.font.color.rgb = RGBColor(0, 100, 200)  # 蓝色
            end_mark.font.size = Pt(9)

            logger.info(f"已添加格式化公式: {formula[:30]}...")

        except Exception as e:
            logger.error(f"格式化公式时出错: {str(e)}")
            # 如果格式化失败，使用简单格式
            simple_run = paragraph.add_run(f"【数学公式】{formula}【公式结束】")
            simple_run.font.name = "Courier New"
            simple_run.font.size = Pt(10)

    def export_to_pdf(self, docx_path: str) -> str:
        """
        将Word文档转换为PDF

        Args:
            docx_path: Word文档路径

        Returns:
            str: PDF文件路径
        """
        try:
            logger.info(f"正在将Word文档转换为PDF: {docx_path}")

            # 构建PDF文件路径（与Word文档同名，但扩展名为.pdf）
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

            # 使用docx2pdf转换
            if docx2pdf_convert is None:
                raise ImportError("docx2pdf模块未安装，无法转换PDF")

            docx2pdf_convert(docx_path, pdf_path)

            logger.info(f"PDF文件已保存到: {pdf_path}")
            return pdf_path
        except Exception as e:
            logger.error(f"转换PDF失败: {str(e)}")
            raise Exception(f"转换PDF失败: {str(e)}")

    def export_to_excel(self, translation_results, original_file_path, target_language):
        """将翻译结果导出为Excel文件"""
        try:
            logger.info(f"正在导出翻译数据到Excel，共 {len(translation_results)} 条记录...")

            # 创建DataFrame
            df = pd.DataFrame({
                '原文': [item['original'] for item in translation_results],
                f'{target_language}翻译': [item['translated'] for item in translation_results]
            })

            # 生成保存路径
            file_name = os.path.basename(original_file_path)
            base_name = os.path.splitext(file_name)[0]
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

            # 使用与主文档相同的输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")

            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)

            # 构建Excel文件路径
            excel_path = os.path.join(output_dir, f"{base_name}_翻译对照表_{timestamp}.xlsx")

            # 保存到Excel
            df.to_excel(excel_path, index=False, engine='openpyxl')

            logger.info(f"翻译对照表已保存至: {excel_path}")
            return excel_path
        except Exception as e:
            logger.error(f"导出Excel文件时出错: {str(e)}")
            return None

    def _extract_and_optimize_text(self, page, page_number: int) -> str:
        """
        提取并优化PDF页面文本，保留原始文档结构

        Args:
            page: pdfplumber页面对象
            page_number: 页码（用于日志）

        Returns:
            str: 优化后的文本
        """
        try:
            # 基本文本提取
            raw_text = page.extract_text() or ""
            if not raw_text.strip():
                return ""

            logger.info(f"第 {page_number} 页原始文本长度: {len(raw_text)} 字符")

            # 检查是否存在文本提取问题（字符间有过多空格）
            char_count = len(raw_text.replace(" ", "").replace("\n", ""))
            space_count = raw_text.count(" ")

            # 如果是学术PDF，尝试更精确的文本提取
            try:
                # 获取页面上的所有文本元素，按照垂直位置排序
                chars = page.chars
                if chars:
                    # 按行分组文本元素
                    lines = self._group_text_by_lines(chars)

                    # 检查是否有可能的标题、段落等结构
                    structured_text = self._reconstruct_text_structure(lines)

                    # 如果结构化文本不为空且长度合理，使用它
                    if structured_text and len(structured_text) >= len(raw_text) * 0.7:
                        logger.info(f"使用结构化文本提取，长度: {len(structured_text)}")
                        return structured_text
            except Exception as struct_err:
                logger.warning(f"结构化文本提取失败: {str(struct_err)}")

            # 如果空格数量远大于字符数量，可能是逐字符提取问题
            if space_count > char_count * 0.5 and char_count > 0:
                logger.warning(f"第 {page_number} 页可能存在逐字符提取问题，尝试优化...")

                # 尝试使用更高级的文本提取方法
                try:
                    # 尝试使用表格提取
                    tables = page.extract_tables()
                    table_text = ""
                    if tables:
                        for table in tables:
                            for row in table:
                                table_text += " ".join([cell or "" for cell in row if cell]) + "\n"

                    # 尝试使用单词提取
                    words = page.extract_words()
                    words_text = " ".join([word.get('text', '') for word in words])

                    # 如果表格文本更长，使用表格文本
                    if len(table_text) > len(words_text) and table_text.strip():
                        logger.info(f"使用表格提取的文本，长度: {len(table_text)}")
                        raw_text = table_text
                    # 如果单词文本更长，使用单词文本
                    elif words_text.strip():
                        logger.info(f"使用单词提取的文本，长度: {len(words_text)}")
                        raw_text = words_text
                except Exception as extract_err:
                    logger.warning(f"尝试替代提取方法失败: {str(extract_err)}")

                # 优化文本
                optimized_text = self._optimize_extracted_text(raw_text)
                logger.info(f"第 {page_number} 页文本优化后长度: {len(optimized_text)} 字符")
                return optimized_text

            # 如果是英文PDF，尝试保留段落结构
            if self.source_lang == "en" and self.preserve_structure:
                # 保留原始换行符，但合并多个连续换行符为两个（表示段落分隔）
                structured_text = re.sub(r'\n{3,}', '\n\n', raw_text)
                return structured_text

            return raw_text

        except Exception as e:
            logger.error(f"提取第 {page_number} 页文本时出错: {str(e)}")
            return ""

    def _group_text_by_lines(self, chars):
        """
        将文本元素按行分组

        Args:
            chars: 页面上的文本元素列表

        Returns:
            list: 按行分组的文本元素
        """
        if not chars:
            return []

        # 按垂直位置排序
        sorted_chars = sorted(chars, key=lambda x: (x['top'], x['x0']))

        # 分组到行
        lines = []
        current_line = []
        current_top = sorted_chars[0]['top']

        for char in sorted_chars:
            # 如果垂直位置差异小于阈值，认为是同一行
            if abs(char['top'] - current_top) < 3:  # 3像素的容差
                current_line.append(char)
            else:
                # 开始新行
                if current_line:
                    lines.append(current_line)
                current_line = [char]
                current_top = char['top']

        # 添加最后一行
        if current_line:
            lines.append(current_line)

        return lines

    def _reconstruct_text_structure(self, lines):
        """
        重建文本结构

        Args:
            lines: 按行分组的文本元素

        Returns:
            str: 重建的结构化文本
        """
        if not lines:
            return ""

        # 分析行间距，识别段落
        result = []
        prev_line_bottom = 0

        for line in lines:
            # 按水平位置排序行内元素
            sorted_line = sorted(line, key=lambda x: x['x0'])

            # 提取行文本
            line_text = "".join([char['text'] for char in sorted_line])

            # 计算行间距
            if prev_line_bottom > 0:
                line_spacing = line[0]['top'] - prev_line_bottom

                # 大间距表示新段落
                if line_spacing > 10:  # 10像素的阈值，可以根据实际情况调整
                    result.append("\n\n")
                # 正常间距表示同一段落内的换行
                else:
                    result.append("\n")

            # 添加行文本
            result.append(line_text)

            # 更新前一行底部位置
            prev_line_bottom = max([char['bottom'] for char in line])

        return "".join(result)

    def _smart_paragraph_detection(self, lines):
        """
        智能段落识别

        Args:
            lines: 文本行列表

        Returns:
            list: 识别出的段落列表
        """
        if not lines:
            return []

        paragraphs = []
        current_paragraph = ""

        # 段落识别特征
        title_patterns = [
            # 标题模式：数字+点+空格开头（如 "1. 引言"）
            r'^\d+\.\s+\w+',
            # 标题模式：罗马数字+点+空格开头
            r'^[IVXivx]+\.\s+\w+',
            # 标题模式：大写字母开头的短行（可能是标题）
            r'^[A-Z][^.!?]{0,50}$',
            # 标题模式：Abstract, Introduction, Conclusion等常见章节名
            r'^(Abstract|Introduction|Conclusion|References|Bibliography|Acknowledgements|Appendix)(\s|:|$)'
        ]

        # 段落结束特征
        paragraph_end_patterns = [
            # 句子结束符
            r'[.!?]$',
            # 引用标记
            r'\[\d+\]$',
            # 空行
            r'^\s*$'
        ]

        for i, line in enumerate(lines):
            line = line.strip()

            # 空行表示段落分隔
            if not line:
                if current_paragraph:
                    paragraphs.append(current_paragraph)
                    current_paragraph = ""
                continue

            # 检查是否是新标题
            is_title = False
            for pattern in title_patterns:
                if re.search(pattern, line):
                    # 如果当前有段落内容，先保存
                    if current_paragraph:
                        paragraphs.append(current_paragraph)
                    # 将标题作为独立段落
                    paragraphs.append(line)
                    current_paragraph = ""
                    is_title = True
                    break

            if is_title:
                continue

            # 检查是否是段落结束
            is_paragraph_end = False
            for pattern in paragraph_end_patterns:
                if re.search(pattern, line):
                    is_paragraph_end = True
                    break

            # 处理普通行
            if current_paragraph:
                # 如果当前行很短且前一行是段落结束，可能是新段落开始
                if len(line) < 40 and is_paragraph_end and i < len(lines) - 1:
                    # 看下一行是否更长，如果是，可能是新段落的第一行
                    if len(lines[i+1].strip()) > len(line):
                        paragraphs.append(current_paragraph)
                        current_paragraph = line
                        continue

                # 否则，将当前行添加到当前段落
                current_paragraph += " " + line
            else:
                current_paragraph = line

            # 如果达到最大段落长度，保存并开始新段落
            if len(current_paragraph) >= self.max_paragraph_length:
                paragraphs.append(current_paragraph)
                current_paragraph = ""

        # 添加最后一个段落
        if current_paragraph:
            paragraphs.append(current_paragraph)

        return paragraphs

    def _optimize_extracted_text(self, text):
        """
        优化提取的文本

        Args:
            text: 原始提取的文本

        Returns:
            str: 优化后的文本
        """
        # 1. 移除过多的空格
        optimized_text = re.sub(r'\s+', ' ', text).strip()

        # 2. 根据源语言进行特定优化
        if self.source_lang == "en":
            # 英文文本优化：修复被错误分割的单词
            # 简单启发式：如果一个单词以小写字母结尾，下一个单词以小写字母开头，可能是被分割的
            optimized_text = re.sub(r'([a-z]) ([a-z])', r'\1\2', optimized_text)

            # 修复常见的连字符分割
            optimized_text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', optimized_text)

        elif self.source_lang == "zh":
            # 中文文本优化：移除中文字符之间的空格
            optimized_text = re.sub(r'([\u4e00-\u9fff])\s+([\u4e00-\u9fff])', r'\1\2', optimized_text)

        # 3. 保留段落结构（如果启用）
        if self.preserve_structure and '\n' in text:
            # 尝试恢复原始段落结构
            paragraphs = text.split('\n\n')
            if len(paragraphs) > 1:
                # 如果原文有明确的段落分隔，保留它们
                optimized_paragraphs = [re.sub(r'\s+', ' ', p).strip() for p in paragraphs if p.strip()]
                optimized_text = '\n\n'.join(optimized_paragraphs)

        return optimized_text

    def _export_used_terminology(self, used_terminology: Dict[str, str]) -> str:
        """导出使用的术语到Excel文件，包含使用统计信息"""
        try:
            logger.info(f"正在导出使用的术语到Excel，共 {len(used_terminology)} 个术语...")

            # 根据翻译方向创建不同的列名
            if self.source_lang == "zh":
                # 中文 → 外语
                source_column = '中文术语'
                target_column = '目标语术语'
            else:
                # 外语 → 中文
                source_column = '外语术语'
                target_column = '中文术语'

            # 获取术语使用统计信息
            usage_stats = self.term_extractor.get_terminology_usage_stats()

            # 如果有统计信息，使用更详细的导出格式
            if usage_stats:
                # 创建包含使用次数的DataFrame
                data = []
                for term, term_info in usage_stats.items():
                    data.append({
                        source_column: term_info['source'],
                        target_column: term_info['target'],
                        '使用次数': term_info['count']
                    })

                # 如果没有统计数据，使用基本的术语列表
                if not data:
                    for source, target in used_terminology.items():
                        data.append({
                            source_column: source,
                            target_column: target,
                            '使用次数': 'N/A'
                        })

                # 创建DataFrame
                df = pd.DataFrame(data)

                # 按使用次数降序排序
                if '使用次数' in df.columns and any(isinstance(x, int) for x in df['使用次数']):
                    df = df.sort_values(by='使用次数', ascending=False)
            else:
                # 创建基本的DataFrame（无使用次数）
                df = pd.DataFrame({
                    source_column: list(used_terminology.keys()),
                    target_column: list(used_terminology.values())
                })

            # 生成保存路径
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

            # 使用与主文档相同的输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")

            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)

            # 构建Excel文件路径
            excel_path = os.path.join(output_dir, f"使用的术语_{timestamp}.xlsx")

            # 保存到Excel
            df.to_excel(excel_path, index=False, engine='openpyxl')

            logger.info(f"使用的术语表已保存至: {excel_path}")
            return excel_path
        except Exception as e:
            logger.error(f"导出术语Excel文件时出错: {str(e)}")
            return None

    def _guess_image_extension(self, img):
        """
        尝试猜测图片的扩展名

        Args:
            img: PDF中的图片对象

        Returns:
            str: 猜测的图片扩展名（如.jpg, .png等）
        """
        # 尝试从图片对象中获取类型信息
        if 'colorspace' in img:
            colorspace = img['colorspace']
            if colorspace == 'DeviceRGB':
                return '.png'
            elif colorspace == 'DeviceCMYK':
                return '.jpg'
            elif colorspace == 'DeviceGray':
                return '.png'

        # 尝试从过滤器信息中猜测
        if 'filter' in img:
            filter_type = img['filter']
            if isinstance(filter_type, list):
                filter_type = filter_type[0] if filter_type else None

            if filter_type == 'DCTDecode':
                return '.jpg'
            elif filter_type == 'FlateDecode':
                return '.png'
            elif filter_type == 'JPXDecode':
                return '.jp2'
            elif filter_type == 'CCITTFaxDecode':
                return '.tiff'

        # 默认返回PNG
        return '.png'

    def _extract_image_info(self, page) -> List[Dict]:
        """
        提取页面中的图片信息，包括图片与文本的相对位置关系

        Args:
            page: PDF页面对象

        Returns:
            List[Dict]: 图片信息列表，每个字典包含图片的详细位置信息和上下文关系
        """
        try:
            images = []
            # 获取页面中的图片对象
            for img in page.images:
                # 提取图片位置信息
                img_info = {
                    'x0': img['x0'],
                    'y0': img['y0'],
                    'x1': img['x1'],
                    'y1': img['y1'],
                    'width': img['width'],
                    'height': img['height'],
                    'type': 'image',
                    'is_figure': True,  # 标记为图片
                    'image_data': img.get('stream', None),  # 尝试获取图片数据流
                    'image_ext': self._guess_image_extension(img),  # 尝试猜测图片扩展名
                    'position': 'unknown',  # 初始化位置关系
                    'context': {}  # 初始化上下文关系
                }
                images.append(img_info)
                logger.info(f"提取到图片: 宽度={img['width']}, 高度={img['height']}, 位置=({img['x0']},{img['y0']})-({img['x1']},{img['y1']})")

            # 获取页面文本和布局信息
            try:
                # 获取页面文本
                page_text = page.extract_text() or ""
                # 获取页面上的所有文本块
                text_blocks = page.extract_words(x_tolerance=3, y_tolerance=3, keep_blank_chars=True)
                # 获取页面上的所有文本行
                lines = self._group_text_by_lines(page.chars)

                # 分析图片与文本的相对位置关系
                for i, img_info in enumerate(images):
                    if img_info['type'] != 'image':
                        continue

                    # 计算图片的中心点
                    img_center_x = (img_info['x0'] + img_info['x1']) / 2
                    img_center_y = (img_info['y0'] + img_info['y1']) / 2

                    # 查找与图片最接近的文本块
                    closest_blocks = []
                    for block in text_blocks:
                        block_center_x = (block['x0'] + block['x1']) / 2
                        block_center_y = (block['top'] + block['bottom']) / 2

                        # 计算距离
                        distance = ((img_center_x - block_center_x) ** 2 +
                                   (img_center_y - block_center_y) ** 2) ** 0.5

                        closest_blocks.append((block, distance))

                    # 按距离排序
                    closest_blocks.sort(key=lambda x: x[1])

                    # 确定图片位置（在文本上方、下方、左侧、右侧或嵌入）
                    if closest_blocks:
                        closest_block = closest_blocks[0][0]

                        # 确定相对位置
                        if img_info['y1'] < closest_block['top']:
                            position = 'above_text'  # 图片在文本上方
                        elif img_info['y0'] > closest_block['bottom']:
                            position = 'below_text'  # 图片在文本下方
                        elif img_info['x1'] < closest_block['x0']:
                            position = 'left_of_text'  # 图片在文本左侧
                        elif img_info['x0'] > closest_block['x1']:
                            position = 'right_of_text'  # 图片在文本右侧
                        else:
                            position = 'embedded'  # 图片嵌入在文本中

                        img_info['position'] = position

                        # 记录上下文信息
                        img_info['context'] = {
                            'closest_text': closest_block['text'],
                            'distance': closest_blocks[0][1],
                            'relative_position': position
                        }

                        logger.info(f"图片 {i+1} 位置关系: {position}, 最近文本: {closest_block['text'][:30]}...")

                    # 查找图片可能的标题（通常在图片下方）
                    potential_captions = []
                    for line_idx, line in enumerate(lines):
                        if not line:
                            continue

                        # 计算行的中心点
                        line_chars = sorted(line, key=lambda x: x['x0'])
                        if not line_chars:
                            continue

                        line_y0 = min(char['top'] for char in line_chars)
                        line_y1 = max(char['bottom'] for char in line_chars)
                        line_center_y = (line_y0 + line_y1) / 2

                        # 检查是否在图片下方且水平位置接近
                        if (line_center_y > img_info['y1'] and
                            abs(img_center_x - sum(char['x0'] for char in line_chars) / len(line_chars)) < img_info['width']):

                            # 提取行文本
                            line_text = "".join(char['text'] for char in sorted(line_chars, key=lambda x: x['x0']))

                            # 检查是否包含图片相关关键词
                            image_keywords = ["图", "Figure", "图片", "Image", "图表", "Chart", "插图", "Illustration"]
                            for keyword in image_keywords:
                                if keyword in line_text and len(line_text) < 200:
                                    potential_captions.append({
                                        'text': line_text,
                                        'line_number': line_idx,
                                        'distance': line_center_y - img_info['y1']
                                    })
                                    break

                    # 如果找到可能的标题，添加到图片信息中
                    if potential_captions:
                        # 按距离排序，选择最近的
                        potential_captions.sort(key=lambda x: x['distance'])
                        img_info['caption'] = potential_captions[0]['text']
                        img_info['caption_line'] = potential_captions[0]['line_number']
                        logger.info(f"图片 {i+1} 可能的标题: {img_info['caption']}")

                # 查找可能包含图片的区域（通过文本内容判断）
                if page_text:
                    # 查找可能的图片标题行
                    image_keywords = ["图", "Figure", "图片", "Image", "图表", "Chart", "插图", "Illustration"]
                    text_lines = page_text.split('\n')

                    for i, line in enumerate(text_lines):
                        for keyword in image_keywords:
                            if keyword in line and len(line) < 200:  # 图片说明通常较短
                                # 找到可能的图片标题行
                                # 检查是否已经作为某个图片的标题被识别
                                is_duplicate = False
                                for img in images:
                                    if img.get('caption') == line:
                                        is_duplicate = True
                                        break

                                if not is_duplicate:
                                    img_info = {
                                        'text': line,
                                        'line_number': i,
                                        'type': 'figure_caption',
                                        'is_figure': True,  # 标记为图片相关内容
                                        'position': 'standalone_caption'  # 独立的标题
                                    }
                                    images.append(img_info)
                                    logger.info(f"提取到独立图片标题: {line[:50]}")
                                break
            except Exception as text_err:
                logger.warning(f"分析图片与文本位置关系时出错: {str(text_err)}")
                logger.exception(text_err)

            # 记录提取到的图片总数
            if images:
                logger.info(f"总共提取到 {len(images)} 个图片相关元素")

            return images
        except Exception as e:
            logger.error(f"提取图片信息时出错: {str(e)}")
            logger.exception(e)
            return []

    def _detect_image_format(self, image_bytes: bytes) -> str:
        """检测图片格式"""
        if not image_bytes or len(image_bytes) < 8:
            return "unknown"

        # 检查常见图片格式的文件头
        if image_bytes.startswith(b'\x89PNG\r\n\x1a\n'):
            return "PNG"
        elif image_bytes.startswith(b'\xff\xd8\xff'):
            return "JPEG"
        elif image_bytes.startswith(b'GIF87a') or image_bytes.startswith(b'GIF89a'):
            return "GIF"
        elif image_bytes.startswith(b'BM'):
            return "BMP"
        elif image_bytes.startswith(b'II*\x00') or image_bytes.startswith(b'MM\x00*'):
            return "TIFF"
        elif image_bytes.startswith(b'RIFF') and b'WEBP' in image_bytes[:12]:
            return "WEBP"
        else:
            return "unknown"

    def _fix_image_data(self, image_bytes: bytes) -> bytes:
        """尝试修复图片数据"""
        if not image_bytes:
            return image_bytes

        # 尝试添加常见的图片头部
        headers_to_try = [
            (b'\x89PNG\r\n\x1a\n', "PNG"),
            (b'\xff\xd8\xff\xe0', "JPEG"),
            (b'GIF89a', "GIF"),
            (b'BM', "BMP")
        ]

        for header, format_name in headers_to_try:
            try:
                # 创建带有正确头部的图片数据
                fixed_data = header + image_bytes[len(header):]

                # 尝试用PIL验证
                from PIL import Image
                import io
                with Image.open(io.BytesIO(fixed_data)) as test_img:
                    test_img.verify()

                logger.info(f"成功修复图片数据为{format_name}格式")
                return fixed_data
            except Exception:
                continue

        # 如果无法修复，返回原始数据
        return image_bytes

    def _create_image_placeholder(self, run, img_index: int, img_info: Dict):
        """创建图片占位符"""
        try:
            # 获取图片尺寸信息
            width = img_info.get('width', 300)
            height = img_info.get('height', 200)

            # 创建图片占位符文本
            placeholder_text = f"[图片 {img_index}]"

            # 添加占位符文本
            run.add_text(placeholder_text)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # 蓝色

            # 添加图片信息
            if width and height:
                info_text = f" (尺寸: {int(width)}×{int(height)})"
                run.add_text(info_text)
                run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色

            logger.info(f"创建了图片占位符: {placeholder_text}")

        except Exception as e:
            logger.error(f"创建图片占位符失败: {str(e)}")
            # 最简单的占位符
            run.add_text(f"[图片 {img_index}]")
