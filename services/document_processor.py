import os
import logging
import time
import json
import re
import traceback
from docx import Document
from docx.shared import RGBColor
from typing import Dict, List, Any, Tuple
from .translator import TranslationService
import pandas as pd
from datetime import datetime
from utils.term_extractor import TermExtractor
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
    logger.warning("docx2pdf模块未安装，DOCX转PDF功能将不可用")

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, translator: TranslationService):
        self.translator = translator
        self.use_terminology = True  # 默认使用术语库
        self.preprocess_terms = True  # 是否预处理术语（默认启用）
        self.term_extractor = TermExtractor()  # 术语提取器
        self.export_pdf = False  # 是否导出PDF
        self.source_lang = "zh"  # 默认源语言为中文
        self.target_lang = "en"  # 默认目标语言为英文
        self.is_cn_to_foreign = True  # 默认翻译方向为中文→外语
        self.progress_callback = None  # 进度回调函数
        self.retry_count = 3  # 翻译失败重试次数
        self.retry_delay = 1  # 重试延迟（秒）
        self.output_format = "bilingual"  # 输出格式：bilingual（双语）或translation_only（仅翻译）

        # 配置日志记录器
        self.logger = logging.getLogger(__name__)
        self.web_logger = logging.getLogger('web_logger')

        # 初始化日志配置
        self.logger.info("DocumentProcessor initialized")
        self.web_logger.info("Translation service ready")
        # 数学公式正则表达式模式
        self.latex_patterns = [
            r'\$\$(.*?)\$\$',  # 行间公式 $$...$$
            r'\$(.*?)\$',      # 行内公式 $...$
            r'\\begin\{equation\}(.*?)\\end\{equation\}',  # equation环境
            r'\\begin\{align\}(.*?)\\end\{align\}',        # align环境
            r'\\begin\{eqnarray\}(.*?)\\end\{eqnarray\}'   # eqnarray环境
        ]

    def set_progress_callback(self, callback):
        """设置进度回调函数"""
        self.progress_callback = callback

    def process_document(self, file_path: str, target_language: str, terminology: Dict, source_lang: str = "zh", target_lang: str = "en") -> str:
        """
        处理文档翻译

        Args:
            file_path: 文件路径
            target_language: 目标语言名称（用于术语表查找）
            terminology: 术语词典
            source_lang: 源语言代码，默认为中文(zh)
            target_lang: 目标语言代码，默认为英文(en)

        Returns:
            str: 输出文件路径
        """
        # 更新进度：开始处理
        self._update_progress(0.01, "开始处理文档...")

        # 设置翻译方向
        self.source_lang = source_lang
        self.target_lang = target_lang

        # 根据源语言和目标语言确定翻译方向
        self.is_cn_to_foreign = source_lang == "zh" or (source_lang == "auto" and target_lang != "zh")
        logger.info(f"翻译方向: {'中文→外语' if self.is_cn_to_foreign else '外语→中文'}")

        # 检查文件是否存在
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            raise Exception("文件不存在")

        # 在程序根目录下创建输出目录
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 获取输出路径，添加时间戳避免文件名冲突
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        time_stamp = time.strftime("%Y%m%d%H%M%S")
        output_file_name = f"{file_name}_带翻译_{time_stamp}.docx"
        output_path = os.path.join(output_dir, output_file_name)

        # 更新进度：检查文件
        self._update_progress(0.05, "检查文件权限...")

        # 复制并处理文档
        try:
            # 检查文件是否可写入
            try:
                with open(file_path, 'rb') as f:
                    # 可以打开源文件
                    pass
            except PermissionError:
                logger.error(f"无法读取源文件，文件可能被占用: {file_path}")
                raise Exception(f"无法读取源文件，文件可能被占用。请确保文件未被其他程序（如Word）打开。")

            # 检查目标目录是否可写入
            try:
                test_file = os.path.join(output_dir, "__test_write__.tmp")
                with open(test_file, 'w') as f:
                    f.write("test")
                os.remove(test_file)
            except (PermissionError, IOError) as e:
                logger.error(f"无法写入输出目录: {output_dir}, 错误: {str(e)}")
                raise Exception(f"无法写入输出目录，请检查权限或以管理员身份运行程序。")

            # 更新进度：复制文档
            self._update_progress(0.1, "复制文档...")
            self.web_logger.info(f"Starting to copy document: {os.path.basename(file_path)}")

            # 创建Word文档
            doc = self._copy_document(file_path, output_path)
            logger.info(f"成功复制原文档到: {output_path}")
            self.web_logger.info(f"Document copied successfully to: {os.path.basename(output_path)}")
        except Exception as e:
            logger.error(f"复制文档失败: {str(e)}")
            raise Exception(f"复制文档失败: {str(e)}")

        # 更新进度：加载术语库
        self._update_progress(0.15, "加载术语库...")

        # 获取选定语言的术语表
        target_terminology = terminology.get(target_language, {})
        logger.info(f"使用 {target_language} 术语表，包含 {len(target_terminology)} 个术语")

        # 如果术语表为空但术语库不为空，尝试使用其他方式获取术语
        if not target_terminology and terminology:
            # 检查是否有英语术语表
            if "英语" in terminology and target_language == "en":
                target_terminology = terminology["英语"]
                logger.info(f"使用英语术语表替代，包含 {len(target_terminology)} 个术语")
            # 检查是否有日语术语表
            elif "日语" in terminology and target_language == "ja":
                target_terminology = terminology["日语"]
                logger.info(f"使用日语术语表替代，包含 {len(target_terminology)} 个术语")
            # 检查是否有韩语术语表
            elif "韩语" in terminology and target_language == "ko":
                target_terminology = terminology["韩语"]
                logger.info(f"使用韩语术语表替代，包含 {len(target_terminology)} 个术语")
            # 检查是否有德语术语表
            elif "德语" in terminology and target_language == "de":
                target_terminology = terminology["德语"]
                logger.info(f"使用德语术语表替代，包含 {len(target_terminology)} 个术语")
            # 检查是否有法语术语表
            elif "法语" in terminology and target_language == "fr":
                target_terminology = terminology["法语"]
                logger.info(f"使用法语术语表替代，包含 {len(target_terminology)} 个术语")
            # 检查是否有西班牙语术语表
            elif "西班牙语" in terminology and target_language == "es":
                target_terminology = terminology["西班牙语"]
                logger.info(f"使用西班牙语术语表替代，包含 {len(target_terminology)} 个术语")
            else:
                logger.warning(f"无法找到匹配的术语表，将使用空术语表")

        # 创建一个列表来收集翻译结果
        translation_results = []

        try:
            # 更新进度：处理段落（与C#版本保持一致，先处理段落）
            self._update_progress(0.2, "处理文档段落...")

            # 处理段落
            self._process_paragraphs(doc, target_terminology, translation_results)

            # 更新进度：处理表格
            self._update_progress(0.4, "处理文档表格...")

            # 处理表格
            self._process_tables(doc, target_terminology, translation_results)

            # 更新进度：保存文档
            self._update_progress(0.8, "保存文档...")

            # 保存文档
            try:
                doc.save(output_path)
                logger.info(f"文件已保存到: {output_path}")
            except PermissionError:
                new_output_path = os.path.join(output_dir, f"{file_name}_带翻译_retry_{time_stamp}.docx")
                logger.warning(f"保存文件失败，尝试使用新文件名: {new_output_path}")
                doc.save(new_output_path)
                logger.info(f"文件已保存到: {new_output_path}")
                output_path = new_output_path

            # 保存JSON结果
            json_output = os.path.join(output_dir, f"{file_name}_翻译结果_{time_stamp}.json")
            with open(json_output, 'w', encoding='utf-8') as f:
                json.dump({"status": "success", "file": output_path}, f, ensure_ascii=False, indent=2)

            # 更新进度：导出Excel
            self._update_progress(0.85, "导出翻译对照表...")

            # 翻译完成后，导出Excel文件
            if translation_results:
                self.export_to_excel(translation_results, file_path, target_language)

            # 如果需要导出PDF，则将Word文档转换为PDF
            if self.export_pdf:
                # 更新进度：导出PDF
                self._update_progress(0.9, "导出PDF文件...")

                pdf_path = self.export_to_pdf(output_path)
                logger.info(f"PDF文件已保存到: {pdf_path}")
                # 更新JSON结果，添加PDF文件路径
                with open(json_output, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
                json_data["pdf_file"] = pdf_path
                with open(json_output, 'w', encoding='utf-8') as f:
                    json.dump(json_data, f, ensure_ascii=False, indent=2)

            # 更新进度：完成
            self._update_progress(1.0, "翻译完成！")

            return output_path

        except Exception as e:
            logger.error(f"翻译过程出错: {str(e)}")
            # 更新进度：出错
            self._update_progress(-1, f"翻译出错: {str(e)}")
            raise

    def _copy_document(self, source_path: str, target_path: str) -> Document:
        """复制文档"""
        import threading

        def load_document():
            """在单独线程中加载文档"""
            try:
                self.logger.info(f"开始加载Word文档: {source_path}")
                self.web_logger.info(f"Loading Word document: {os.path.basename(source_path)}")

                # 检查文件大小
                file_size = os.path.getsize(source_path)
                self.logger.info(f"文档大小: {file_size / (1024*1024):.2f} MB")

                if file_size > 50 * 1024 * 1024:  # 50MB
                    self.logger.warning("文档较大，加载可能需要较长时间")
                    self.web_logger.info("Large document detected, loading may take longer...")

                # 加载文档
                doc = Document(source_path)
                self.logger.info("文档加载成功，开始保存副本")
                self.web_logger.info("Document loaded successfully, saving copy...")

                # 保存副本
                doc.save(target_path)
                self.logger.info(f"文档副本已保存到: {target_path}")

                # 重新加载副本以确保数据完整性
                return Document(target_path)

            except Exception as e:
                self.logger.error(f"加载文档时出错: {str(e)}")
                self.web_logger.error(f"Error loading document: {str(e)}")
                raise

        try:
            # 使用线程加载文档，避免阻塞主线程
            result = [None]
            exception = [None]

            def worker():
                try:
                    result[0] = load_document()
                except Exception as e:
                    exception[0] = e

            thread = threading.Thread(target=worker)
            thread.daemon = True
            thread.start()

            # 等待最多60秒
            thread.join(timeout=60)

            if thread.is_alive():
                self.logger.error("文档加载超时（60秒）")
                self.web_logger.error("Document loading timeout (60 seconds)")
                raise TimeoutError("文档加载超时，请检查文档是否被其他程序占用或文档过大")

            if exception[0]:
                raise exception[0]

            if result[0] is None:
                raise Exception("文档加载失败，未知错误")

            return result[0]

        except Exception as e:
            error_msg = str(e)
            if "PermissionError" in error_msg or "Permission denied" in error_msg:
                raise Exception("无法访问文档文件，请确保文件未被Word或其他程序打开，或以管理员身份运行程序")
            elif "TimeoutError" in error_msg or "timeout" in error_msg.lower():
                raise Exception("文档加载超时，可能是文档过大或被其他程序占用，请稍后重试")
            elif "BadZipFile" in error_msg or "zipfile" in error_msg.lower():
                raise Exception("文档文件损坏或格式不正确，请检查文件是否为有效的Word文档")
            else:
                raise Exception(f"加载文档失败: {error_msg}")

    def _process_tables(self, doc: Document, terminology: Dict, translation_results: list) -> None:
        """处理文档中的表格"""
        self.logger.info("开始处理文档表格")
        self.web_logger.info("Processing document tables...")

        # 如果启用了术语预处理，先收集所有使用的术语
        used_terminology = {}
        if self.preprocess_terms:
            self.logger.info("=== 表格术语预处理已启用，开始收集术语 ===")
            self.logger.info(f"翻译方向: {self.source_lang} -> {self.target_lang}")
            self.logger.info(f"术语库大小: {len(terminology)} 个术语")

            # 显示术语库样本
            if terminology:
                sample_terms = list(terminology.items())[:5]
                self.logger.info(f"术语库样本（前5个）: {sample_terms}")

            # 对于外语→中文翻译模式，预先将术语库键值对调并缓存
            if self.source_lang != "zh":
                self.logger.info("外语→中文翻译模式，预先对调术语库键值并缓存...")
                self.reversed_terminology = self._create_reversed_terminology(terminology)
                self.logger.info(f"对调后的术语库大小: {len(self.reversed_terminology)} 个术语")

                # 显示对调后的术语库样本
                if self.reversed_terminology:
                    reversed_sample = list(self.reversed_terminology.items())[:5]
                    self.logger.info(f"对调后术语库样本（前5个）: {reversed_sample}")
            else:
                self.reversed_terminology = None

            # 先遍历所有表格内容，收集使用的术语
            table_count = len(doc.tables)
            self.logger.info(f"开始分析 {table_count} 个表格中的术语...")

            for table_idx, table in enumerate(doc.tables, 1):
                self.web_logger.info(f"Processing table {table_idx}/{table_count}")
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            try:
                                # 根据翻译方向选择不同的术语提取方法
                                if self.source_lang == "zh":
                                    # 中文 → 外语
                                    self.logger.info(f"从中文表格单元格提取术语: {cell.text[:50]}...")
                                    cell_terms = self.term_extractor.extract_terms(cell.text, terminology)
                                    self.logger.info(f"从中文文本提取术语: {len(cell_terms)} 个")
                                else:
                                    # 外语 → 中文，使用缓存的反向术语库
                                    self.logger.info(f"从外语表格单元格提取术语: {cell.text[:50]}...")
                                    if hasattr(self, 'reversed_terminology') and self.reversed_terminology:
                                        # 使用缓存的反向术语库进行高效匹配
                                        cell_terms = self.term_extractor.extract_foreign_terms_from_reversed_dict(cell.text, self.reversed_terminology)
                                        self.logger.info(f"从外语文本提取术语（使用缓存）: {len(cell_terms)} 个")
                                    else:
                                        # 回退到原始方法
                                        cell_terms = self.term_extractor.extract_foreign_terms_by_chinese_values(cell.text, terminology)
                                        self.logger.info(f"从外语文本提取术语（原始方法）: {len(cell_terms)} 个")

                                # 显示提取到的术语
                                if cell_terms:
                                    self.logger.info(f"表格单元格提取到的术语: {list(cell_terms.items())[:3]}")

                                # 更新使用的术语词典
                                used_terminology.update(cell_terms)
                            except Exception as e:
                                self.logger.error(f"术语提取失败: {str(e)}")
                                self.web_logger.error(f"Term extraction failed: {str(e)}")

            self.logger.info(f"从表格中提取了 {len(used_terminology)} 个术语")
            self.web_logger.info(f"Extracted {len(used_terminology)} terms from tables")

        # 处理表格翻译（与C#版本保持一致，直接遍历所有单元格）
        total_tables = len(doc.tables)
        logger.info(f"开始处理 {total_tables} 个表格")

        # 收集所有单元格
        all_cells = []
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows, 1):
                for cell_idx, cell in enumerate(row.cells, 1):
                    if cell.text.strip():
                        all_cells.append({
                            'cell': cell,
                            'table_idx': table_idx,
                            'row_idx': row_idx,
                            'cell_idx': cell_idx
                        })

        logger.info(f"共找到 {len(all_cells)} 个非空单元格")

        # 处理每个单元格
        for cell_info in all_cells:
            cell = cell_info['cell']
            table_idx = cell_info['table_idx']
            row_idx = cell_info['row_idx']
            cell_idx = cell_info['cell_idx']

            # 更完整地提取单元格文本，确保不遗漏内容
            cell_text = self._extract_complete_cell_text(cell).strip()
            logger.debug(f"检查表格 {table_idx} 行 {row_idx} 列 {cell_idx}: '{cell_text}'")

            # 验证文本提取的完整性
            if not self._validate_cell_text_completeness(cell, cell_text):
                logger.warning(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 文本提取可能不完整，尝试重新提取")
                # 尝试使用备用方法提取
                alternative_text = cell.text.strip()
                if len(alternative_text) > len(cell_text):
                    logger.info(f"使用备用方法获得更完整的文本: '{alternative_text}'")
                    cell_text = alternative_text

            # 使用诊断方法进行详细的翻译决策分析
            diagnosis = self._diagnose_cell_translation_decision(cell_text, table_idx, row_idx, cell_idx)

            # 如果不需要翻译，跳过
            if not diagnosis['should_translate']:
                continue

            logger.info(f"正在翻译表格 {table_idx} 行 {row_idx} 列 {cell_idx}: {cell_text[:50]}...")
            logger.info(f"单元格包含 {len(cell.paragraphs)} 个段落")

            # 先保存原始格式信息，再根据模式处理内容
            original_format = self._save_format_info(cell.paragraphs)

            if self.output_format == "translation_only":
                # 仅翻译模式：保存原文并清空单元格
                original_text = cell_text
                for para in cell.paragraphs:
                    para.clear()
            else:
                # 双语对照模式：保持原文不变
                original_text = cell_text

            # 检查单元格中是否包含数学公式
            # 使用原始文本进行公式提取，确保不遗漏内容
            text, formulas = self._extract_latex_formulas(original_text)

            # 检查是否需要翻译（数值、单位等可能不需要翻译）
            if self._should_skip_translation(text):
                logger.info(f"单元格内容无需翻译: {text}")
                translation = text  # 保持原文
            else:
                # 翻译单元格内容（不包含公式部分）
                translation = ""
                logger.info(f"开始翻译单元格内容: {text[:50]}...")
                logger.info(f"preprocess_terms: {self.preprocess_terms}")
                logger.info(f"source_lang: {self.source_lang}")
                logger.info(f"target_lang: {self.target_lang}")

                # 使用带重试机制的翻译方法
                translation = self._translate_cell_with_retry(text, terminology, table_idx, row_idx, cell_idx)

            try:
                # 将公式重新插入到翻译后的文本中
                if formulas:
                    translation = self._restore_latex_formulas(translation, formulas)

                logger.info(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译完成: {translation[:50]}...")

                # 验证翻译结果质量
                location = f"表格 {table_idx} 行 {row_idx} 列 {cell_idx}"
                validation_issues = self._validate_translation_result(original_text, translation, location)
                if validation_issues:
                    for issue in validation_issues:
                        logger.warning(f"翻译质量问题: {issue}")
                        self.web_logger.warning(f"Translation quality issue: {issue}")

                # 收集翻译结果
                translation_results.append({
                    'original': cell_text if self.output_format == "bilingual" else original_text,
                    'translated': translation,
                    'location': location,
                    'validation_issues': validation_issues  # 添加验证问题信息
                })

                # 使用C#版本一致的方法更新单元格文本
                self._update_table_cell_text(cell.paragraphs, original_text, translation)
                logger.info(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 处理完成")

            except Exception as e:
                logger.error(f"处理表格 {table_idx} 行 {row_idx} 列 {cell_idx} 的翻译结果时失败: {str(e)}")
                # 即使添加翻译失败，也要继续处理下一个单元格
                continue

    def _process_paragraphs(self, doc: Document, terminology: Dict, translation_results: list) -> None:
        """处理文档中的段落"""
        self.logger.info("开始处理文档段落")
        self.web_logger.info("Processing document paragraphs...")

        # 如果启用了术语预处理，先收集所有使用的术语
        used_terminology = {}
        if self.preprocess_terms:
            self.logger.info("=== 术语预处理已启用，开始收集段落术语 ===")
            self.logger.info(f"翻译方向: {self.source_lang} -> {self.target_lang}")
            self.logger.info(f"术语库大小: {len(terminology)} 个术语")

            # 显示术语库样本
            if terminology:
                sample_terms = list(terminology.items())[:5]
                self.logger.info(f"术语库样本（前5个）: {sample_terms}")

            # 对于外语→中文翻译模式，预先将术语库键值对调并缓存
            if self.source_lang != "zh":
                self.logger.info("外语→中文翻译模式，预先对调术语库键值并缓存...")
                self.reversed_terminology = self._create_reversed_terminology(terminology)
                self.logger.info(f"对调后的术语库大小: {len(self.reversed_terminology)} 个术语")

                # 显示对调后的术语库样本
                if self.reversed_terminology:
                    reversed_sample = list(self.reversed_terminology.items())[:5]
                    self.logger.info(f"对调后术语库样本（前5个）: {reversed_sample}")
            else:
                self.reversed_terminology = None

            # 先遍历所有段落，收集使用的术语
            para_count = len(doc.paragraphs)
            self.logger.info(f"开始分析 {para_count} 个段落中的术语...")

            for para_idx, paragraph in enumerate(doc.paragraphs, 1):
                self.web_logger.info(f"Processing paragraph {para_idx}/{para_count}")
                if paragraph.text.strip():
                    try:
                        # 根据翻译方向选择不同的术语提取方法
                        if self.source_lang == "zh":
                            # 中文 → 外语
                            self.logger.info(f"从中文段落提取术语: {paragraph.text[:50]}...")
                            para_terms = self.term_extractor.extract_terms(paragraph.text, terminology)
                            self.logger.info(f"从中文段落提取术语: {len(para_terms)} 个")
                        else:
                            # 外语 → 中文，使用缓存的反向术语库
                            self.logger.info(f"从外语段落提取术语: {paragraph.text[:50]}...")
                            if self.reversed_terminology:
                                # 使用缓存的反向术语库进行高效匹配
                                para_terms = self.term_extractor.extract_foreign_terms_from_reversed_dict(paragraph.text, self.reversed_terminology)
                                self.logger.info(f"从外语段落提取术语（使用缓存）: {len(para_terms)} 个")
                            else:
                                # 回退到原始方法
                                para_terms = self.term_extractor.extract_foreign_terms_by_chinese_values(paragraph.text, terminology)
                                self.logger.info(f"从外语段落提取术语（原始方法）: {len(para_terms)} 个")

                        # 显示提取到的术语
                        if para_terms:
                            self.logger.info(f"段落提取到的术语: {list(para_terms.items())[:3]}")

                        # 更新使用的术语词典
                        used_terminology.update(para_terms)
                    except Exception as e:
                        self.logger.error(f"段落术语提取失败: {str(e)}")
                        self.web_logger.error(f"Paragraph term extraction failed: {str(e)}")

            self.logger.info(f"从段落中提取了 {len(used_terminology)} 个术语")
            self.web_logger.info(f"Extracted {len(used_terminology)} terms from paragraphs")

            # 如果有使用的术语，导出到Excel文件
            if used_terminology:
                self._export_used_terminology(used_terminology)

        # 处理段落翻译
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraph_count += 1
                logger.info(f"正在翻译段落 {paragraph_count}: {paragraph.text[:50]}...")

                # 保存原文文本（在所有模式下都需要）
                original_text = paragraph.text

                # 如果是仅翻译模式，清空原文段落
                if self.output_format == "translation_only":
                    paragraph.clear()
                    # 保存原始格式信息
                    original_format = self._save_format_info([paragraph])
                else:
                    # 双语对照模式，保存原始格式信息
                    original_format = self._save_format_info([paragraph])

                # 检查段落中是否包含数学公式
                text, formulas = self._extract_latex_formulas(paragraph.text if self.output_format == "bilingual" else original_text)

                # 翻译段落内容（不包含公式部分）
                try:
                    if self.preprocess_terms:
                        # 使用术语预处理方式翻译
                        if self.is_cn_to_foreign:
                            # 中文 → 外语
                            # 检查是否有可用的术语
                            if not used_terminology:
                                # 如果没有术语，使用常规翻译
                                logger.info("未找到匹配术语，使用常规翻译（带术语库）")
                                translation = self.translator.translate_text(text, terminology, self.source_lang, self.target_lang)
                            else:
                                # 直接使用翻译器的内置术语处理功能
                                logger.info(f"使用翻译器内置术语处理功能，找到 {len(used_terminology)} 个匹配术语")
                                # 记录术语样本（仅记录前5个术语，避免日志过大）
                                terms_sample = list(used_terminology.items())[:5]
                                logger.info(f"术语样本（前5个）: {terms_sample}")

                                translation = self.translator.translate_text(text, used_terminology, self.source_lang, self.target_lang)
                                logger.info(f"最终翻译结果前100个字符: {translation[:100]}")
                        else:
                            # 外语 → 中文
                            # 检查是否有可用的术语
                            if not used_terminology:
                                # 如果没有术语，使用常规翻译
                                logger.info("未找到匹配术语，使用常规翻译")
                                translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)
                            else:
                                # 使用term_extractor的占位符系统进行术语预处理
                                # used_terminology已经是 {外语术语: 中文术语} 格式，直接使用
                                reverse_terminology = used_terminology

                                logger.info(f"使用术语预处理方式翻译，找到 {len(reverse_terminology)} 个匹配术语")
                                # 记录术语样本（仅记录前5个术语，避免日志过大）
                                terms_sample = list(reverse_terminology.items())[:5]
                                logger.info(f"术语样本（前5个）: {terms_sample}")

                                processed_text = self.term_extractor.replace_foreign_terms_with_placeholders(text, reverse_terminology)
                                logger.info(f"替换后的文本前100个字符: {processed_text[:100]}")

                                # 翻译处理后的文本（不使用术语库，因为已经预处理了）
                                logger.info("开始翻译含占位符的文本...")
                                translated_with_placeholders = self.translator.translate_text(processed_text, None, self.source_lang, self.target_lang)
                                logger.info(f"翻译后的文本前100个字符: {translated_with_placeholders[:100]}")

                                # 将占位符替换回中文术语
                                logger.info("开始将占位符替换回中文术语...")
                                translation = self.term_extractor.restore_placeholders_with_chinese_terms(translated_with_placeholders)
                                logger.info(f"最终翻译结果前100个字符: {translation[:100]}")
                    else:
                        # 使用常规方式翻译
                        if self.is_cn_to_foreign:
                            # 中文 → 外语，使用术语库
                            translation = self.translator.translate_text(text, terminology, self.source_lang, self.target_lang)
                        else:
                            # 外语 → 中文，不使用术语库（因为术语库是中文→外语格式）
                            translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)
                except Exception as e:
                    logger.error(f"翻译段落内容失败: {str(e)}")
                    # 返回错误信息而不是抛出异常，这样可以继续处理其他段落
                    translation = f"翻译失败: {str(e)}"

                # 将公式重新插入到翻译后的文本中
                if formulas:
                    translation = self._restore_latex_formulas(translation, formulas)

                logger.info(f"段落翻译完成: {translation[:50]}")

                # 验证翻译结果质量
                location = f"段落 {paragraph_count}"
                validation_issues = self._validate_translation_result(original_text, translation, location)
                if validation_issues:
                    for issue in validation_issues:
                        logger.warning(f"翻译质量问题: {issue}")
                        self.web_logger.warning(f"Translation quality issue: {issue}")

                # 收集翻译结果
                translation_results.append({
                    'original': paragraph.text if self.output_format == "bilingual" else original_text,
                    'translated': translation,
                    'location': location,
                    'validation_issues': validation_issues  # 添加验证问题信息
                })

                # 在原文后添加翻译
                self._add_translation_with_format([paragraph], translation, original_format)
                time.sleep(0.1)  # 添加小延迟，避免API请求过快

    def _extract_latex_formulas(self, text: str) -> Tuple[str, List[Tuple[str, str]]]:
        """
        提取文本中的LaTeX公式，并用占位符替换

        Args:
            text: 原始文本

        Returns:
            Tuple[str, List[Tuple[str, str]]]: 替换后的文本和公式列表(占位符, 公式)
        """
        formulas = []
        processed_text = text

        # 遍历所有公式模式
        for i, pattern in enumerate(self.latex_patterns):
            # 查找所有匹配的公式
            matches = re.findall(pattern, processed_text, re.DOTALL)

            # 替换每个公式为占位符
            for j, match in enumerate(matches):
                # 获取完整的公式文本
                if pattern.startswith(r'\$\$'):
                    full_formula = f"$${match}$$"
                elif pattern.startswith(r'\$'):
                    full_formula = f"${match}$"
                elif 'equation' in pattern:
                    full_formula = f"\\begin{{equation}}{match}\\end{{equation}}"
                elif 'align' in pattern:
                    full_formula = f"\\begin{{align}}{match}\\end{{align}}"
                elif 'eqnarray' in pattern:
                    full_formula = f"\\begin{{eqnarray}}{match}\\end{{eqnarray}}"
                else:
                    full_formula = match

                # 创建唯一的占位符
                placeholder = f"[LATEX_FORMULA_{i}_{j}]"

                # 保存公式和占位符
                formulas.append((placeholder, full_formula))

                # 替换文本中的公式为占位符
                processed_text = processed_text.replace(full_formula, placeholder, 1)

        return processed_text, formulas

    def _restore_latex_formulas(self, text: str, formulas: List[Tuple[str, str]]) -> str:
        """
        将占位符替换回LaTeX公式

        Args:
            text: 包含占位符的文本
            formulas: 公式列表(占位符, 公式)

        Returns:
            str: 恢复公式后的文本
        """
        result = text

        # 替换所有占位符为原始公式
        for placeholder, formula in formulas:
            result = result.replace(placeholder, formula)

        return result

    def export_to_pdf(self, docx_path: str) -> str:
        """
        将Word文档转换为PDF

        Args:
            docx_path: Word文档路径

        Returns:
            str: PDF文件路径
        """
        try:
            logger.info(f"正在将Word文档转换为PDF: {docx_path}")

            # 构建PDF文件路径（与Word文档同名，但扩展名为.pdf）
            pdf_path = os.path.splitext(docx_path)[0] + ".pdf"

            # 使用docx2pdf转换
            if not DOCX2PDF_AVAILABLE:
                raise ImportError("docx2pdf模块未安装，无法转换为PDF")
            docx2pdf_convert(docx_path, pdf_path)

            logger.info(f"PDF文件已保存到: {pdf_path}")
            return pdf_path
        except Exception as e:
            logger.error(f"转换PDF失败: {str(e)}")
            raise Exception(f"转换PDF失败: {str(e)}")

    def _save_format_info(self, paragraphs: List[Any]) -> List[Dict]:
        """保存段落格式信息"""
        format_info = []
        for paragraph in paragraphs:
            para_format = {
                'style_name': paragraph.style.name,
                'runs': []
            }
            for run in paragraph.runs:
                run_format = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_size': run.font.size,
                    'font_name': run.font.name
                }
                para_format['runs'].append(run_format)
            format_info.append(para_format)
        return format_info

    def _add_translation_with_format(self, paragraphs: List[Any], translation: str, original_format: List[Dict]) -> None:
        """添加带格式的翻译文本"""
        # 修复：只在第一个有内容的段落中添加翻译，避免重复
        translation_added = False

        for paragraph in paragraphs:
            # 只在第一个有内容的段落中添加翻译
            if not translation_added and paragraph.text.strip():
                # 根据输出格式决定如何添加翻译
                if self.output_format == "bilingual":
                    # 双语对照模式：在原文后添加翻译
                    paragraph.add_run('\n')  # 添加换行
                    run = paragraph.add_run(translation)
                else:
                    # 仅翻译模式：直接添加翻译（不添加换行）
                    run = paragraph.add_run(translation)

                # 设置翻译文本字体
                run.font.name = 'Calibri'

                # 保持与原文相同的格式
                if original_format and original_format[0]['runs']:
                    first_run = original_format[0]['runs'][0]
                    run.bold = first_run['bold']
                    run.italic = first_run['italic']
                    run.underline = first_run['underline']
                    if first_run['font_size']:
                        run.font.size = first_run['font_size']

                translation_added = True
                break  # 添加完翻译后立即退出循环

    def _extract_complete_cell_text(self, cell) -> str:
        """
        更完整地提取单元格文本，确保不遗漏任何内容（增强版）

        Args:
            cell: 表格单元格对象

        Returns:
            str: 完整的单元格文本
        """
        if not cell or not cell.paragraphs:
            return ""

        # 收集所有段落的文本
        all_text_parts = []

        logger.debug(f"单元格包含 {len(cell.paragraphs)} 个段落")

        for para_idx, paragraph in enumerate(cell.paragraphs):
            logger.debug(f"处理段落 {para_idx + 1}: runs数量={len(paragraph.runs)}")

            if not paragraph.runs:
                # 如果段落没有runs，直接获取段落文本
                para_text = paragraph.text.strip()
                if para_text:
                    all_text_parts.append(para_text)
                    logger.debug(f"段落 {para_idx + 1} (无runs): '{para_text}'")
            else:
                # 如果段落有runs，逐个收集run的文本
                para_parts = []
                for run_idx, run in enumerate(paragraph.runs):
                    run_text = run.text
                    if run_text:
                        para_parts.append(run_text)
                        logger.debug(f"段落 {para_idx + 1} run {run_idx + 1}: '{run_text}'")

                if para_parts:
                    para_text = ''.join(para_parts).strip()
                    if para_text:
                        all_text_parts.append(para_text)
                        logger.debug(f"段落 {para_idx + 1} 合并后: '{para_text}'")

        # 合并所有文本部分，使用空格分隔而不是换行，避免句子被截断
        if len(all_text_parts) > 1:
            # 多个段落，检查是否应该用空格还是换行连接
            complete_text = ' '.join(all_text_parts).strip()
            logger.debug(f"多段落合并 (空格连接): '{complete_text}'")

            # 如果文本看起来像是连续的句子，使用空格连接
            # 如果文本看起来像是独立的行，使用换行连接
            if any('\n' in part for part in all_text_parts):
                complete_text = '\n'.join(all_text_parts).strip()
                logger.debug(f"多段落合并 (换行连接): '{complete_text}'")
        else:
            complete_text = '\n'.join(all_text_parts).strip()

        # 如果提取的文本为空，尝试使用原始的cell.text作为备用
        if not complete_text:
            complete_text = cell.text.strip()
            logger.debug(f"使用备用方法提取: '{complete_text}'")

        logger.debug(f"最终提取的完整文本: '{complete_text}'")
        return complete_text

    def _validate_cell_text_completeness(self, cell, extracted_text: str) -> bool:
        """
        验证提取的单元格文本是否完整

        Args:
            cell: 表格单元格对象
            extracted_text: 提取的文本

        Returns:
            bool: True表示文本完整，False表示可能有遗漏
        """
        try:
            # 统计原始单元格中的字符数
            total_chars = 0
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        total_chars += len(run.text.strip())

            # 比较提取文本的字符数
            extracted_chars = len(extracted_text.strip())

            # 如果提取的字符数明显少于原始字符数，可能有遗漏
            if total_chars > 0 and extracted_chars < total_chars * 0.8:
                logger.warning(f"单元格文本可能不完整: 原始{total_chars}字符，提取{extracted_chars}字符")
                return False

            return True
        except Exception as e:
            logger.error(f"验证单元格文本完整性失败: {str(e)}")
            return True  # 出错时假设完整

    def _contains_chinese_content(self, text: str) -> bool:
        """
        检查文本是否包含需要翻译的中文内容

        Args:
            text: 要检查的文本

        Returns:
            bool: True表示包含中文内容，False表示不包含
        """
        if not text or not text.strip():
            return False

        # 检查是否包含中文字符
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        if chinese_chars:
            logger.debug(f"发现中文内容: {text} - 中文字符: {chinese_chars[:10]}")
            return True

        return False

    def _should_skip_translation(self, text: str) -> bool:
        """
        判断文本是否应该跳过翻译（增强版）

        Args:
            text: 要检查的文本

        Returns:
            bool: True表示应该跳过翻译，False表示需要翻译
        """
        if not text or not text.strip():
            return True

        text = text.strip()

        # 1. 纯数字（包括小数）
        if re.match(r'^[\d\.\,\-\+]+$', text):
            logger.debug(f"跳过翻译（纯数字）: {text}")
            return True

        # 2. 数字+单位组合（如 "≥50μs", "0.2-0.4Ω.cm", "20＜x＞50μs"）
        # 注意：不能包含中文字符，避免误判中文术语
        if re.match(r'^[≥≤<>＜＞\d\.\,\-\+\s]*[μΩa-zA-Z]*[μΩa-zA-Z\.\s]*$', text) and not re.search(r'[\u4e00-\u9fff]', text):
            logger.debug(f"跳过翻译（数字+单位）: {text}")
            return True

        # 3. 纯符号或特殊字符
        if re.match(r'^[≥≤<>＜＞\-\+\=\(\)\[\]\{\}\s]+$', text):
            logger.debug(f"跳过翻译（纯符号）: {text}")
            return True

        # 4. 很短的文本（1个字符），可能是代码或标识符
        if len(text) <= 1:
            logger.debug(f"跳过翻译（过短）: {text}")
            return True

        # 5. 检查是否包含需要翻译的中文内容（重点优化）
        if self.source_lang == "zh":
            # 中文→外语翻译：检查是否包含中文字符
            chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
            if chinese_chars:
                # 包含中文字符，需要翻译
                logger.debug(f"需要翻译（包含中文）: {text} - 中文字符: {chinese_chars[:5]}")
                return False

            # 不包含中文，检查是否已经是英文
            if re.match(r'^[a-zA-Z\s\d\.\,\-\+\(\)\[\]]+$', text):
                logger.debug(f"跳过翻译（已是英文）: {text}")
                return True

        # 6. 检查是否包含需要翻译的英文内容
        elif self.source_lang == "en":
            # 英文→中文翻译：检查是否主要是英文
            english_chars = len(re.findall(r'[a-zA-Z]', text))
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
            total_meaningful_chars = english_chars + chinese_chars

            if total_meaningful_chars > 0:
                english_ratio = english_chars / total_meaningful_chars
                if english_ratio > 0.5:  # 50%以上是英文，需要翻译
                    logger.debug(f"需要翻译（主要是英文）: {text}")
                    return False
                elif chinese_chars > 0:
                    # 主要是中文，可能不需要翻译
                    logger.debug(f"跳过翻译（主要是中文）: {text}")
                    return True

        # 7. 特殊情况：空白或只包含标点符号
        if re.match(r'^[\s\.\,\;\:\!\?\(\)\[\]\{\}\"\']+$', text):
            logger.debug(f"跳过翻译（只有标点符号）: {text}")
            return True

        # 8. 默认情况：如果包含任何文字内容，都尝试翻译
        if re.search(r'[\u4e00-\u9fff]|[a-zA-Z]', text):
            logger.debug(f"需要翻译（包含文字内容）: {text}")
            return False

        logger.debug(f"跳过翻译（其他情况）: {text}")
        return True

    def _update_table_cell_text(self, paragraphs: List[Any], original_text: str, translated_text: str) -> None:
        """更新表格单元格文本（与C#版本保持一致）"""
        if not paragraphs:
            return

        # 找到第一个有内容的段落
        target_paragraph = None
        for paragraph in paragraphs:
            if paragraph.text.strip():
                target_paragraph = paragraph
                break

        if target_paragraph is None:
            return

        # 清除所有段落的内容，避免重复
        for paragraph in paragraphs:
            # 清除所有run元素
            for run in paragraph.runs[:]:  # 使用切片复制避免修改列表时的问题
                paragraph._element.remove(run._element)

        # 只在第一个段落中添加内容
        if self.output_format == "bilingual":
            # 双语对照模式：显示原文和译文
            if original_text and original_text.strip():
                original_run = target_paragraph.add_run(original_text)
                # 添加换行
                target_paragraph.add_run('\n')

            # 添加翻译文本
            if translated_text and translated_text.strip():
                translated_run = target_paragraph.add_run(translated_text)
                # 设置翻译文本的字体颜色为蓝色以便区分
                translated_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        else:
            # 仅翻译结果模式：只显示翻译文本
            if translated_text and translated_text.strip():
                new_run = target_paragraph.add_run(translated_text)
            else:
                # 如果翻译文本为空，保留原文
                logger.warning(f"翻译文本为空，保留原文: {original_text[:50]}...")
                if original_text and original_text.strip():
                    new_run = target_paragraph.add_run(original_text)

    def _is_already_translated(self, text: str) -> bool:
        """检查文本是否已经包含翻译内容（增强版）"""
        logger.debug(f"开始检查是否已翻译: '{text}'")

        if not text or not text.strip():
            logger.debug("文本为空，返回False")
            return False

        text = text.strip()

        # 1. 检查是否包含明显的双语对照标记
        bilingual_markers = [
            "原文：", "译文：", "英文：", "中文：",
            "Original:", "Translation:", "English:", "Chinese:"
        ]

        for marker in bilingual_markers:
            if marker in text:
                logger.debug(f"检测到双语标记，已翻译: {marker}")
                return True

        logger.debug("未检测到双语标记")

        # 2. 检查是否包含常见的双语对照模式特征
        lines = text.split('\n')
        lines = [line.strip() for line in lines if line.strip()]
        logger.debug(f"文本分行数: {len(lines)}")

        # 3. 如果有多行，检查是否符合双语对照格式
        if len(lines) >= 2:
            # 检查是否有明显的原文-译文模式
            for i in range(len(lines) - 1):
                current_line = lines[i]
                next_line = lines[i + 1]

                if current_line and next_line:
                    # 检查是否是中文-英文或英文-中文的组合
                    if (self._is_chinese_text(current_line) and self._is_english_text(next_line)) or \
                       (self._is_english_text(current_line) and self._is_chinese_text(next_line)):
                        logger.debug(f"检测到双语对照格式，已翻译")
                        return True

        # 4. 如果只有一行，检查是否包含中英文混合（可能已翻译）
        if len(lines) == 1:
            logger.debug("单行文本，进入混合语言检测")
            # 更严格的混合语言检测 - 提高阈值避免误判
            has_mixed = self._contains_mixed_languages(text)
            if has_mixed:
                # 进一步检查是否真的是翻译内容
                chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
                english_chars = len(re.findall(r'[a-zA-Z]', text))
                total_chars = len(text.strip())

                # 更严格的条件：需要明显的双语特征才认为已翻译
                # 1. 中英文字符都要足够多
                # 2. 总长度要合理（避免误判短文本）
                # 3. 中英文比例要相对平衡
                # 4. 增加更严格的阈值，避免误判简单混合
                if (chinese_chars >= 4 and english_chars >= 6 and total_chars >= 15 and
                    min(chinese_chars, english_chars) / max(chinese_chars, english_chars) > 0.4):
                    logger.debug(f"检测到明显的混合语言内容，可能已翻译: 中文{chinese_chars}字符，英文{english_chars}字符，总长度{total_chars}")
                    return True
                else:
                    logger.debug(f"混合语言检测未达到阈值: 中文{chinese_chars}字符，英文{english_chars}字符，总长度{total_chars}")
                    # 不要在这里返回False，继续检查括号翻译格式

        # 5. 检查是否包含翻译结果的特征模式
        # 例如：包含括号中的翻译 "位错(dislocation)"
        bracket_pattern1 = r'[\u4e00-\u9fff]+\s*\([a-zA-Z\s]+\)'
        bracket_pattern2 = r'[a-zA-Z\s]+\s*\([\u4e00-\u9fff\s]+\)'

        match1 = re.search(bracket_pattern1, text)
        match2 = re.search(bracket_pattern2, text)

        logger.debug(f"括号翻译检测 - 文本: '{text}'")
        logger.debug(f"括号翻译检测 - 模式1匹配: {bool(match1)} {f'-> {match1.group()}' if match1 else ''}")
        logger.debug(f"括号翻译检测 - 模式2匹配: {bool(match2)} {f'-> {match2.group()}' if match2 else ''}")

        if match1 or match2:
            logger.debug(f"检测到括号翻译格式，已翻译")
            return True

        return False

    def _diagnose_cell_translation_decision(self, cell_text: str, table_idx: int, row_idx: int, cell_idx: int) -> dict:
        """
        诊断表格单元格的翻译决策，提供详细信息

        Args:
            cell_text: 单元格文本
            table_idx: 表格索引
            row_idx: 行索引
            cell_idx: 列索引

        Returns:
            dict: 包含决策信息的字典
        """
        diagnosis = {
            'should_translate': False,
            'reason': '',
            'details': {}
        }

        # 基本信息
        diagnosis['details']['text_length'] = len(cell_text.strip())
        diagnosis['details']['chinese_chars'] = len(re.findall(r'[\u4e00-\u9fff]', cell_text))
        diagnosis['details']['english_chars'] = len(re.findall(r'[a-zA-Z]', cell_text))
        diagnosis['details']['has_chinese'] = self._contains_chinese_content(cell_text)
        diagnosis['details']['should_skip'] = self._should_skip_translation(cell_text)
        diagnosis['details']['is_already_translated'] = self._is_already_translated(cell_text)

        # 决策逻辑
        if not cell_text.strip():
            diagnosis['reason'] = '空文本'
        elif diagnosis['details']['should_skip']:
            diagnosis['reason'] = '满足跳过条件（数字、符号等）'
        elif diagnosis['details']['is_already_translated']:
            diagnosis['reason'] = '检测为已翻译内容'
        elif self.source_lang == "zh" and not diagnosis['details']['has_chinese']:
            diagnosis['reason'] = '中文翻译模式但无中文内容'
        else:
            diagnosis['should_translate'] = True
            diagnosis['reason'] = '需要翻译'

        # 记录详细日志
        logger.info(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译决策:")
        logger.info(f"  文本: '{cell_text[:100]}{'...' if len(cell_text) > 100 else ''}'")
        logger.info(f"  长度: {diagnosis['details']['text_length']}, 中文字符: {diagnosis['details']['chinese_chars']}, 英文字符: {diagnosis['details']['english_chars']}")
        logger.info(f"  决策: {'需要翻译' if diagnosis['should_translate'] else '跳过翻译'}")
        logger.info(f"  原因: {diagnosis['reason']}")

        return diagnosis

    def _validate_translation_result(self, original_text: str, translated_text: str, location: str) -> list:
        """
        验证翻译结果的质量，检测潜在问题

        Args:
            original_text: 原文
            translated_text: 译文
            location: 位置信息

        Returns:
            list: 发现的问题列表
        """
        issues = []

        if not translated_text or not translated_text.strip():
            issues.append(f"{location}: 翻译结果为空")
            return issues

        # 1. 检查是否包含思考过程残留
        thinking_indicators = [
            "让我", "我来", "我需要", "首先", "根据", "考虑到", "分析", "思考",
            "这个术语", "在材料科学", "在英文中", "最合适", "最准确", "最恰当",
            "应该翻译", "可以翻译", "比较合适", "在这个语境"
        ]

        for indicator in thinking_indicators:
            if indicator in translated_text:
                issues.append(f"{location}: 翻译结果包含思考过程 - '{indicator}'")
                break

        # 2. 检查是否包含提示词残留
        prompt_indicators = [
            "翻译结果", "译文", "以下是", "这是", "请将", "将以下",
            "translation:", "translate:", "result:"
        ]

        for indicator in prompt_indicators:
            if indicator.lower() in translated_text.lower():
                issues.append(f"{location}: 翻译结果包含提示词 - '{indicator}'")
                break

        # 3. 检查中文→英文翻译是否完整
        if self.source_lang == "zh" and self.target_lang == "en":
            # 检查原文中的中文是否都被翻译
            original_chinese = re.findall(r'[\u4e00-\u9fff]', original_text)
            translated_chinese = re.findall(r'[\u4e00-\u9fff]', translated_text)

            if len(original_chinese) > 0 and len(translated_chinese) > len(original_chinese) * 0.5:
                issues.append(f"{location}: 可能存在未翻译的中文内容")

            # 检查译文是否主要包含英文
            english_chars = len(re.findall(r'[a-zA-Z]', translated_text))
            total_chars = len(re.sub(r'\s', '', translated_text))
            if total_chars > 0 and english_chars / total_chars < 0.3:
                issues.append(f"{location}: 译文英文字符比例过低，可能翻译不完整")

        # 4. 检查英文→中文翻译是否完整
        elif self.source_lang == "en" and self.target_lang == "zh":
            # 检查译文是否包含足够的中文
            chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', translated_text))
            total_chars = len(re.sub(r'\s', '', translated_text))
            if total_chars > 0 and chinese_chars / total_chars < 0.3:
                issues.append(f"{location}: 译文中文字符比例过低，可能翻译不完整")

        # 5. 检查是否原文和译文完全相同（可能未翻译）
        if original_text.strip() == translated_text.strip() and len(original_text.strip()) > 3:
            issues.append(f"{location}: 原文和译文完全相同，可能未翻译")

        # 6. 检查长度异常
        original_len = len(original_text.strip())
        translated_len = len(translated_text.strip())

        if translated_len < original_len * 0.2:
            issues.append(f"{location}: 译文过短，可能翻译不完整")
        elif translated_len > original_len * 5:
            issues.append(f"{location}: 译文过长，可能包含多余内容")

        return issues

    def _contains_mixed_languages(self, text: str) -> bool:
        """检查文本是否包含混合语言"""
        if not text or not text.strip():
            return False

        has_chinese = False
        has_english = False

        for char in text:
            if '\u4e00' <= char <= '\u9fff':  # 中文字符范围
                has_chinese = True
            elif char.isalpha() and char.isascii():  # 英文字符
                has_english = True

            if has_chinese and has_english:
                return True

        return False

    def _is_chinese_text(self, text: str) -> bool:
        """检查文本是否主要是中文"""
        if not text or not text.strip():
            return False

        chinese_count = 0
        total_chars = 0

        for char in text:
            if not char.isspace() and char not in '.,;:!?()[]{}""''':
                total_chars += 1
                if '\u4e00' <= char <= '\u9fff':
                    chinese_count += 1

        return total_chars > 0 and chinese_count / total_chars > 0.5

    def _is_english_text(self, text: str) -> bool:
        """检查文本是否主要是英文"""
        if not text or not text.strip():
            return False

        english_count = 0
        total_chars = 0

        for char in text:
            if not char.isspace() and char not in '.,;:!?()[]{}""''':
                total_chars += 1
                if char.isalpha() and char.isascii():
                    english_count += 1

        return total_chars > 0 and english_count / total_chars > 0.5

    def export_to_excel(self, translation_results, original_file_path, target_language):
        """将翻译结果导出为Excel文件"""
        try:
            logger.info(f"正在导出翻译数据到Excel，共 {len(translation_results)} 条记录...")

            # 创建DataFrame
            df = pd.DataFrame({
                '原文': [item['original'] for item in translation_results],
                f'{target_language}翻译': [item['translated'] for item in translation_results]
            })

            # 生成保存路径
            file_name = os.path.basename(original_file_path)
            base_name = os.path.splitext(file_name)[0]
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

            # 使用与主文档相同的输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")

            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)

            # 构建Excel文件路径
            excel_path = os.path.join(output_dir, f"{base_name}_翻译对照表_{timestamp}.xlsx")

            # 保存到Excel
            df.to_excel(excel_path, index=False, engine='openpyxl')

            logger.info(f"翻译对照表已保存至: {excel_path}")
            return excel_path
        except Exception as e:
            logger.error(f"导出Excel文件时出错: {str(e)}")
            return None

    def _export_used_terminology(self, used_terminology: Dict[str, str]) -> str:
        """导出使用的术语到Excel文件，包含使用统计信息"""
        try:
            logger.info(f"正在导出使用的术语到Excel，共 {len(used_terminology)} 个术语...")

            # 根据翻译方向创建不同的列名
            if self.source_lang == "zh":
                # 中文 → 外语
                source_column = '中文术语'
                target_column = '目标语术语'
            else:
                # 外语 → 中文
                source_column = '外语术语'
                target_column = '中文术语'

            # 获取术语使用统计信息
            usage_stats = self.term_extractor.get_terminology_usage_stats()

            # 如果有统计信息，使用更详细的导出格式
            if usage_stats:
                # 创建包含使用次数的DataFrame
                data = []
                for term, term_info in usage_stats.items():
                    data.append({
                        source_column: term_info['source'],
                        target_column: term_info['target'],
                        '使用次数': term_info['count']
                    })

                # 如果没有统计数据，使用基本的术语列表
                if not data:
                    for source, target in used_terminology.items():
                        data.append({
                            source_column: source,
                            target_column: target,
                            '使用次数': 'N/A'
                        })

                # 创建DataFrame
                df = pd.DataFrame(data)

                # 按使用次数降序排序
                if '使用次数' in df.columns and any(isinstance(x, int) for x in df['使用次数']):
                    df = df.sort_values(by='使用次数', ascending=False)
            else:
                # 创建基本的DataFrame（无使用次数）
                df = pd.DataFrame({
                    source_column: list(used_terminology.keys()),
                    target_column: list(used_terminology.values())
                })

            # 生成保存路径
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

            # 使用与主文档相同的输出目录
            output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "输出")

            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)

            # 构建Excel文件路径
            excel_path = os.path.join(output_dir, f"使用的术语_{timestamp}.xlsx")

            # 保存到Excel
            df.to_excel(excel_path, index=False, engine='openpyxl')

            logger.info(f"使用的术语表已保存至: {excel_path}")
            return excel_path
        except Exception as e:
            logger.error(f"导出术语Excel文件时出错: {str(e)}")
            return None

    def _create_reversed_terminology(self, terminology: Dict[str, str]) -> Dict[str, str]:
        """
        创建反向术语库（外语术语: 中文术语）

        Args:
            terminology: 原始术语库 {中文术语: 外语术语}

        Returns:
            Dict[str, str]: 反向术语库 {外语术语: 中文术语}
        """
        reversed_terminology = {}

        for cn_term, foreign_term in terminology.items():
            # 确保外语术语不为空
            if foreign_term and foreign_term.strip():
                # 如果有多个中文术语对应同一个外语术语，选择最长的中文术语
                if foreign_term in reversed_terminology:
                    existing_cn_term = reversed_terminology[foreign_term]
                    if len(cn_term) > len(existing_cn_term):
                        reversed_terminology[foreign_term] = cn_term
                        self.logger.debug(f"术语冲突，选择更长的中文术语: {foreign_term} -> {cn_term} (替换 {existing_cn_term})")
                else:
                    reversed_terminology[foreign_term] = cn_term

        self.logger.info(f"成功创建反向术语库，原始术语库: {len(terminology)} 个，反向术语库: {len(reversed_terminology)} 个")
        return reversed_terminology

    def _update_progress(self, progress: float, message: str = ""):
        """更新进度"""
        # 记录进度到日志
        if progress >= 0:
            logger.info(f"进度: {progress*100:.1f}% - {message}")
            self.web_logger.info(f"Progress: {progress*100:.1f}% - {message}")
        else:
            logger.error(f"进度: 错误 - {message}")
            self.web_logger.error(f"Progress: Error - {message}")

        if self.progress_callback:
            import asyncio
            import threading

            def safe_callback():
                """安全的回调函数执行"""
                try:
                    # 尝试异步调用回调函数
                    if asyncio.iscoroutinefunction(self.progress_callback):
                        # 检查是否已有事件循环在运行
                        try:
                            loop = asyncio.get_event_loop()
                            if loop.is_running():
                                # 如果当前事件循环正在运行，创建一个任务而不是新的事件循环
                                asyncio.create_task(self.progress_callback(progress, message))
                            else:
                                # 如果当前事件循环不在运行，使用它来运行协程
                                loop.run_until_complete(self.progress_callback(progress, message))
                        except RuntimeError:
                            # 如果没有事件循环，创建一个新的
                            loop = asyncio.new_event_loop()
                            asyncio.set_event_loop(loop)
                            try:
                                loop.run_until_complete(self.progress_callback(progress, message))
                            finally:
                                loop.close()
                    else:
                        # 同步回调函数
                        self.progress_callback(progress, message)
                except Exception as e:
                    logger.error(f"更新进度失败: {str(e)}")

            # 在单独线程中执行回调，避免阻塞主线程
            try:
                callback_thread = threading.Thread(target=safe_callback)
                callback_thread.daemon = True
                callback_thread.start()

                # 等待最多1秒，避免进度更新阻塞太久
                callback_thread.join(timeout=1.0)

                if callback_thread.is_alive():
                    logger.warning("进度回调执行超时，跳过此次更新")

            except Exception as e:
                logger.error(f"执行进度回调时出错: {str(e)}")

    def _translate_with_retry(self, text: str, terminology: Dict = None) -> str:
        """带重试机制的翻译"""
        if not text.strip():
            return text

        # 增加自适应重试延迟
        retry_delay = self.retry_delay
        max_retry_delay = 30  # 最大重试延迟30秒

        for attempt in range(self.retry_count):
            try:
                # 使用翻译服务进行翻译
                translation = self.translator.translate_text(text, terminology, self.source_lang, self.target_lang)

                # 检查翻译结果是否为空或异常短
                if not translation or (len(translation) < len(text) * 0.1 and len(text) > 50):
                    logger.warning(f"翻译结果异常短 (尝试 {attempt+1}/{self.retry_count}): 原文长度 {len(text)}，译文长度 {len(translation or '')}")
                    if attempt < self.retry_count - 1:
                        logger.info(f"等待 {retry_delay} 秒后重试...")
                        time.sleep(retry_delay)
                        # 增加重试延迟（指数退避）
                        retry_delay = min(retry_delay * 1.5, max_retry_delay)
                        continue

                # 检查翻译结果是否包含错误信息
                error_indicators = ["翻译失败", "translation failed", "error", "错误", "请求超时", "timeout"]
                if any(indicator in translation.lower() for indicator in error_indicators):
                    logger.warning(f"翻译结果包含错误信息 (尝试 {attempt+1}/{self.retry_count}): {translation[:100]}")
                    if attempt < self.retry_count - 1:
                        logger.info(f"等待 {retry_delay} 秒后重试...")
                        time.sleep(retry_delay)
                        # 增加重试延迟（指数退避）
                        retry_delay = min(retry_delay * 1.5, max_retry_delay)
                        continue

                return translation
            except Exception as e:
                logger.warning(f"翻译失败 (尝试 {attempt+1}/{self.retry_count}): {str(e)}")
                if attempt < self.retry_count - 1:
                    logger.info(f"等待 {retry_delay} 秒后重试...")
                    time.sleep(retry_delay)
                    # 增加重试延迟（指数退避）
                    retry_delay = min(retry_delay * 1.5, max_retry_delay)
                else:
                    # 最后一次尝试失败，抛出异常
                    raise Exception(f"翻译失败，已重试 {self.retry_count} 次: {str(e)}")

    def translate_paragraph(self, text: str, target_language: str, terminology: Dict) -> str:
        """翻译段落文本"""
        if not text.strip():
            return text

        try:
            # 根据术语库开关决定是否使用术语库
            if self.use_terminology:
                # 获取目标语言的术语库
                terms_dict = terminology.get(target_language, {})

                # 如果目标语言不在术语库中，尝试使用语言代码映射
                if not terms_dict and target_language in ['en', 'ja', 'ko', 'fr', 'de', 'es', 'ru']:
                    # 映射语言代码到术语库中的语言名称
                    language_map = {
                        'en': '英语',
                        'ja': '日语',
                        'ko': '韩语',
                        'fr': '法语',
                        'de': '德语',
                        'es': '西班牙语',
                        'ru': '俄语'
                    }
                    mapped_language = language_map.get(target_language)
                    if mapped_language and mapped_language in terminology:
                        terms_dict = terminology.get(mapped_language, {})
                        logger.info(f"使用映射后的语言名称 '{mapped_language}' 获取术语库")

                # 记录术语库大小
                logger.info(f"使用{target_language}术语库，包含 {len(terms_dict)} 个术语")

                # 记录术语预处理状态
                logger.info(f"术语预处理功能状态: {'启用' if self.preprocess_terms else '禁用'}")
                logger.info(f"翻译方向: {'中文→外语' if self.is_cn_to_foreign else '外语→中文'}")

                # 如果术语库为空，直接使用常规翻译
                if not terms_dict:
                    logger.warning(f"术语库为空，使用常规翻译")
                    return self._translate_with_retry(text, None)

                # 如果启用了术语预处理
                if self.preprocess_terms:
                    try:
                        # 记录术语库内容（仅记录前5个术语，避免日志过大）
                        terms_sample = list(terms_dict.items())[:5]
                        logger.info(f"术语库样本（前5个）: {terms_sample}")

                        # 根据翻译方向选择不同的术语处理方法
                        if self.is_cn_to_foreign:
                            # 中文 → 外语
                            logger.info(f"使用中文→外语术语预处理流程")

                            # 提取文本中使用的术语
                            logger.info(f"开始从中文文本中提取术语，文本前50个字符: {text[:50]}")
                            used_terms = self.term_extractor.extract_terms(text, terms_dict)
                            logger.info(f"从中文文本中提取了 {len(used_terms)} 个术语")

                            # 记录提取到的术语（仅记录前5个，避免日志过大）
                            if used_terms:
                                terms_sample = list(used_terms.items())[:5]
                                logger.info(f"提取到的术语样本（前5个）: {terms_sample}")
                            else:
                                logger.warning("未从文本中提取到任何术语")

                            # 如果没有找到术语，使用常规翻译
                            if not used_terms:
                                logger.info("未找到匹配术语，使用常规翻译（带术语库）")
                                return self._translate_with_retry(text, terms_dict)

                            # 直接使用翻译器的内置术语处理功能
                            logger.info(f"使用翻译器内置术语处理功能，找到 {len(used_terms)} 个匹配术语")
                            # 记录术语样本（仅记录前5个术语，避免日志过大）
                            terms_sample = list(used_terms.items())[:5]
                            logger.info(f"术语样本（前5个）: {terms_sample}")

                            result = self._translate_with_retry(text, used_terms)
                            logger.info(f"最终翻译结果: {result[:100]}...")
                            return result
                        else:
                            # 外语 → 中文
                            logger.info(f"使用外语→中文术语预处理流程")

                            try:
                                # 提取文本中使用的术语，优先使用缓存的反向术语库
                                logger.info(f"开始从外语文本中提取术语，文本前50个字符: {text[:50]}")

                                # 检查是否有缓存的反向术语库
                                if hasattr(self, 'reversed_terminology') and self.reversed_terminology:
                                    # 使用缓存的反向术语库进行高效匹配
                                    logger.info("使用缓存的反向术语库进行术语提取")
                                    used_terms = self.term_extractor.extract_foreign_terms_from_reversed_dict(text, self.reversed_terminology)
                                    logger.info(f"从外语文本中提取了 {len(used_terms)} 个术语（使用缓存）")
                                else:
                                    # 回退到原始方法，但先创建反向术语库缓存
                                    logger.info("缓存不可用，创建反向术语库并使用")
                                    self.reversed_terminology = self._create_reversed_terminology(terms_dict)
                                    if self.reversed_terminology:
                                        used_terms = self.term_extractor.extract_foreign_terms_from_reversed_dict(text, self.reversed_terminology)
                                        logger.info(f"从外语文本中提取了 {len(used_terms)} 个术语（新建缓存）")
                                    else:
                                        # 最后的回退方案
                                        used_terms = self.term_extractor.extract_foreign_terms_by_chinese_values(text, terms_dict)
                                        logger.info(f"从外语文本中提取了 {len(used_terms)} 个术语（原始方法）")

                                # 记录提取到的术语（仅记录前5个，避免日志过大）
                                if used_terms:
                                    terms_sample = list(used_terms.items())[:5]
                                    logger.info(f"提取到的术语样本（前5个）: {terms_sample}")
                                else:
                                    logger.warning("未从外语文本中提取到任何术语")

                                # 如果没有找到术语，使用常规翻译
                                if not used_terms:
                                    logger.info("未找到匹配术语，使用常规翻译（不带术语库）")
                                    return self._translate_with_retry(text, None)

                                # 使用term_extractor的占位符系统进行术语预处理
                                # used_terms已经是 {外语术语: 中文术语} 格式，直接使用
                                reverse_terminology = used_terms

                                logger.info(f"使用术语预处理方式翻译，找到 {len(reverse_terminology)} 个匹配术语")
                                # 记录术语样本（仅记录前5个术语，避免日志过大）
                                terms_sample = list(reverse_terminology.items())[:5]
                                logger.info(f"术语样本（前5个）: {terms_sample}")

                                processed_text = self.term_extractor.replace_foreign_terms_with_placeholders(text, reverse_terminology)
                                logger.info(f"已将术语替换为占位符: {processed_text[:100]}...")

                                # 翻译处理后的文本（不使用术语库，因为已经预处理了）
                                logger.info("开始翻译含占位符的文本...")
                                translated_with_placeholders = self._translate_with_retry(processed_text, None)
                                logger.info(f"带占位符的翻译结果: {translated_with_placeholders[:100]}...")

                                # 将占位符替换回中文术语
                                logger.info("开始将占位符替换回中文术语...")
                                result = self.term_extractor.restore_placeholders_with_chinese_terms(translated_with_placeholders)
                                logger.info(f"最终翻译结果（替换回术语）: {result[:100]}...")
                                return result
                            except Exception as e:
                                logger.error(f"外语→中文术语处理失败: {str(e)}")
                                logger.error(f"错误详情: {str(e.__class__.__name__)}: {str(e)}")
                                # 如果术语处理失败，使用常规翻译
                                logger.info("术语处理失败，回退到常规翻译（不带术语库）")
                                return self._translate_with_retry(text, None)
                    except Exception as e:
                        logger.error(f"术语预处理失败: {str(e)}")
                        # 如果术语预处理失败，回退到常规翻译
                        logger.info("术语预处理失败，回退到常规翻译（带术语库）")
                        return self._translate_with_retry(text, terms_dict)
                else:
                    # 使用常规方式翻译
                    logger.info("使用常规方式翻译（带术语库）")
                    return self._translate_with_retry(text, terms_dict)
            else:
                # 不使用术语库的翻译
                logger.info("不使用术语库进行翻译")
                return self._translate_with_retry(text, None)

        except Exception as e:
            logger.error(f"翻译失败: {str(e)}")
            # 返回错误信息而不是抛出异常，这样可以继续处理其他段落
            return f"翻译失败: {str(e)}"

    def _translate_cell_with_retry(self, text: str, terminology: Dict, table_idx: int, row_idx: int, cell_idx: int, max_retries: int = 3) -> str:
        """
        带重试机制的表格单元格翻译方法

        Args:
            text: 要翻译的文本
            terminology: 术语词典
            table_idx: 表格索引
            row_idx: 行索引
            cell_idx: 列索引
            max_retries: 最大重试次数

        Returns:
            str: 翻译结果
        """
        last_error = None

        for attempt in range(max_retries):
            try:
                logger.info(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译尝试 {attempt + 1}/{max_retries}")

                if self.preprocess_terms:
                    # 使用术语预处理方式翻译
                    if self.source_lang == "zh":
                        # 中文 → 外语：使用新的直接术语替换策略
                        cell_terminology = self.term_extractor.extract_terms(text, terminology)
                        logger.info(f"从单元格提取到 {len(cell_terminology)} 个术语")

                        if cell_terminology:
                            terms_sample = list(cell_terminology.items())[:5]
                            logger.info(f"术语样本（前5个）: {terms_sample}")
                            logger.info("使用新的直接术语替换策略进行翻译")
                            # 直接使用提取的术语进行翻译，翻译器内部会处理术语替换
                            translation = self.translator.translate_text(text, cell_terminology, self.source_lang, self.target_lang)
                        else:
                            logger.info("未找到匹配术语，使用常规翻译（带完整术语库）")
                            translation = self.translator.translate_text(text, terminology, self.source_lang, self.target_lang)
                    else:
                        # 外语 → 中文：使用新的直接术语替换策略
                        cell_terminology = self.term_extractor.extract_foreign_terms(text, terminology)
                        logger.info(f"从单元格提取到 {len(cell_terminology)} 个外语术语")

                        if cell_terminology:
                            terms_sample = list(cell_terminology.items())[:5]
                            logger.info(f"术语样本（前5个）: {terms_sample}")
                            logger.info("使用新的直接术语替换策略进行翻译")
                            translation = self.translator.translate_text(text, cell_terminology, self.source_lang, self.target_lang)
                        else:
                            logger.info("未找到匹配术语，使用常规翻译")
                            translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)
                else:
                    # 不使用术语预处理，直接翻译
                    logger.info("不使用术语预处理，使用常规翻译")
                    if self.source_lang == "zh":
                        translation = self.translator.translate_text(text, terminology, self.source_lang, self.target_lang)
                    else:
                        translation = self.translator.translate_text(text, None, self.source_lang, self.target_lang)

                # 验证翻译结果
                if translation and translation.strip():
                    # 检查是否包含占位符残留
                    import re
                    placeholder_patterns = [r'\[术语\d+\]', r'\[Term\s*\d+\]']
                    has_placeholders = any(re.search(pattern, translation) for pattern in placeholder_patterns)

                    if has_placeholders:
                        logger.warning(f"翻译结果包含占位符残留，尝试重新翻译")
                        if attempt < max_retries - 1:
                            continue
                        else:
                            logger.error(f"多次重试后仍有占位符残留: {translation}")

                    logger.info(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译成功")
                    return translation
                else:
                    raise Exception("翻译结果为空")

            except Exception as e:
                last_error = e
                logger.warning(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译尝试 {attempt + 1} 失败: {str(e)}")

                if attempt < max_retries - 1:
                    import time
                    wait_time = (attempt + 1) * 2  # 递增等待时间
                    logger.info(f"等待 {wait_time} 秒后重试...")
                    time.sleep(wait_time)
                else:
                    logger.error(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 所有重试均失败")

        # 所有重试都失败，返回原文
        logger.error(f"表格 {table_idx} 行 {row_idx} 列 {cell_idx} 翻译最终失败: {str(last_error)}")
        return text  # 返回原文而不是错误信息