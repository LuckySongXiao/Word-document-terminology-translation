#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复表格4最后单元格的不完整翻译
"""

import os
import sys
import json
import logging
from docx import Document
from services.zhipuai_translator import ZhipuAITranslator
from services.ollama_translator import OllamaTranslator

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def load_api_config():
    """加载API配置"""
    try:
        with open('API_config/zhipu_api.json', 'r', encoding='utf-8') as f:
            zhipu_config = json.load(f)
        return zhipu_config
    except Exception as e:
        logger.error(f"加载API配置失败: {e}")
        return None

def get_translator():
    """获取翻译器"""
    # 尝试使用智谱AI
    config = load_api_config()
    if config and config.get('api_key'):
        try:
            translator = ZhipuAITranslator(config['api_key'])
            # 测试连接
            test_result = translator.translate("测试", {}, "zh", "en")
            if test_result and test_result != "测试":
                logger.info("使用智谱AI翻译器")
                return translator
        except Exception as e:
            logger.warning(f"智谱AI翻译器初始化失败: {e}")
    
    # 回退到本地Ollama
    try:
        # 从配置文件加载Ollama配置
        with open('config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        ollama_config = config.get("fallback_translator", {})
        model = ollama_config.get("model", "qwen2.5:7b")
        api_url = ollama_config.get("api_url", "http://localhost:11434")
        
        translator = OllamaTranslator(
            model=model,
            api_url=api_url,
            model_list_timeout=10,
            translate_timeout=60
        )
        logger.info("使用本地Ollama翻译器")
        return translator
    except Exception as e:
        logger.error(f"Ollama翻译器初始化失败: {e}")
        return None

def fix_incomplete_translation():
    """修复不完整的翻译"""
    
    # 文档路径
    output_dir = "输出"
    doc_name = "单晶电阻率管控技术标准_带翻译_20250610215844.docx"
    doc_path = os.path.join(output_dir, doc_name)
    
    if not os.path.exists(doc_path):
        logger.error(f"文档不存在: {doc_path}")
        return False
    
    # 获取翻译器
    translator = get_translator()
    if not translator:
        logger.error("无法初始化翻译器")
        return False
    
    try:
        # 打开文档
        doc = Document(doc_path)
        logger.info(f"打开文档: {doc_path}")
        
        # 找到表格4（最后一个表格）
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        table = doc.tables[3]  # 第4个表格（索引3）
        logger.info(f"处理表格4，行数: {len(table.rows)}")
        
        # 找到最后一行
        last_row = table.rows[-1]
        logger.info(f"最后一行单元格数: {len(last_row.cells)}")
        
        # 检查并修复最后一个单元格
        if len(last_row.cells) >= 1:
            last_cell = last_row.cells[-1]
            cell_text = last_cell.text.strip()
            
            logger.info(f"最后一个单元格当前内容: {cell_text[:100]}...")
            
            # 分析当前内容
            lines = [line.strip() for line in cell_text.split('\n') if line.strip()]
            chinese_lines = []
            english_lines = []
            
            import re
            for line in lines:
                has_chinese = bool(re.search(r'[\u4e00-\u9fff]', line))
                has_english = bool(re.search(r'[A-Za-z]', line))
                
                if has_chinese and not has_english:
                    chinese_lines.append(line)
                elif has_english and not has_chinese:
                    english_lines.append(line)
                elif has_chinese and has_english:
                    # 混合行，需要进一步分析
                    if line.startswith(('备注', '1、', '2、', '3、')):
                        chinese_lines.append(line)
                    else:
                        english_lines.append(line)
            
            logger.info(f"中文行数: {len(chinese_lines)}")
            logger.info(f"英文行数: {len(english_lines)}")
            
            # 找出未翻译的中文内容
            untranslated_lines = []
            for chinese_line in chinese_lines:
                # 检查这行中文是否有对应的英文翻译
                has_translation = False
                
                # 简单的匹配逻辑：如果中文行包含"1、"，检查是否有对应的"1."英文翻译
                if "1、" in chinese_line:
                    has_translation = any("1." in eng_line for eng_line in english_lines)
                elif "2、" in chinese_line:
                    has_translation = any("2." in eng_line for eng_line in english_lines)
                
                if not has_translation:
                    untranslated_lines.append(chinese_line)
                    logger.info(f"发现未翻译内容: {chinese_line}")
            
            # 翻译未翻译的内容
            if untranslated_lines:
                logger.info(f"需要翻译 {len(untranslated_lines)} 行内容")
                
                new_translations = []
                for line in untranslated_lines:
                    try:
                        logger.info(f"正在翻译: {line}")
                        translated = translator.translate(line, {}, "zh", "en")
                        if translated and translated != line:
                            new_translations.append(translated)
                            logger.info(f"翻译结果: {translated}")
                        else:
                            logger.warning(f"翻译失败或无变化: {line}")
                    except Exception as e:
                        logger.error(f"翻译出错: {e}")
                
                # 更新单元格内容
                if new_translations:
                    # 重新构建单元格内容
                    new_content_lines = []
                    
                    # 添加所有中文内容
                    for chinese_line in chinese_lines:
                        new_content_lines.append(chinese_line)
                    
                    # 添加所有英文翻译（包括原有的和新的）
                    for eng_line in english_lines:
                        new_content_lines.append(eng_line)
                    
                    for new_trans in new_translations:
                        new_content_lines.append(new_trans)
                    
                    # 清空现有内容并添加新内容
                    for para in last_cell.paragraphs:
                        para.clear()
                    
                    # 添加新内容
                    if last_cell.paragraphs:
                        para = last_cell.paragraphs[0]
                    else:
                        para = last_cell.add_paragraph()
                    
                    para.add_run('\n'.join(new_content_lines))
                    
                    logger.info("单元格内容已更新")
                else:
                    logger.warning("没有成功翻译任何内容")
            else:
                logger.info("没有发现未翻译的内容")
        
        # 保存修复后的文档
        fixed_doc_name = doc_name.replace(".docx", "_修复完整翻译.docx")
        fixed_doc_path = os.path.join(output_dir, fixed_doc_name)
        doc.save(fixed_doc_path)
        logger.info(f"修复后的文档已保存: {fixed_doc_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"修复文档时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    logger.info("开始修复表格翻译不完整问题...")
    
    success = fix_incomplete_translation()
    
    if success:
        logger.info("✅ 修复完成！")
        logger.info("请检查输出文件夹中的修复后文档")
    else:
        logger.error("❌ 修复失败！")

if __name__ == "__main__":
    main()
