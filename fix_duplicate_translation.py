#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复表格重复翻译问题
"""

import os
import sys
from docx import Document
import logging
import re

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def clean_duplicate_translations_in_table(doc_path, output_path):
    """清理表格中的重复翻译"""
    
    try:
        # 打开文档
        doc = Document(doc_path)
        logger.info(f"打开文档: {doc_path}")
        
        # 检查表格数量
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        # 处理表格4（索引3）
        table = doc.tables[3]
        logger.info(f"处理表格4，行数: {len(table.rows)}")
        
        # 分析表格结构，找出重复的行
        duplicate_rows = []
        
        for row_idx, row in enumerate(table.rows):
            logger.info(f"检查第{row_idx+1}行，单元格数: {len(row.cells)}")
            
            # 检查这一行是否所有单元格都包含相同的内容
            if len(row.cells) > 1:
                first_cell_text = row.cells[0].text.strip()
                
                # 如果第一个单元格包含"备注"，检查是否所有单元格都相同
                if "备注" in first_cell_text and len(first_cell_text) > 10:
                    all_same = True
                    for cell in row.cells[1:]:
                        if cell.text.strip() != first_cell_text:
                            all_same = False
                            break
                    
                    if all_same:
                        duplicate_rows.append(row_idx)
                        logger.info(f"发现重复行: 第{row_idx+1}行")
        
        logger.info(f"发现 {len(duplicate_rows)} 个重复行: {[r+1 for r in duplicate_rows]}")
        
        # 如果发现重复行，进行清理
        if duplicate_rows:
            # 保留第一个重复行，清空其他重复行
            keep_row = duplicate_rows[0]
            logger.info(f"保留第{keep_row+1}行，清空其他重复行")
            
            for row_idx in duplicate_rows[1:]:
                row = table.rows[row_idx]
                logger.info(f"清空第{row_idx+1}行")
                
                # 清空这一行的所有单元格
                for cell in row.cells:
                    # 清空单元格内容
                    for paragraph in cell.paragraphs:
                        paragraph.clear()
                    
                    # 如果没有段落，添加一个空段落
                    if not cell.paragraphs:
                        cell.add_paragraph()
        
        # 保存修复后的文档
        doc.save(output_path)
        logger.info(f"修复后的文档已保存: {output_path}")
        
        return True
        
    except Exception as e:
        logger.error(f"修复文档时出错: {e}")
        import traceback
        traceback.print_exc()
        return False

def verify_fix(doc_path):
    """验证修复效果"""
    
    try:
        doc = Document(doc_path)
        
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        table = doc.tables[3]
        logger.info(f"验证表格4，行数: {len(table.rows)}")
        
        # 检查是否还有重复行
        duplicate_count = 0
        non_empty_rows = 0
        
        for row_idx, row in enumerate(table.rows):
            # 检查这一行是否包含内容
            has_content = False
            for cell in row.cells:
                if cell.text.strip():
                    has_content = True
                    break
            
            if has_content:
                non_empty_rows += 1
                
                # 检查是否所有单元格都相同
                if len(row.cells) > 1:
                    first_cell_text = row.cells[0].text.strip()
                    if "备注" in first_cell_text and len(first_cell_text) > 10:
                        all_same = True
                        for cell in row.cells[1:]:
                            if cell.text.strip() != first_cell_text:
                                all_same = False
                                break
                        
                        if all_same:
                            duplicate_count += 1
                            logger.info(f"第{row_idx+1}行仍然重复")
        
        logger.info(f"验证结果: 非空行数={non_empty_rows}, 重复行数={duplicate_count}")
        
        if duplicate_count <= 1:
            logger.info("✅ 修复成功！重复翻译问题已解决")
            return True
        else:
            logger.warning(f"❌ 仍有 {duplicate_count} 个重复行")
            return False
            
    except Exception as e:
        logger.error(f"验证时出错: {e}")
        return False

def main():
    logger.info("开始修复表格重复翻译问题...")
    
    # 输入文档路径
    output_dir = "输出"
    
    # 查找最新的翻译文档
    import glob
    pattern = os.path.join(output_dir, "单晶电阻率管控技术标准_带翻译_*.docx")
    files = glob.glob(pattern)
    
    if not files:
        logger.error("未找到翻译文档")
        return False
    
    # 按修改时间排序，获取最新的文件
    latest_file = max(files, key=os.path.getmtime)
    logger.info(f"处理最新文档: {latest_file}")
    
    # 生成输出文件名
    base_name = os.path.splitext(os.path.basename(latest_file))[0]
    output_name = f"{base_name}_修复重复翻译.docx"
    output_path = os.path.join(output_dir, output_name)
    
    # 修复重复翻译
    success = clean_duplicate_translations_in_table(latest_file, output_path)
    
    if success:
        # 验证修复效果
        verify_success = verify_fix(output_path)
        
        if verify_success:
            logger.info("🎉 修复完成！重复翻译问题已解决")
            logger.info(f"修复后的文档: {output_path}")
            return True
        else:
            logger.error("修复验证失败")
            return False
    else:
        logger.error("修复失败")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
