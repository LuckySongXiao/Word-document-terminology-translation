#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查输出文档中表格翻译的完整性
"""

import os
import sys
from docx import Document
import logging

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def check_document_tables(doc_path):
    """检查文档中的表格翻译情况"""
    try:
        doc = Document(doc_path)
        logger.info(f"正在检查文档: {doc_path}")
        
        if not doc.tables:
            logger.info("文档中没有表格")
            return
            
        logger.info(f"文档包含 {len(doc.tables)} 个表格")
        
        for table_idx, table in enumerate(doc.tables, 1):
            logger.info(f"\n=== 表格 {table_idx} ===")
            logger.info(f"表格行数: {len(table.rows)}")
            
            # 检查每一行的单元格
            for row_idx, row in enumerate(table.rows):
                logger.info(f"  第 {row_idx + 1} 行，单元格数: {len(row.cells)}")
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        logger.info(f"    单元格 [{row_idx + 1}, {cell_idx + 1}]: {cell_text[:100]}...")
                    else:
                        logger.info(f"    单元格 [{row_idx + 1}, {cell_idx + 1}]: [空]")
            
            # 特别检查最后一行的最后两个单元格
            if table.rows:
                last_row = table.rows[-1]
                logger.info(f"\n  === 最后一行检查 ===")
                logger.info(f"  最后一行单元格数: {len(last_row.cells)}")
                
                if len(last_row.cells) >= 2:
                    # 检查最后两个单元格
                    second_last_cell = last_row.cells[-2]
                    last_cell = last_row.cells[-1]
                    
                    second_last_text = second_last_cell.text.strip()
                    last_text = last_cell.text.strip()
                    
                    logger.info(f"  倒数第二个单元格: {second_last_text}")
                    logger.info(f"  最后一个单元格: {last_text}")
                    
                    # 检查是否包含中文（可能未翻译）
                    def contains_chinese(text):
                        return any('\u4e00' <= char <= '\u9fff' for char in text)
                    
                    def contains_english(text):
                        return any(char.isalpha() and ord(char) < 128 for char in text)
                    
                    if second_last_text:
                        has_chinese = contains_chinese(second_last_text)
                        has_english = contains_english(second_last_text)
                        logger.info(f"  倒数第二个单元格 - 包含中文: {has_chinese}, 包含英文: {has_english}")
                        
                        if has_chinese and not has_english:
                            logger.warning(f"  ⚠️ 倒数第二个单元格可能未翻译（仅包含中文）")
                        elif has_chinese and has_english:
                            # 检查是否是双语对照格式
                            lines = second_last_text.split('\n')
                            chinese_lines = [line for line in lines if contains_chinese(line)]
                            english_lines = [line for line in lines if contains_english(line) and not contains_chinese(line)]
                            logger.info(f"  倒数第二个单元格 - 中文行数: {len(chinese_lines)}, 纯英文行数: {len(english_lines)}")

                            if len(chinese_lines) > len(english_lines):
                                logger.warning(f"  ⚠️ 倒数第二个单元格可能翻译不完整（中文行多于英文行）")
                                logger.info(f"  中文内容: {chinese_lines}")
                                logger.info(f"  英文内容: {english_lines}")
                    
                    if last_text:
                        has_chinese = contains_chinese(last_text)
                        has_english = contains_english(last_text)
                        logger.info(f"  最后一个单元格 - 包含中文: {has_chinese}, 包含英文: {has_english}")

                        if has_chinese and not has_english:
                            logger.warning(f"  ⚠️ 最后一个单元格可能未翻译（仅包含中文）")
                        elif has_chinese and has_english:
                            # 检查是否是双语对照格式
                            lines = last_text.split('\n')
                            chinese_lines = [line for line in lines if contains_chinese(line)]
                            english_lines = [line for line in lines if contains_english(line) and not contains_chinese(line)]
                            logger.info(f"  最后一个单元格 - 中文行数: {len(chinese_lines)}, 纯英文行数: {len(english_lines)}")

                            if len(chinese_lines) > len(english_lines):
                                logger.warning(f"  ⚠️ 最后一个单元格可能翻译不完整（中文行多于英文行）")
                                logger.info(f"  中文内容: {chinese_lines}")
                                logger.info(f"  英文内容: {english_lines}")
                else:
                    logger.info(f"  最后一行单元格数不足2个")
                    
    except Exception as e:
        logger.error(f"检查文档失败: {str(e)}")
        import traceback
        traceback.print_exc()

def main():
    # 检查指定的文档
    output_dir = "输出"
    doc_name = "单晶电阻率管控技术标准_带翻译_20250610230213.docx"
    doc_path = os.path.join(output_dir, doc_name)

    if not os.path.exists(doc_path):
        logger.error(f"指定文档不存在: {doc_path}")
        # 查找最新的翻译文档作为备选
        import glob
        pattern = os.path.join(output_dir, "单晶电阻率管控技术标准_带翻译_*.docx")
        files = glob.glob(pattern)

        if files:
            latest_file = max(files, key=os.path.getmtime)
            logger.info(f"使用最新文档: {latest_file}")
            doc_path = latest_file
        else:
            logger.error("未找到任何翻译文档")
            return

    logger.info(f"检查文档: {doc_path}")

    # 检查文档修改时间
    import time
    mtime = os.path.getmtime(doc_path)
    mtime_str = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(mtime))
    logger.info(f"文档修改时间: {mtime_str}")

    check_document_tables(doc_path)

if __name__ == "__main__":
    main()
