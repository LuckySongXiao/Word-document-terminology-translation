#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试合并单元格修复效果
"""

import os
import sys
import json
import logging
from services.document_processor import DocumentProcessor
from services.zhipuai_translator import ZhipuAITranslator
from services.ollama_translator import OllamaTranslator

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def load_api_config():
    """加载API配置"""
    try:
        with open('API_config/zhipu_api.json', 'r', encoding='utf-8') as f:
            zhipu_config = json.load(f)
        return zhipu_config
    except Exception as e:
        logger.error(f"加载API配置失败: {e}")
        return None

def get_translator():
    """获取翻译器"""
    # 尝试使用智谱AI
    config = load_api_config()
    if config and config.get('api_key'):
        try:
            translator = ZhipuAITranslator(config['api_key'])
            # 测试连接
            test_result = translator.translate("测试", {}, "zh", "en")
            if test_result and test_result != "测试":
                logger.info("使用智谱AI翻译器")
                return translator
        except Exception as e:
            logger.warning(f"智谱AI翻译器初始化失败: {e}")
    
    # 回退到本地Ollama
    try:
        # 从配置文件加载Ollama配置
        with open('config.json', 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        ollama_config = config.get("fallback_translator", {})
        model = ollama_config.get("model", "qwen2.5:7b")
        api_url = ollama_config.get("api_url", "http://localhost:11434")
        
        translator = OllamaTranslator(
            model=model,
            api_url=api_url,
            model_list_timeout=10,
            translate_timeout=60
        )
        logger.info("使用本地Ollama翻译器")
        return translator
    except Exception as e:
        logger.error(f"Ollama翻译器初始化失败: {e}")
        return None

def test_merged_cell_fix():
    """测试合并单元格修复效果"""
    
    logger.info("=== 测试合并单元格修复效果 ===")
    
    # 检查原始文档是否存在
    input_file = "uploads/单晶电阻率管控技术标准.docx"
    if not os.path.exists(input_file):
        logger.error(f"原始文档不存在: {input_file}")
        return False
    
    try:
        translator = get_translator()
        if not translator:
            logger.error("无法初始化翻译器")
            return False
        
        # 创建文档处理器
        processor = DocumentProcessor(translator)
        processor.source_lang = "zh"
        processor.target_lang = "en"
        processor.output_format = "bilingual"
        
        # 加载术语库
        terminology_file = "术语库.xlsx"
        if os.path.exists(terminology_file):
            terminology = processor.load_terminology(terminology_file)
            logger.info(f"加载术语库: {len(terminology)} 个术语")
        else:
            terminology = {}
            logger.warning("术语库文件不存在，使用空术语库")
        
        # 翻译文档
        logger.info("开始翻译文档（使用合并单元格修复）...")
        output_path = processor.process_document(input_file, terminology)
        
        if output_path and os.path.exists(output_path):
            logger.info(f"✅ 文档翻译成功: {output_path}")
            
            # 验证修复效果
            return verify_merged_cell_fix(output_path)
        else:
            logger.error("❌ 文档翻译失败")
            return False
            
    except Exception as e:
        logger.error(f"文档翻译测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def verify_merged_cell_fix(doc_path):
    """验证合并单元格修复效果"""
    
    logger.info(f"验证修复效果: {doc_path}")
    
    try:
        from docx import Document
        doc = Document(doc_path)
        
        if len(doc.tables) < 4:
            logger.error("文档中表格数量不足")
            return False
        
        table = doc.tables[3]  # 表格4
        logger.info(f"检查表格4，行数: {len(table.rows)}")
        
        # 检查是否还有重复翻译问题
        duplicate_rows = []
        
        for row_idx, row in enumerate(table.rows):
            if len(row.cells) > 1:
                # 检查这一行是否所有单元格都包含相同内容
                first_cell_text = row.cells[0].text.strip()
                
                if first_cell_text and len(first_cell_text) > 20:  # 只检查有实质内容的单元格
                    all_same = True
                    for cell in row.cells[1:]:
                        if cell.text.strip() != first_cell_text:
                            all_same = False
                            break
                    
                    if all_same:
                        duplicate_rows.append(row_idx)
                        logger.warning(f"第{row_idx+1}行仍然存在重复翻译")
        
        if duplicate_rows:
            logger.error(f"❌ 修复失败！仍有 {len(duplicate_rows)} 行存在重复翻译: {[r+1 for r in duplicate_rows]}")
            return False
        else:
            logger.info("✅ 修复成功！未发现重复翻译问题")
            
            # 检查翻译完整性
            return check_translation_completeness(table)
        
    except Exception as e:
        logger.error(f"验证修复效果时出错: {e}")
        return False

def check_translation_completeness(table):
    """检查翻译完整性"""
    
    logger.info("检查翻译完整性...")
    
    import re
    
    incomplete_cells = []
    
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            
            if cell_text and len(cell_text) > 10:  # 只检查有实质内容的单元格
                # 检查是否包含中文
                has_chinese = bool(re.search(r'[\u4e00-\u9fff]', cell_text))
                # 检查是否包含英文
                has_english = bool(re.search(r'[A-Za-z]', cell_text))
                
                if has_chinese and not has_english:
                    # 只有中文，没有英文翻译
                    incomplete_cells.append((row_idx + 1, cell_idx + 1, cell_text[:50]))
                    logger.warning(f"单元格 [{row_idx + 1}, {cell_idx + 1}] 缺少英文翻译: {cell_text[:50]}...")
    
    if incomplete_cells:
        logger.error(f"❌ 发现 {len(incomplete_cells)} 个单元格缺少翻译")
        for row, col, text in incomplete_cells:
            logger.error(f"  单元格 [{row}, {col}]: {text}...")
        return False
    else:
        logger.info("✅ 所有单元格翻译完整")
        return True

def main():
    logger.info("开始测试合并单元格修复效果...")
    
    # 测试修复效果
    success = test_merged_cell_fix()
    
    logger.info("\n=== 测试结果总结 ===")
    if success:
        logger.info("🎉 合并单元格修复测试通过！")
        logger.info("✅ 重复翻译问题已解决")
        logger.info("✅ 翻译完整性良好")
        return True
    else:
        logger.error("💥 合并单元格修复测试失败")
        logger.error("需要进一步调试和修复")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
